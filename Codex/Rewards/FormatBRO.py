import os, sys, re, time, queue, shutil, threading
from tkinter import filedialog, messagebox, ttk
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from tkinter import font as tkfont
from send2trash import send2trash
from docx.oxml import OxmlElement
from zipfile import BadZipFile
from datetime import datetime
from docx.oxml.ns import qn
from docx import Document
from pathlib import Path
import tkinter as tk


def install_frozen_executable_icon(root, retry_ms: int = 250) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    _set_frozen_executable_icon(root)
    try:
        root.after(retry_ms, lambda: _set_frozen_executable_icon(root))
    except Exception:
        pass


def _set_frozen_executable_icon(root) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    try:
        import ctypes
        from ctypes import wintypes

        wm_seticon = 0x0080
        icon_small = 0
        icon_big = 1
        gclp_hicon = -14
        gclp_hiconsm = -34

        shell32 = ctypes.WinDLL("shell32", use_last_error=True)
        user32 = ctypes.WinDLL("user32", use_last_error=True)
        long_ptr = (
            ctypes.c_longlong if ctypes.sizeof(ctypes.c_void_p) == 8 else ctypes.c_long
        )

        shell32.ExtractIconExW.argtypes = (
            wintypes.LPCWSTR,
            ctypes.c_int,
            ctypes.POINTER(wintypes.HICON),
            ctypes.POINTER(wintypes.HICON),
            wintypes.UINT,
        )
        shell32.ExtractIconExW.restype = wintypes.UINT
        user32.SendMessageW.argtypes = (
            wintypes.HWND,
            wintypes.UINT,
            wintypes.WPARAM,
            wintypes.LPARAM,
        )
        user32.SendMessageW.restype = long_ptr
        user32.GetParent.argtypes = (wintypes.HWND,)
        user32.GetParent.restype = wintypes.HWND

        try:
            set_class_icon = user32.SetClassLongPtrW
        except AttributeError:
            set_class_icon = user32.SetClassLongW
        set_class_icon.argtypes = (wintypes.HWND, ctypes.c_int, long_ptr)
        set_class_icon.restype = long_ptr

        icon_handles = getattr(root, "_frozen_executable_icon_handles", None)
        if icon_handles is None:
            large_icons = (wintypes.HICON * 1)()
            small_icons = (wintypes.HICON * 1)()
            icon_count = shell32.ExtractIconExW(
                sys.executable,
                0,
                large_icons,
                small_icons,
                1,
            )
            if icon_count <= 0:
                return
            icon_handles = (small_icons[0], large_icons[0])
            root._frozen_executable_icon_handles = icon_handles
        small_icon, large_icon = icon_handles

        root.update_idletasks()
        hwnds = []
        for raw_hwnd in (root.winfo_id(), root.wm_frame()):
            try:
                hwnd = int(str(raw_hwnd), 0)
            except (TypeError, ValueError):
                continue
            while hwnd and hwnd not in hwnds:
                hwnds.append(hwnd)
                hwnd = user32.GetParent(hwnd)

        for hwnd in hwnds:
            if small_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_small, small_icon)
                set_class_icon(hwnd, gclp_hiconsm, small_icon)
            if large_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_big, large_icon)
                set_class_icon(hwnd, gclp_hicon, large_icon)
    except Exception:
        pass


def install_dark_title_bar(window, retry_ms: int = 80) -> None:
    if sys.platform != "win32":
        return

    def _apply() -> None:
        try:
            import ctypes

            hwnd = ctypes.windll.user32.GetParent(window.winfo_id())
            value = ctypes.c_int(1)
            for attr in (20, 19):
                ctypes.windll.dwmapi.DwmSetWindowAttribute(
                    hwnd,
                    attr,
                    ctypes.byref(value),
                    ctypes.sizeof(value),
                )
        except Exception:
            pass

    try:
        window.update_idletasks()
        _apply()
        window.after(retry_ms, _apply)
    except Exception:
        pass

# pywin32 для конвертації .doc -> .docx
try:
    import win32com.client
except ImportError:
    win32com = None

try:
    import pythoncom
except ImportError:
    pythoncom = None

# ================= ГЛОБАЛЬНІ КОНФІГ-ЗМІННІ =================

BASE_DIR: Path | None = None
ROOT_DIR: Path | None = None
BACKUP_ROOT: Path | None = None

SKIP_FILES = {"БР.docx", "БР.doc"}

SKIP_WORDS = {"приданий із", "приданий з", "прийнятий в оперативне підпорядкування з", "прийнятий в оперативне підпорядкування із"}

TEXT_TO_UPPER = [
    "3 мб", " мр:", " бро", "б/к", " мб ",
]

TEXT_TO_LOWER = [
    "Старший", "Молодший", "Сержант", "Солдат", "Головний",
]

TEXT_TO_BOLD = [
    "Виконував бойові (спеціальні) завдання в межах до ротного опорного пункту включно",
    "Виконував бойові (спеціальні) завдання,",
    "Виконував бойові (спеціальні) завдання згідно з бойовими наказами (розпорядженнями)",
]
TEXT_REPLACEMENTS = {
    "ВЗводу": "взводу", " ;": ";", " ,": ",", " .": ".", " мр:": " МР:", " :": ":", "—": "–", 
    "2026.": "2026", "військової частини А0998": " ", "ст.": "старший",
    "ТВО ": "Тимчасово виконуючий обов'язки ", " )": ")", "( ": "(",
}

TEXT_TO_ITALIC = [ "(на 70 тис. грн.)", "(на 100 тис. грн.)", "(на 30 тис. грн.)"]
ITALIC_PREFIXES_WITH_COLON = {
    "(на 100 тис. грн)",
    "(на 70 тис. грн)",
    "(на 30 тис. грн)",
}

PATTERN_DATA = re.compile(
    r"що\s+\d{2}\.\d{2}\.\d{4}\s+на\s+підставі",
    re.IGNORECASE,
)

REPORT_BODY_MARK = "Дійсним доповідаю"
DATE_PATTERN_RE = re.compile(r"(\d{2}\.\d{2}\.\d{4})")
RANKS = {"сержант", "солдат", "лейтенант"}

# ================= ЛОГУВАННЯ =================

LOG_LINES = []

def log(msg: str):
    print(msg)
    LOG_LINES.append(str(msg))


def configure_launch_styles(root):
    colors = {
        "window": "#F4F7FB",
        "panel": "#FFFFFF",
        "border": "#D6DEE8",
        "text": "#15202B",
        "muted": "#617084",
        "accent": "#0F766E",
        "accent_dark": "#0B5D56",
        "header": "#103C3A",
        "header_badge": "#0F766E",
    }

    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except Exception:
        pass

    try:
        style.configure("LaunchRoot.TFrame", background=colors["window"])
        style.configure("LaunchPanel.TFrame", background=colors["panel"])
        style.configure("LaunchSection.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        style.configure("LaunchField.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10, "bold"))
        style.configure("LaunchMuted.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("LaunchStatus.TLabel", background=colors["window"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("ProgressHeader.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        style.configure("ProgressBody.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10))
        style.configure("ProgressFile.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("ProgressCount.TLabel", background=colors["panel"], foreground=colors["accent_dark"], font=("Segoe UI", 10, "bold"))
        style.configure(
            "Launch.TEntry",
            fieldbackground="#FFFFFF",
            foreground=colors["text"],
            bordercolor=colors["border"],
            lightcolor=colors["border"],
            darkcolor=colors["border"],
            padding=(8, 6),
        )
        style.map(
            "Launch.TEntry",
            fieldbackground=[("readonly", "#FFFFFF")],
            foreground=[("readonly", colors["text"])],
        )
        style.configure(
            "LaunchBrowse.TButton",
            font=("Segoe UI", 10, "bold"),
            padding=(18, 10),
            foreground="#FFFFFF",
            background="#053D39",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchBrowse.TButton",
            background=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
        )
        style.configure(
            "LaunchPrimary.TButton",
            font=("Segoe UI", 11, "bold"),
            padding=(20, 14),
            foreground="#FFFFFF",
            background="#053D39",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.configure(
            "LaunchPrimaryHover.TButton",
            font=("Segoe UI", 11, "bold"),
            padding=(20, 14),
            foreground="#FFFFFF",
            background=colors["accent_dark"],
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchPrimary.TButton",
            background=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
            bordercolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
        )
        style.map(
            "LaunchPrimaryHover.TButton",
            background=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
            bordercolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
        )
        style.configure(
            "Launch.Horizontal.TProgressbar",
            background=colors["header_badge"],
            troughcolor="#E8EEF5",
            bordercolor=colors["border"],
            lightcolor=colors["header_badge"],
            darkcolor=colors["header_badge"],
            thickness=12,
        )
    except Exception:
        pass

    return colors


def _widget_exists(widget) -> bool:
    try:
        return bool(widget.winfo_exists())
    except Exception:
        return False


def _prepare_dialog_parent(root) -> None:
    if not _widget_exists(root):
        return
    try:
        root.lift()
        root.attributes("-topmost", True)
        root.update()
    except Exception:
        pass


def format_path_for_entry_display(path: str) -> str:
    path = path.strip()
    if not path:
        return ""

    normalized = Path(path)
    drive_prefix = ""
    if normalized.drive and len(normalized.drive) == 2 and normalized.drive[1] == ":":
        drive_prefix = f"{normalized.drive}/"

    name = normalized.name
    parent_name = normalized.parent.name
    if parent_name and name:
        return f"{drive_prefix}.../{parent_name}/{name}"
    if name:
        return f"{drive_prefix}{name}" if drive_prefix else f".../{name}"
    if drive_prefix:
        return drive_prefix
    return path.replace("\\", "/")


def fit_path_for_entry_display(path: str, entry) -> str:
    display_path = format_path_for_entry_display(path)
    if not display_path:
        return ""

    try:
        entry.update_idletasks()
        font_name = entry.cget("font") or "TkDefaultFont"
        try:
            font = tkfont.nametofont(font_name)
        except Exception:
            font = tkfont.nametofont("TkDefaultFont")

        available_width = max(120, entry.winfo_width() - 14)
        if font.measure(display_path) <= available_width:
            return display_path

        parts = display_path.split("/")
        if len(parts) >= 3:
            prefix = "/".join(parts[:-1]) + "/..."
            shortened_name = parts[-1]
            while len(shortened_name) > 1 and font.measure(prefix + shortened_name) > available_width:
                shortened_name = shortened_name[1:]
            return prefix + shortened_name

        ellipsis = "..."
        shortened_path = display_path
        while len(shortened_path) > 1 and font.measure(ellipsis + shortened_path) > available_width:
            shortened_path = shortened_path[1:]
        return ellipsis + shortened_path
    except Exception:
        return display_path


def refresh_path_entry_display(source_var, display_var, entry) -> None:
    if not _widget_exists(entry):
        return
    display_var.set(fit_path_for_entry_display(source_var.get(), entry))


def bind_primary_button_hover(button) -> None:
    def on_enter(_event) -> None:
        button.configure(style="LaunchPrimaryHover.TButton")
        try:
            button.focus_force()
        except Exception:
            pass

    def on_leave(_event) -> None:
        button.configure(style="LaunchPrimary.TButton")
        try:
            button.winfo_toplevel().focus_force()
        except Exception:
            pass

    button.bind("<Enter>", on_enter, add="+")
    button.bind("<Leave>", on_leave, add="+")


def choose_folder(root, source_var, display_var, entry, title: str) -> None:
    _prepare_dialog_parent(root)
    path = filedialog.askdirectory(title=title, parent=root)
    if path:
        source_var.set(path)
        refresh_path_entry_display(source_var, display_var, entry)
    _prepare_dialog_parent(root)


def ask_folder_settings_window(
    parent,
    *,
    app_name: str,
    subtitle: str,
    hint: str,
    browse_title: str,
) -> Path | None:
    dialog = parent
    created_dialog = False
    if dialog is None:
        dialog = tk.Tk()
        created_dialog = True

    selected = {"path": None}
    dialog.withdraw()
    install_frozen_executable_icon(dialog)
    dialog.title(f"НАЛАШТУВАННЯ - FormatBRO:")
    dialog.attributes("-topmost", True)
    dialog.resizable(False, False)

    colors = configure_launch_styles(dialog)
    dialog.configure(bg=colors["window"])
    install_dark_title_bar(dialog)

    path_var = tk.StringVar(value="")
    path_display_var = tk.StringVar(value="")
    done_var = tk.BooleanVar(master=dialog, value=False)
    start_button = None

    def update_launch_state(*_args) -> None:
        pass

    def on_start() -> None:
        folder_text = path_var.get().strip()
        if not folder_text:
            _prepare_dialog_parent(dialog)
            messagebox.showerror("Помилка", "Потрібно вказати папку для обробки.", parent=dialog)
            return

        folder = Path(folder_text).expanduser().resolve()
        if not folder.exists() or not folder.is_dir():
            _prepare_dialog_parent(dialog)
            messagebox.showerror("Помилка", f"Це не папка:\n{folder}", parent=dialog)
            return

        selected["path"] = folder
        dialog.withdraw()
        done_var.set(True)

    def on_cancel() -> None:
        selected["path"] = None
        dialog.withdraw()
        done_var.set(True)

    dialog.protocol("WM_DELETE_WINDOW", on_cancel)
    dialog.bind("<Escape>", lambda _event: on_cancel())
    dialog.bind("<Return>", lambda _event: on_start())

    surface = ttk.Frame(dialog, style="LaunchRoot.TFrame")
    surface.grid(row=0, column=0, sticky="nsew")
    surface.grid_columnconfigure(0, weight=1)

    header = tk.Frame(surface, bg=colors["header"], padx=22, pady=12)
    header.grid(row=0, column=0, sticky="we")
    header.grid_columnconfigure(0, weight=1)
    tk.Label(
        header,
        text=app_name,
        bg=colors["header"],
        fg="#FFFFFF",
        font=("Segoe UI Semibold", 16),
        anchor="w",
    ).grid(row=0, column=0, sticky="w")
    tk.Label(
        header,
        text=subtitle,
        bg=colors["header"],
        fg="#D7FBF5",
        font=("Segoe UI", 9),
        anchor="w",
    ).grid(row=1, column=0, sticky="w", pady=(2, 6))
    badge_shell = tk.Frame(header, bg=colors["accent_dark"], padx=1, pady=1)
    badge = tk.Label(
        badge_shell,
        text="НАЛАШТУВАННЯ",
        bg="#053D39",
        fg="#FFFFFF",
        font=("Segoe UI", 12, "bold"),
        padx=10,
        pady=5,
    )
    badge.pack()
    badge_shell.place(relx=1.0, x=5, y=5, anchor="ne")

    body = ttk.Frame(surface, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
    body.grid(row=1, column=0, sticky="nsew")
    body.grid_columnconfigure(0, weight=1)

    panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
    panel_shell.grid(row=0, column=0, sticky="we")
    settings = ttk.Frame(panel_shell, style="LaunchPanel.TFrame", padding=(16, 14, 16, 12))
    settings.pack(fill="both", expand=True)
    settings.grid_columnconfigure(1, weight=1)

    ttk.Label(settings, text="Папка для обробки", style="LaunchSection.TLabel").grid(
        row=0, column=0, columnspan=3, sticky="w"
    )
    ttk.Label(settings, text=hint, style="LaunchMuted.TLabel", wraplength=520).grid(
        row=1, column=0, columnspan=3, pady=(4, 5), sticky="w"
    )
    ttk.Label(settings, text="ПАПКА:", style="LaunchField.TLabel").grid(
        row=2, column=0, padx=(0, 8), pady=7, sticky="e"
    )
    path_entry = ttk.Entry(settings, textvariable=path_display_var, width=48, state="readonly", style="Launch.TEntry")
    path_entry.grid(row=2, column=1, padx=(0, 14), pady=7, sticky="we")
    ttk.Button(
        settings,
        text="ОБРАТИ",
        width=14,
        style="LaunchBrowse.TButton",
        command=lambda: choose_folder(dialog, path_var, path_display_var, path_entry, browse_title),
    ).grid(row=2, column=2, pady=7, sticky="nsew")

    footer = ttk.Frame(body, style="LaunchRoot.TFrame")
    footer.grid(row=1, column=0, pady=(10, 0), sticky="we")
    footer.grid_columnconfigure(0, weight=1)
    start_button = ttk.Button(
        footer,
        text="ЗАПУСТИТИ",
        style="LaunchPrimary.TButton",
        command=on_start,
        padding=(18, 18),
    )
    start_button.grid(row=2, column=0, sticky="we")
    bind_primary_button_hover(start_button)

    path_var.trace_add("write", lambda *_args: (refresh_path_entry_display(path_var, path_display_var, path_entry), update_launch_state()))
    path_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(path_var, path_display_var, path_entry))
    update_launch_state()

    dialog.update_idletasks()
    dialog.geometry(f"{dialog.winfo_reqwidth()}x{dialog.winfo_reqheight()}")
    dialog.deiconify()
    dialog.lift()
    dialog.grab_set()
    dialog.focus_force()
    dialog.wait_variable(done_var)

    try:
        dialog.grab_release()
    except Exception:
        pass

    if created_dialog and _widget_exists(dialog):
        dialog.destroy()

    return selected["path"]


class ProgressWindow:
    def __init__(self, owner: tk.Tk):
        self.dialog = tk.Toplevel(owner)
        self.dialog.withdraw()
        install_frozen_executable_icon(self.dialog)
        self.dialog.title("FormatBRO - Обробка")
        self.dialog.resizable(False, False)
        self.dialog.attributes("-topmost", True)
        colors = configure_launch_styles(self.dialog)
        self.dialog.configure(bg=colors["window"])
        install_dark_title_bar(self.dialog)

        container = ttk.Frame(self.dialog, style="LaunchRoot.TFrame")
        container.pack(fill="both", expand=True)

        header = tk.Frame(container, bg=colors["header"], padx=22, pady=10)
        header.pack(fill="x")
        header.grid_columnconfigure(0, weight=1)
        tk.Label(
            header,
            text="FormatBRO",
            bg=colors["header"],
            fg="#FFFFFF",
            font=("Segoe UI Semibold", 15),
            anchor="w",
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            header,
            text="Патч документів та чистка форматування виконуються.",
            bg=colors["header"],
            fg="#D7FBF5",
            font=("Segoe UI", 9),
            anchor="w",
        ).grid(row=1, column=0, sticky="w", pady=(3, 4))
        badge_shell = tk.Frame(header, bg=colors["accent_dark"], padx=1, pady=1)
        badge = tk.Label(
            badge_shell,
            text="ОБРОБКА",
            bg="#053D39",
            fg="#FFFFFF",
            font=("Segoe UI", 12, "bold"),
            padx=10,
            pady=5,
        )
        badge.pack()
        badge_shell.place(relx=1.0, x=5, y=5, anchor="ne")

        body = ttk.Frame(container, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
        body.pack(fill="both", expand=True)
        panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
        panel_shell.pack(fill="both", expand=True)
        frame = ttk.Frame(panel_shell, style="LaunchPanel.TFrame", padding=(16, 14, 16, 12))
        frame.pack(fill="both", expand=True)
        frame.grid_columnconfigure(0, weight=1)

        self.header_var = tk.StringVar(value="Підготовка...")
        self.detail_var = tk.StringVar(value="Будь ласка, зачекайте")
        self.file_var = tk.StringVar(value="")
        self.count_var = tk.StringVar(value="0 / 0")

        ttk.Label(frame, textvariable=self.header_var, style="ProgressHeader.TLabel").grid(
            row=0, column=0, padx=(0, 12), sticky="w"
        )
        ttk.Label(frame, textvariable=self.count_var, style="ProgressCount.TLabel").grid(
            row=0, column=1, sticky="e"
        )
        ttk.Label(frame, textvariable=self.detail_var, style="ProgressBody.TLabel", wraplength=500).grid(
            row=1, column=0, columnspan=2, pady=(8, 6), sticky="w"
        )
        ttk.Label(frame, textvariable=self.file_var, style="ProgressFile.TLabel", wraplength=500).grid(
            row=2, column=0, columnspan=2, pady=(0, 12), sticky="w"
        )
        self.progress = ttk.Progressbar(
            frame,
            orient="horizontal",
            mode="determinate",
            maximum=100,
            style="Launch.Horizontal.TProgressbar",
        )
        self.progress.grid(row=3, column=0, columnspan=2, sticky="we")
        self.dialog.update_idletasks()
        width = max(560, self.dialog.winfo_reqwidth())
        height = max(248, self.dialog.winfo_reqheight())
        self.dialog.geometry(f"{width}x{height}")
        self.dialog.deiconify()
        self.dialog.lift()
        self.refresh()

    def update(self, header=None, detail=None, current=None, total=None, file_name=None):
        try:
            if not self.dialog.winfo_exists():
                return
            if header is not None:
                self.header_var.set(header)
            if detail is not None:
                self.detail_var.set(detail)
            if file_name is not None:
                self.file_var.set(file_name)
            if current is not None and total is not None:
                safe_total = total if total > 0 else 1
                pct = max(0.0, min(100.0, (current / safe_total) * 100.0))
                self.progress["value"] = pct
                self.count_var.set(f"{current} / {total}")
            self.refresh()
        except tk.TclError:
            return

    def refresh(self):
        try:
            self.dialog.update_idletasks()
            self.dialog.update()
        except Exception:
            pass

    def close(self):
        try:
            self.dialog.destroy()
        except Exception:
            pass


class ProgressUpdateProxy:
    def __init__(self, updates):
        self.updates = updates

    def update(self, header=None, detail=None, current=None, total=None, file_name=None):
        self.updates.put(
            {
                "header": header,
                "detail": detail,
                "current": current,
                "total": total,
                "file_name": file_name,
            }
        )


def apply_pending_progress_updates(progress: ProgressWindow, updates) -> None:
    while True:
        try:
            payload = updates.get_nowait()
        except queue.Empty:
            return
        progress.update(**payload)


def run_format_workflow_with_progress(progress: ProgressWindow) -> None:
    updates = queue.Queue()
    result = {"error": None}
    progress_proxy = ProgressUpdateProxy(updates)

    def worker() -> None:
        global BASE_DIR, ROOT_DIR, BACKUP_ROOT
        com_initialized = False
        try:
            if pythoncom is not None:
                pythoncom.CoInitialize()
                com_initialized = True
            make_backup(BASE_DIR, progress=progress_proxy)

            try:
                patch_reports(progress=progress_proxy)
            except Exception as e:
                log(f"[PATCH ERROR] Патч документів зламався, але чистка продовжується: {e}")

            try:
                convert_all_docs_to_docx(
                    ROOT_DIR,
                    progress=progress_proxy,
                    stage_header="Етап 3.5/4: Перепровіряємо...",
                )
            except Exception as e:
                log(f"[DOC->DOCX ERROR] Конвертація дала збій, продовжую як є: {e}")

            clean_process_all_docx(ROOT_DIR, progress=progress_proxy)
        except Exception as exc:
            result["error"] = exc
        finally:
            if com_initialized:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()

    while thread.is_alive():
        apply_pending_progress_updates(progress, updates)
        progress.refresh()
        time.sleep(0.03)

    thread.join()
    apply_pending_progress_updates(progress, updates)
    progress.refresh()

    if result["error"] is not None:
        raise result["error"]

# =================================================
#              SAFE XML REMOVE HELPERS
# =================================================

def _safe_remove_from_parent(node) -> bool:
    try:
        parent = node.getparent()
        if parent is None:
            return False
        if node not in list(parent):
            return False
        parent.remove(node)
        return True
    except Exception:
        return False


def delete_paragraph(paragraph):
    try:
        p = paragraph._element
        _safe_remove_from_parent(p)
    except Exception:
        return

def normalize_apostrophes(text: str) -> str:
    if not text:
        return text

    replacements = {
        "’": "'",
        "‘": "'",
        "ʼ": "'",
        "ʹ": "'",
        "＇": "'",
        "´": "'",
        "`": "'",
    }

    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def normalize_apostrophes_in_document(document):
    def fix_paragraphs(paragraphs):
        for p in paragraphs:
            if not p.text:
                continue
            new = normalize_apostrophes(p.text)
            if new != p.text:
                p.text = new

    fix_paragraphs(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                fix_paragraphs(cell.paragraphs)

DATE_START_RE = re.compile(r"^\s*\d{2}\.\d{2}\.\d{4}\b")
DATE_ONLY_RE  = re.compile(r"^\s*\d{2}\.\d{2}\.\d{4}\s*$")


def is_empty_paragraph(p) -> bool:
    return not ((p.text or "").strip())


def is_date_paragraph(p) -> bool:
    t = p.text or ""
    return bool(DATE_START_RE.match(t))


def insert_empty_paragraph_before(paragraph, count: int = 1):
    p = paragraph._p
    for _ in range(count):
        new_p = OxmlElement("w:p")
        p.addprevious(new_p)


def normalize_date_block_spacing(document, empty_before_count: int = 2):
    # збираємо індекси "датних" абзаців
    paras = list(document.paragraphs)
    date_idxs = [i for i, p in enumerate(paras) if is_date_paragraph(p)]
    if not date_idxs:
        return

    # обробляємо знизу вгору (щоб видалення не ламало наступні індекси)
    for date_idx in reversed(date_idxs):
        paras = list(document.paragraphs)
        if date_idx >= len(paras):
            continue

        date_p = paras[date_idx]
        date_elm = date_p._p

        # знайти 2 попередні НЕпорожні абзаци
        p2 = None
        p1 = None

        i = date_idx - 1
        while i >= 0:
            if not is_empty_paragraph(paras[i]):
                p2 = paras[i]
                break
            i -= 1

        if p2 is None:
            continue

        i = (paras.index(p2)) - 1
        while i >= 0:
            if not is_empty_paragraph(paras[i]):
                p1 = paras[i]
                break
            i -= 1

        if p1 is None:
            continue

        p1_elm = p1._p
        p2_elm = p2._p

        # refresh + мапа element -> paragraph
        paras = list(document.paragraphs)
        elm2para = {p._p: p for p in paras}

        # якщо якісь параграфи вже "поїхали", пробуємо знайти по елементам
        if p1_elm not in elm2para or p2_elm not in elm2para or date_elm not in elm2para:
            continue

        p1 = elm2para[p1_elm]
        p2 = elm2para[p2_elm]
        date_p = elm2para[date_elm]

        # --- (1) видалити порожні абзаци між p1, p2, date_p
        paras = list(document.paragraphs)
        idx_p1 = next((i for i, p in enumerate(paras) if p._p is p1_elm), None)
        idx_p2 = next((i for i, p in enumerate(paras) if p._p is p2_elm), None)
        idx_dt = next((i for i, p in enumerate(paras) if p._p is date_elm), None)
        if None in (idx_p1, idx_p2, idx_dt):
            continue

        # між p1 і p2
        for j in range(idx_p2 - 1, idx_p1, -1):
            if is_empty_paragraph(paras[j]):
                delete_paragraph(paras[j])

        # refresh
        paras = list(document.paragraphs)
        idx_p1 = next((i for i, p in enumerate(paras) if p._p is p1_elm), None)
        idx_p2 = next((i for i, p in enumerate(paras) if p._p is p2_elm), None)
        idx_dt = next((i for i, p in enumerate(paras) if p._p is date_elm), None)
        if None in (idx_p1, idx_p2, idx_dt):
            continue

        # між p2 і date
        for j in range(idx_dt - 1, idx_p2, -1):
            if is_empty_paragraph(paras[j]):
                delete_paragraph(paras[j])

        # refresh
        paras = list(document.paragraphs)
        idx_p1 = next((i for i, p in enumerate(paras) if p._p is p1_elm), None)
        idx_dt = next((i for i, p in enumerate(paras) if p._p is date_elm), None)
        if None in (idx_p1, idx_dt):
            continue

        # --- (2) прибрати всі пусті абзаци над блоком і вставити рівно 2 пусті
        # видаляємо пусті безпосередньо перед p1 (поки не буде непорожній або початок)
        paras = list(document.paragraphs)
        idx_p1 = next((i for i, p in enumerate(paras) if p._p is p1_elm), None)
        if idx_p1 is None:
            continue

        k = idx_p1 - 1
        while k >= 0 and is_empty_paragraph(paras[k]):
            delete_paragraph(paras[k])
            paras = list(document.paragraphs)
            idx_p1 = next((i for i, p in enumerate(paras) if p._p is p1_elm), None)
            if idx_p1 is None:
                break
            k = idx_p1 - 1

        # вставляємо рівно empty_before_count пустих перед p1
        paras = list(document.paragraphs)
        elm2para = {p._p: p for p in paras}
        if p1_elm not in elm2para:
            continue
        insert_empty_paragraph_before(elm2para[p1_elm], count=empty_before_count)

        # --- (3) прибрати всі пусті абзаци після date_p
        paras = list(document.paragraphs)
        idx_dt = next((i for i, p in enumerate(paras) if p._p is date_elm), None)
        if idx_dt is None:
            continue

        k = idx_dt + 1
        while k < len(paras) and is_empty_paragraph(paras[k]):
            delete_paragraph(paras[k])
            paras = list(document.paragraphs)
            idx_dt = next((i for i, p in enumerate(paras) if p._p is date_elm), None)
            if idx_dt is None:
                break
            k = idx_dt + 1

        # --- (4) відступи = 0 для p1, p2, date_p
        paras = list(document.paragraphs)
        elm2para = {p._p: p for p in paras}
        for elm in (p1_elm, p2_elm, date_elm):
            p = elm2para.get(elm)
            if not p:
                continue
            pf = p.paragraph_format
            pf.left_indent = Cm(0)
            pf.first_line_indent = Cm(0)
            pf.hanging_indent = None

# =================================================
#              ЧАСТИНА 0: BACKUP ПАПКИ
# =================================================

def make_backup(base_dir: Path, progress: ProgressWindow | None = None) -> Path:
    global BACKUP_ROOT
    backup_dest = BACKUP_ROOT

    log(f"[BACKUP] Створюю резервну копію у: {backup_dest}")

    try:
        backup_dest.mkdir(parents=True, exist_ok=True)

        items = list(base_dir.rglob("*"))
        if progress:
            progress.update(
                header="Етап 1/4: Backup",
                detail="Створення резервної копії...",
                current=0,
                total=len(items),
                file_name="",
            )

        for idx, item in enumerate(items, 1):
            if BACKUP_ROOT in item.parents or item == BACKUP_ROOT:
                if progress:
                    progress.update(current=idx, total=len(items), file_name=item.name)
                continue

            rel = item.relative_to(base_dir)
            dest_path = backup_dest / rel

            if item.is_dir():
                dest_path.mkdir(parents=True, exist_ok=True)
            else:
                dest_path.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(item, dest_path)
            if progress:
                progress.update(current=idx, total=len(items), file_name=item.name)

    except Exception as e:
        log(f"[BACKUP ERROR] Не вдалось зробити backup: {e}")

    return backup_dest

def remove_all_page_breaks(doc: Document):
    for para in doc.paragraphs:
        for run in para.runs:
            r_elm = run._element
            br_elems = list(r_elm.findall(".//w:br", r_elm.nsmap))
            for br in br_elems:
                br_type = br.get(qn("w:type"))
                if br_type == "page":
                    _safe_remove_from_parent(br)

def paragraph_text_from_elm(p_elm) -> str:
    texts = []
    for node in p_elm.iter():
        if node.tag == qn("w:t"):
            if node.text:
                texts.append(node.text)
    return "".join(texts)

def is_empty_paragraph_elm(p_elm) -> bool:
    if p_elm.tag != qn("w:p"):
        return False
    text = paragraph_text_from_elm(p_elm).strip()
    return text == ""

def cleanup_empty_paragraphs_before_breaks(doc: Document):
    for para in list(doc.paragraphs):
        p_elm = para._p

        has_page_break = False
        for node in p_elm.iter():
            if node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
                has_page_break = True
                break

        if not has_page_break:
            continue

        prev = p_elm.getprevious()
        while prev is not None and prev.tag == qn("w:p") and is_empty_paragraph_elm(prev):
            parent = prev.getparent()
            if parent is not None and prev in list(parent):
                parent.remove(prev)
            prev = p_elm.getprevious()

def apply_times_new_roman_12(doc: Document):
    def fix_run(run):
        if not run.text:
            return
        try:
            f = run.font
            f.name = "Times New Roman"
            f.size = Pt(12)
        except Exception as e:
            log(f"[FONT WARN] Не вдалось змінити шрифт в одному з run: {e}")

    for para in doc.paragraphs:
        for run in para.runs:
            fix_run(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        fix_run(run)

def insert_page_break_before(paragraph):
    p = paragraph._p
    new_p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    new_p.append(r)
    p.addprevious(new_p)

def find_next_report_date(paragraphs, start_idx):
    for i in range(start_idx + 1, len(paragraphs)):
        m = DATE_PATTERN_RE.search(paragraphs[i].text or "")
        if m:
            return m.group(1)
    return None

def replace_report_date_in_paragraph_text(text: str, report_date: str) -> tuple[str, bool]:
    replacement = f"що {report_date} на підставі"
    new_text, cnt = PATTERN_DATA.subn(replacement, text, count=1)
    if cnt == 0:
        return text, False

    new_text = re.sub(r"[ \t\u00A0]+", " ", new_text)
    return new_text, True

def smart_replace_report_date_in_paragraph(para, report_date: str) -> bool:
    full_text = para.text or ""
    new_text, matched = replace_report_date_in_paragraph_text(full_text, report_date)
    if not matched:
        return False
    if new_text == full_text:
        return True

    if len(para.runs) == 1:
        para.runs[0].text = new_text
        return True

    for run in para.runs:
        t = run.text or ""
        m = PATTERN_DATA.search(t)
        if m:
            before = t[:m.start()]
            after = t[m.end():]
            patched = before + f"що {report_date} на підставі" + after
            if patched != t:
                run.text = patched
            return True

    para.text = new_text
    return True

def process_docx_patch(path: Path):
    try:
        doc = Document(path)
    except BadZipFile:
        log(f"[SKIP] {path} не є коректним DOCX (BadZipFile)")
        return
    except Exception as e:
        log(f"[ERROR] Не вдалось відкрити {path}: {e}")
        return

    remove_all_page_breaks(doc)

    paragraphs = doc.paragraphs
    first_komandiru_seen = False

    for idx, para in enumerate(paragraphs):
        text = para.text or ""

        if text.lstrip().startswith("Командиру 3"):
            if not first_komandiru_seen:
                first_komandiru_seen = True
            else:
                insert_page_break_before(para)

        if not text:
            continue

        if REPORT_BODY_MARK in text:
            report_date = find_next_report_date(paragraphs, idx)
            if not report_date:
                log(f"[WARN] Не знайшов дату рапорту після абзацу {idx} у {path.name}")
                continue

            if smart_replace_report_date_in_paragraph(para, report_date):
                log(f"[OK] Замінив дату у шапці рапорту на {report_date}")
            else:
                log(f"[WARN] Не знайшов PATTERN_DATA в абзаці з рапортом ({path.name})")

    cleanup_empty_paragraphs_before_breaks(doc)
    apply_times_new_roman_12(doc)

    try:
        doc.save(path)
    except Exception as e:
        log(f"[ERROR] Не вдалось зберегти {path}: {e}")

def convert_all_docs_to_docx(
    root_dir: Path,
    progress: ProgressWindow | None = None,
    stage_header: str = "Конвертація DOC -> DOCX",
):
    if win32com is None:
        log("[WARN] pywin32 не встановлено, конвертація .doc пропущена")
        if progress:
            progress.update(header=stage_header, detail="pywin32 не встановлено, етап пропущено", current=1, total=1)
        return

    doc_files = [p for p in root_dir.glob("*.doc") if p.is_file()]
    if progress:
        progress.update(
            header=stage_header,
            detail="Конвертація .doc у .docx...",
            current=0,
            total=len(doc_files),
            file_name="",
        )

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        for i, path in enumerate(doc_files, 1):
            if progress:
                progress.update(current=i - 1, total=len(doc_files), file_name=path.name)
            if path.name.startswith("~$"):
                log(f"[SKIP] {path} (тимчасовий файл Word)")
                if progress:
                    progress.update(current=i, total=len(doc_files), file_name=path.name)
                continue
            if path.name in SKIP_FILES:
                log(f"[SKIP] {path} (у SKIP_FILES)")
                if progress:
                    progress.update(current=i, total=len(doc_files), file_name=path.name)
                continue

            new_path = path.with_suffix(".docx")
            try:
                doc = word.Documents.Open(str(path))
                doc.SaveAs(str(new_path), FileFormat=16)  # wdFormatXMLDocument
                doc.Close()
                path.unlink()
                log(f"[OK] Конвертовано {path.name} -> {new_path.name} (старий .doc видалено)")
            except Exception as e:
                log(f"[ERROR] Конвертація {path}: {e}")
            if progress:
                progress.update(current=i, total=len(doc_files), file_name=path.name)
    finally:
        word.Quit()

def patch_reports(progress: ProgressWindow | None = None):
    global ROOT_DIR, BACKUP_ROOT

    if ROOT_DIR is None:
        raise RuntimeError("ROOT_DIR не задано.")
    if not ROOT_DIR.exists():
        raise FileNotFoundError(f"ROOT_DIR не існує: {ROOT_DIR}")

    convert_all_docs_to_docx(ROOT_DIR, progress=progress, stage_header="Етап 2/4: DOC -> DOCX")

    log(f"[PATCH] Обробка .docx у папці (без підпапок): {ROOT_DIR}")

    docx_files = [p for p in ROOT_DIR.glob("*.docx") if p.is_file()]
    if progress:
        progress.update(
            header="Етап 3/4: Patch DOCX",
            detail="Патч структури документів...",
            current=0,
            total=len(docx_files),
            file_name="",
        )

    for i, path in enumerate(docx_files, 1):
        if progress:
            progress.update(current=i - 1, total=len(docx_files), file_name=path.name)
        if BACKUP_ROOT in path.parents:
            if progress:
                progress.update(current=i, total=len(docx_files), file_name=path.name)
            continue
        if path.name.startswith("~$"):
            log(f"[SKIP] {path} (тимчасовий файл Word)")
            if progress:
                progress.update(current=i, total=len(docx_files), file_name=path.name)
            continue
        if path.name in SKIP_FILES:
            log(f"[SKIP] {path} (у SKIP_FILES)")
            if progress:
                progress.update(current=i, total=len(docx_files), file_name=path.name)
            continue

        log(f"[PATCH] {path.name}")
        process_docx_patch(path)
        if progress:
            progress.update(current=i, total=len(docx_files), file_name=path.name)

# =================================================
#      ЧАСТИНА 2: МАСОВА ЧИСТКА/ФОРМАТУВАННЯ
# =================================================

def italicize_matches_in_paragraph(paragraph, targets, ignore_case=True) -> bool:
    text = paragraph.text or ""
    if not text.strip() or not targets:
        return False

    # зберемо regex типу: (word1|word2|word3)
    escaped = [re.escape(t) for t in targets if t]
    if not escaped:
        return False

    flags = re.IGNORECASE if ignore_case else 0
    pattern = re.compile("|".join(escaped), flags)

    matches = list(pattern.finditer(text))
    if not matches:
        return False

    # Перебудовуємо абзац: звичайний текст + курсивні шматки
    paragraph.text = ""  # очищаємо всі runs

    pos = 0
    for m in matches:
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])  # звичайний
        r = paragraph.add_run(text[m.start():m.end()])  # збіг
        r.italic = True
        pos = m.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])

    return True

def apply_italic(document, to_italic=None, ignore_case=True):
    if to_italic is None:
        to_italic = []

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            italicize_matches_in_paragraph(p, to_italic, ignore_case=ignore_case)

    process_paragraphs(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

def clean_remove_page_breaks(document):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            r = run._r
            br_elems = r.findall(".//w:br", r.nsmap)
            for br in br_elems:
                br_type = br.get(qn('w:type'))
                if br_type == 'page':
                    _safe_remove_from_parent(br)

def _clear_shading_and_color_in_paragraph(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is not None:
        shd_elems = p_pr.findall(".//w:shd", p_pr.nsmap)
        for shd in shd_elems:
            _safe_remove_from_parent(shd)

    for run in paragraph.runs:
        r_pr = run._r.rPr
        if r_pr is not None:
            shd_elems = r_pr.findall(".//w:shd", r_pr.nsmap)
            for shd in shd_elems:
                _safe_remove_from_parent(shd)

            hl_elems = r_pr.findall(".//w:highlight", r_pr.nsmap)
            for hl in hl_elems:
                _safe_remove_from_parent(hl)

        font = run.font
        font.color.rgb = RGBColor(0, 0, 0)
        if hasattr(font.color, "theme_color"):
            font.color.theme_color = None
        if hasattr(font, "highlight_color"):
            font.highlight_color = None

def _clear_shading_highlight_in_element(root_elm):
    for tag in ("w:shd", "w:highlight"):
        for node in root_elm.xpath(f".//{tag}"):
            _safe_remove_from_parent(node)

def clear_shading_and_color(document):
    for paragraph in document.paragraphs:
        _clear_shading_and_color_in_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.tcPr
                if tc_pr is not None:
                    shd_elems = tc_pr.findall(".//w:shd", tc_pr.nsmap)
                    for shd in shd_elems:
                        _safe_remove_from_parent(shd)
                for paragraph in cell.paragraphs:
                    _clear_shading_and_color_in_paragraph(paragraph)

    # Видаляємо заливку/підсвічування з основного тіла документа.
    _clear_shading_highlight_in_element(document._element)

    # Також чистимо стилі: інколи заливка "сидить" у стилях.
    try:
        _clear_shading_highlight_in_element(document.styles._element)
    except Exception:
        pass

    # І заголовки/колонтитули, якщо вони є.
    try:
        for section in document.sections:
            for part in (
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
                section.even_page_header,
                section.even_page_footer,
            ):
                for paragraph in part.paragraphs:
                    _clear_shading_and_color_in_paragraph(paragraph)
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                _clear_shading_and_color_in_paragraph(paragraph)
                _clear_shading_highlight_in_element(part._element)
    except Exception:
        pass

def normalize_font(document, font_name="Times New Roman", font_size_pt=12):
    size = Pt(font_size_pt)

    try:
        normal_style = document.styles['Normal']
        normal_style.font.name = font_name
        normal_style.font.size = size
    except KeyError:
        pass

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = size

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = size

def fix_left_indents(document):
    def reset_paragraph(paragraph):
        paragraph.paragraph_format.left_indent = Cm(0)

    for paragraph in document.paragraphs:
        reset_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    reset_paragraph(paragraph)

def set_margins(document, left_cm=2.5, other_cm=1.5):
    for section in document.sections:
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(other_cm)
        section.top_margin = Cm(other_cm)
        section.bottom_margin = Cm(other_cm)

def justify_all_paragraphs(document):
    for paragraph in document.paragraphs:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def fix_space_before_parenthesis(document):
    pattern = re.compile(r"(\S)\(")

    def fix_paragraph(paragraph):
        original = paragraph.text
        fixed = pattern.sub(r"\1 (", original)
        if fixed != original:
            paragraph.text = fixed

    for paragraph in document.paragraphs:
        fix_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fix_paragraph(paragraph)

def normalize_commandiru_block(document):
    line1 = "Командиру 3 механізованого батальйону"

    def norm(s: str) -> str:
        return " ".join((s or "").replace("\n", " ").split())

    def looks_like_line1(text: str) -> bool:
        t = norm(text)
        # достатньо стійкий "якір" на line1
        return t.startswith("Командиру 3") and "механізован" in t and "батальйон" in t

    def normalize_in_paragraphs(paragraphs):
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            t1 = p.text or ""

            # якщо це звертання до командира — нормалізуємо ТІЛЬКИ до line1
            if looks_like_line1(t1):
                p.text = line1
                i += 1
                continue

            i += 1

    normalize_in_paragraphs(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                normalize_in_paragraphs(cell.paragraphs)

def set_last_four_nonempty_indent_rules(document, fourth_indent_cm: float = 1.0):
    paras = list(document.paragraphs)
    if not paras:
        return

    # зібрати індекси непорожніх абзаців знизу вверх
    nonempty_idxs = []
    for i in range(len(paras) - 1, -1, -1):
        if (paras[i].text or "").strip():
            nonempty_idxs.append(i)
            if len(nonempty_idxs) == 4:
                break

    if not nonempty_idxs:
        return

    # nonempty_idxs: [last, second_last, third_last, fourth_last]
    # 3 останні -> indent 0
    for i in nonempty_idxs[:3]:
        pf = paras[i].paragraph_format
        pf.first_line_indent = Cm(0)

    # 4-й знизу -> indent 1 см (якщо існує)
    if len(nonempty_idxs) >= 4:
        i4 = nonempty_idxs[3]
        pf = paras[i4].paragraph_format
        pf.first_line_indent = Cm(fourth_indent_cm)

def fix_last_daily_signature_line(document, spacer: str = "\t") -> bool:
    paras = list(document.paragraphs)
    nonempty = [p for p in paras if (p.text or "").strip()]
    if len(nonempty) < 2:
        return False

    p = nonempty[-2]  # передостанній непорожній
    text = (p.text or "").replace("\u00A0", " ").strip()

    # мінімальна перевірка: має бути хоча б 3 токени (посада/звання + ім'я + прізвище)
    tokens = text.split()
    if len(tokens) < 3:
        return False

    # беремо 2 останні токени як "ім'я прізвище"
    name_surname = " ".join(tokens[-2:])
    left = " ".join(tokens[:-2]).rstrip()

    # якщо вже є таб між лівою частиною і ПІБ, не чіпаємо
    if "\t" in (p.text or ""):
        return False

    # Перебудовуємо абзац через runs, щоб spacer точно зберігся
    p.text = ""
    p.add_run(left)
    p.add_run(spacer)
    p.add_run(name_surname)
    return True

def center_raport(document):
    def process_paragraphs(paragraphs):
        for p in paragraphs:
            stripped = (p.text or "").strip()
            if stripped.casefold() == "рапорт":
                p.text = "РАПОРТ"
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.left_indent = Cm(0)
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.hanging_indent = None
                if p.paragraph_format.tab_stops is not None:
                    try:
                        p.paragraph_format.tab_stops.clear_all()
                    except Exception:
                        pass

    process_paragraphs(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

def insert_page_break_before_paragraph(paragraph):
    paragraph.paragraph_format.page_break_before = True

def handle_commandiru_3(document):
    target_substr = "Командиру 3"

    while True:
        paragraphs = list(document.paragraphs)
        changed = False

        for i in range(len(paragraphs) - 1):
            cur = paragraphs[i]
            nxt = paragraphs[i + 1]

            if (cur.text or "").strip() == "" and target_substr in (nxt.text or ""):
                delete_paragraph(cur)
                changed = True
                break

        if not changed:
            break

    paragraphs = list(document.paragraphs)
    for p in paragraphs:
        if target_substr in (p.text or ""):
            insert_page_break_before_paragraph(p)

def remove_trailing_empty_paragraphs(document):
    while document.paragraphs:
        last = document.paragraphs[-1]
        if (last.text or "").strip() == "":
            delete_paragraph(last)
        else:
            break

def apply_text_transforms(document, to_upper=None, to_lower=None, replacements=None):
    if to_upper is None:
        to_upper = []
    if to_lower is None:
        to_lower = []
    if replacements is None:
        replacements = {}

    def transform_paragraph(paragraph):
        t = paragraph.text
        if not t:
            return
        new_t = t

        for old, new in replacements.items():
            if old in new_t:
                new_t = new_t.replace(old, new)

        for pattern in to_upper:
            if pattern in new_t:
                new_t = new_t.replace(pattern, pattern.upper())

        for pattern in to_lower:
            if pattern and re.search(re.escape(pattern), new_t, flags=re.IGNORECASE):
                new_t = re.sub(re.escape(pattern), pattern.lower(), new_t, flags=re.IGNORECASE)

        if new_t != t:
            paragraph.text = new_t

    for paragraph in document.paragraphs:
        transform_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    transform_paragraph(paragraph)

def uppercase_short_parentheses_content(text: str, max_len: int = 8, skip_words=None) -> str:
    if not text:
        return text
    if skip_words is None:
        skip_words = set()

    skip_words_ci = sorted((s.lower() for s in skip_words if s), key=len, reverse=True)
    paren_re = re.compile(r"\(([^()]*)\)")
    placeholder_re = re.compile(r"__skip_phrase_\d+__")

    def _replace_paren(m):
        inner = m.group(1)
        if not inner:
            return m.group(0)

        protected = []
        work = inner

        # Фрази зі skip_words залишаємо без змін, але інший текст у цих дужках можна підіймати.
        for phrase in skip_words_ci:
            rx = re.compile(re.escape(phrase), flags=re.IGNORECASE)

            def _stash(sm):
                key = f"__skip_phrase_{len(protected)}__"
                protected.append((key, sm.group(0)))
                return key

            work = rx.sub(_stash, work)

        # Рахуємо загальну довжину вмісту дужок (разом із пробілами/пунктуацією),
        # ігноруючи фрази зі SKIP_WORDS.
        effective = placeholder_re.sub("", work).strip()
        if len(effective) > max_len:
            return m.group(0)

        work = work.lower()

        for key, original in protected:
            work = work.replace(key, original)

        return f"({work})"

    return paren_re.sub(_replace_paren, text)

def apply_uppercase_short_parentheses_content(document, max_len: int = 8, skip_words=None):
    if skip_words is None:
        skip_words = set()

    def process_paragraph(paragraph):
        t = paragraph.text or ""
        if not t:
            return
        new_t = uppercase_short_parentheses_content(t, max_len=max_len, skip_words=skip_words)
        if new_t != t:
            paragraph.text = new_t

    for paragraph in document.paragraphs:
        process_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

def set_line_spacing_one(document):
    def apply(p):
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)

    for p in document.paragraphs:
        apply(p)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply(p)

def apply_bold(document, to_bold=None):
    if to_bold is None:
        to_bold = []

    def bold_paragraph(paragraph):
        if not to_bold:
            return
        for run in paragraph.runs:
            text = run.text or ""
            if not text:
                continue
            if any(pat in text for pat in to_bold):
                run.bold = True

    for paragraph in document.paragraphs:
        bold_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    bold_paragraph(paragraph)

def bold_prefix_before_colon(document, max_words=12):
    delim_re = re.compile(r"\s*:\s*|\s*[—–]\s*")

    def process_paragraph(p):
        text = p.text or ""
        if not text.strip():
            return

        text2 = text.replace("\u00A0", " ")
        m = delim_re.search(text2)
        if not m:
            return

        before = text2[:m.start()]
        delim  = text2[m.start():m.end()]
        after  = text2[m.end():]

        if not before.strip():
            return
        if len(before.split()) > max_words:
            return

        p.text = ""
        run_bold = p.add_run(before + delim)
        if ":" in delim and before.strip() in ITALIC_PREFIXES_WITH_COLON:
            run_bold.bold = False
            run_bold.italic = True
        else:
            run_bold.bold = True
        if after:
            p.add_run(after)

    for p in document.paragraphs:
        process_paragraph(p)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

def normalize_names_after_rank(document):
    def smart_capitalize(s: str) -> str:
        parts = s.split("-")
        return "-".join(p[:1].upper() + p[1:].lower() if p else p for p in parts)

    def split_punct(token: str):
        # відділяємо кінцеву пунктуацію типу: "солдат," -> ("солдат", ",")
        m = re.match(r"^(.+?)([,\.;:\)\]]+)?$", token)
        if not m:
            return token, ""
        return m.group(1), (m.group(2) or "")

    def fix_text(text: str) -> str:
        tokens = text.split()
        i = 0
        while i < len(tokens):
            base, punct = split_punct(tokens[i])
            t = base.lower()

            if t in RANKS and i + 3 < len(tokens):
                tokens[i + 1] = tokens[i + 1].upper()

                # 2) прізвище/ім'я/по-батькові -> лише перша велика
                tokens[i + 2] = smart_capitalize(tokens[i + 2])
                tokens[i + 3] = smart_capitalize(tokens[i + 3])

                i += 4
            else:
                i += 1

        return " ".join(tokens)

    # --- решту твоєї функції лишаємо як є ---
    paragraphs = list(document.paragraphs)
    page_break_indices = {idx for idx, p in enumerate(paragraphs) if p.paragraph_format.page_break_before}

    def is_protected(idx: int) -> bool:
        if idx in page_break_indices:
            return True
        if any((idx + k) in page_break_indices for k in (1, 2, 3)):
            return True
        return False

    for idx, p in enumerate(paragraphs):
        if is_protected(idx) or not p.text:
            continue
        fixed = fix_text(p.text)
        if fixed != p.text:
            p.text = fixed

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text:
                        continue
                    fixed = fix_text(p.text)
                    if fixed != p.text:
                        p.text = fixed

def normalize_rank_paragraph_endings(document):
    trailing_re = re.compile(r"[\s,;\.]+$")
    token_clean_re = re.compile(r"^[^\w\u0400-\u04FF]*|[^\w\u0400-\u04FF]*$")

    def first_two_have_rank(text: str) -> bool:
        tokens = (text or "").replace("\u00A0", " ").strip().split()
        for tok in tokens[:2]:
            cleaned = token_clean_re.sub("", tok).lower()
            if cleaned in RANKS:
                return True
        return False

    def is_date_paragraph_text(text: str) -> bool:
        return bool(DATE_ONLY_RE.match((text or "").strip()))

    paragraphs = list(document.paragraphs)
    for i, p in enumerate(paragraphs):
        current = p.text or ""
        if not current.strip():
            continue
        if not first_two_have_rank(current):
            continue

        next_text = paragraphs[i + 1].text if i + 1 < len(paragraphs) else ""
        if is_date_paragraph_text(next_text):
            continue

        next_has_rank = first_two_have_rank(next_text)
        suffix = ";" if next_has_rank else "."
        fixed = trailing_re.sub("", current) + suffix
        if fixed != current:
            p.text = fixed

def format_signature_block_above_page_break(document, spaces_between=10):
    def is_empty(p):
        return not (p.text or "").strip()

    while True:
        paragraphs = list(document.paragraphs)
        break_indices = [i for i, p in enumerate(paragraphs) if p.paragraph_format.page_break_before]
        if not break_indices:
            break

        changed_any = False

        for idx in reversed(break_indices):
            paragraphs = list(document.paragraphs)
            if idx >= len(paragraphs):
                continue

            p_break = paragraphs[idx]
            nonempty = []
            all_in_zone = []
            i = idx - 1
            if i < 0:
                continue

            while i >= 0 and len(nonempty) < 3:
                q = paragraphs[i]
                all_in_zone.append((i, q))
                if not is_empty(q):
                    nonempty.append((i, q))
                i -= 1

            if not nonempty:
                continue

            for j, q in all_in_zone:
                if is_empty(q):
                    delete_paragraph(q)
                    changed_any = True

            paragraphs = list(document.paragraphs)
            try:
                idx2 = paragraphs.index(p_break)
            except ValueError:
                continue
            if idx2 - 1 < 0:
                continue

            nonempty = []
            i = idx2 - 1
            while i >= 0 and len(nonempty) < 3:
                q = paragraphs[i]
                if not is_empty(q):
                    nonempty.append((i, q))
                i -= 1
            if not nonempty:
                continue

            nonempty_sorted = sorted(nonempty, key=lambda t: t[0])

            if len(nonempty_sorted) == 1:
                target = nonempty_sorted[0][1]
            elif len(nonempty_sorted) == 2:
                target = nonempty_sorted[1][1]
            else:
                target = nonempty_sorted[1][1]

            text = (target.text or "").strip()
            tokens = text.split()
            if len(tokens) < 3:
                continue

            left = " ".join(tokens[:-2])
            right = " ".join(tokens[-2:])
            target.text = f"{left}{' ' * spaces_between}{right}"

        if not changed_any:
            break

def fix_first_line_indent_except_specials(document):
    paragraphs = list(document.paragraphs)
    page_break_indices = {i for i, p in enumerate(paragraphs) if p.paragraph_format.page_break_before}

    last_nonempty_indices = []
    for idx in range(len(paragraphs) - 1, -1, -1):
        if (paragraphs[idx].text or "").strip():
            last_nonempty_indices.append(idx)
            if len(last_nonempty_indices) == 4:
                break
    last_nonempty_set = set(last_nonempty_indices)

    def is_protected(i, p):
        text = (p.text or "").strip()
        if text == "РАПОРТ":
            return True
        if "Командиру 3" in text:
            return True
        if i in page_break_indices:
            return True
        if any((i + k) in page_break_indices for k in (1, 2, 3, 4)):
            return True
        if i in last_nonempty_set:
            return True
        return False

    for i, p in enumerate(paragraphs):
        if is_protected(i, p):
            continue
        if not (p.text or "").strip():
            continue

        stripped = p.text.lstrip()
        if stripped != p.text:
            p.text = stripped

        p.paragraph_format.first_line_indent = Cm(1)

    for idx in last_nonempty_set:
        paragraphs[idx].paragraph_format.first_line_indent = Cm(0)

def ensure_final_page_break(document):
    paragraphs = document.paragraphs
    if not paragraphs:
        return

    last_idx = None
    for i in range(len(paragraphs) - 1, -1, -1):
        if (paragraphs[i].text or "").strip():
            last_idx = i
            break
    if last_idx is None:
        return

    new_paragraph = document.add_paragraph()
    new_paragraph.paragraph_format.page_break_before = True

def final_format_commandiru_indent(document):
    def process_paragraphs(paragraphs):
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            txt1 = " ".join(((p.text or "").replace("\n", " ")).split())

            if "Командиру 3" in txt1:
                pf = p.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.left_indent = Cm(9)
                pf.first_line_indent = Cm(0)
                pf.hanging_indent = None
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)

                if i + 1 < len(paragraphs):
                    p2 = paragraphs[i + 1]
                    txt2 = " ".join(((p2.text or "").replace("\n", " ")).split())
                    if "військової частини А0998" in txt2:
                        pf2 = p2.paragraph_format
                        pf2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        pf2.left_indent = Cm(9)
                        pf2.first_line_indent = Cm(0)
                        pf2.hanging_indent = None
                        pf2.space_before = Pt(0)
                        pf2.space_after = Pt(0)
                i += 1
            else:
                i += 1

    process_paragraphs(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

def reset_left_indent_for_three_above_page_break(document):
    paragraphs = list(document.paragraphs)
    page_break_indices = {i for i, p in enumerate(paragraphs) if p.paragraph_format.page_break_before}

    for j in page_break_indices:
        for i in (j - 1, j - 2, j - 3):
            if i < 0 or i >= len(paragraphs):
                continue
            p = paragraphs[i]
            if not (p.text or "").strip():
                continue
            p.paragraph_format.left_indent = Cm(0)

    nonempty_indices = []
    for idx in range(len(paragraphs) - 1, -1, -1):
        if (paragraphs[idx].text or "").strip():
            nonempty_indices.append(idx)
            if len(nonempty_indices) == 5:
                break

    for idx in nonempty_indices:
        paragraphs[idx].paragraph_format.left_indent = Cm(0)

def set_doc_spacing_before_after(document, before_pt=2, after_pt=2):
    def apply(paragraph):
        pf = paragraph.paragraph_format
        pf.space_before = Pt(before_pt)
        pf.space_after  = Pt(after_pt)

    for p in document.paragraphs:
        apply(p)

    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        apply(p)
                    walk_tables(cell.tables)

    walk_tables(document.tables)

def hard_fix_commandiru_indent(document):
    def wipe_indent_paragraph(p):
        try:
            p.style = p.part.document.styles['Normal']
        except Exception:
            pass

        pf = p.paragraph_format

        pPr = p._p.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            _safe_remove_from_parent(numPr)

        if pf.tab_stops is not None:
            try:
                pf.tab_stops.clear_all()
            except Exception:
                pass

        pf.left_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        pf.hanging_indent = None

        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent = Cm(9)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    def process_paragraphs(paragraphs):
        paragraphs = list(paragraphs)
        for i, p in enumerate(paragraphs):
            txt = (p.text or "").replace("\n", " ").strip()
            if "Командиру 3" in txt:
                wipe_indent_paragraph(p)
                if i + 1 < len(paragraphs):
                    p2 = paragraphs[i + 1]
                    if "військової частини А0998" in (p2.text or ""):
                        wipe_indent_paragraph(p2)

    process_paragraphs(document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

def mark_all_text_as_ukrainian(document, lang="uk-UA"):
    def set_run_lang(run):
        rPr = run._r.get_or_add_rPr()
        lang_elm = rPr.find(qn("w:lang"))
        if lang_elm is None:
            lang_elm = OxmlElement("w:lang")
            rPr.append(lang_elm)

        lang_elm.set(qn("w:val"), lang)
        lang_elm.set(qn("w:eastAsia"), lang)
        lang_elm.set(qn("w:bidi"), lang)

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                if run.text:
                    set_run_lang(run)

    process_paragraphs(document.paragraphs)

    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)
                    walk_tables(cell.tables)

    walk_tables(document.tables)

def clean_process_docx_file(path: Path):
    document = Document(path)

    clean_remove_page_breaks(document)
    fix_space_before_parenthesis(document)
    apply_text_transforms(
        document,
        to_upper=TEXT_TO_UPPER,
        to_lower=TEXT_TO_LOWER,
        replacements=TEXT_REPLACEMENTS,
    )
    apply_uppercase_short_parentheses_content(document, max_len=8, skip_words=SKIP_WORDS)
    normalize_apostrophes_in_document(document)
    justify_all_paragraphs(document)
    center_raport(document)
    normalize_commandiru_block(document)
    handle_commandiru_3(document)
    clear_shading_and_color(document)
    fix_left_indents(document)
    set_margins(document, left_cm=2.5, other_cm=1.5)
    apply_italic(document, to_italic=TEXT_TO_ITALIC)
    apply_bold(document, to_bold=TEXT_TO_BOLD)
    bold_prefix_before_colon(document, max_words=12)
    normalize_names_after_rank(document)
    normalize_rank_paragraph_endings(document)
    fix_first_line_indent_except_specials(document)
    reset_left_indent_for_three_above_page_break(document)
    remove_trailing_empty_paragraphs(document)
    format_signature_block_above_page_break(document)
    final_format_commandiru_indent(document)
    hard_fix_commandiru_indent(document)
    set_doc_spacing_before_after(document, before_pt=2, after_pt=2)
    set_line_spacing_one(document)
    normalize_font(document, font_name="Times New Roman", font_size_pt=12)
    ensure_final_page_break(document)
    normalize_date_block_spacing(document, empty_before_count=1)
    set_last_four_nonempty_indent_rules(document, fourth_indent_cm=1.0)
    fix_last_daily_signature_line(document, spacer=(" " * 55) + "\t")
    mark_all_text_as_ukrainian(document)

    document.save(path)

def clean_process_all_docx(root_folder: Path, progress: ProgressWindow | None = None):
    filenames = os.listdir(root_folder)
    if progress:
        progress.update(
            header="Етап 4/4: Clean & Format",
            detail="Фінальна чистка та форматування DOCX...",
            current=0,
            total=len(filenames),
            file_name="",
        )

    for i, filename in enumerate(filenames, 1):
        if progress:
            progress.update(current=i - 1, total=len(filenames), file_name=filename)
        full_path = root_folder / filename

        if not full_path.is_file():
            if progress:
                progress.update(current=i, total=len(filenames), file_name=filename)
            continue
        if not filename.lower().endswith(".docx"):
            if progress:
                progress.update(current=i, total=len(filenames), file_name=filename)
            continue
        if filename.startswith("~$"):
            if progress:
                progress.update(current=i, total=len(filenames), file_name=filename)
            continue
        if filename in SKIP_FILES:
            if progress:
                progress.update(current=i, total=len(filenames), file_name=filename)
            continue
        if BACKUP_ROOT and BACKUP_ROOT in full_path.parents:
            if progress:
                progress.update(current=i, total=len(filenames), file_name=filename)
            continue

        log(f"[CLEAN] {full_path.name}")
        try:
            clean_process_docx_file(full_path)
        except Exception as e:
            log(f"[CLEAN ERROR] {full_path.name}: {e}")
        if progress:
            progress.update(current=i, total=len(filenames), file_name=filename)

# =================================================
#                 1 ВІКНО НАЛАШТУВАНЬ
# =================================================

def ask_all_settings_one_window(parent: tk.Misc | None = None) -> Path:
    folder = ask_folder_settings_window(
        parent,
        app_name="FormatBRO",
        subtitle="Патч та чистка форматування у вибраній директорії...",
        hint="Вкажіть папку з БРО-рапортами, після чого запустіть обробку...",
        browse_title="Виберіть папку з БРО-рапортами",
    )
    if not folder:
        raise SystemExit("Скасовано.")

    return folder

# =================================================
#                     MAIN
# =================================================

def main():
    global BASE_DIR, ROOT_DIR, BACKUP_ROOT

    ui_root = tk.Tk()
    ui_root.withdraw()
    install_frozen_executable_icon(ui_root)
    install_dark_title_bar(ui_root)
    ui_root.attributes("-topmost", True)

    BASE_DIR = ask_all_settings_one_window(parent=ui_root)

    ROOT_DIR = BASE_DIR
    script_name = Path(__file__).stem
    timestamp = datetime.now().strftime("%Y.%m.%d %H.%M")
    BACKUP_ROOT = BASE_DIR / f"Backup - {script_name} - {timestamp}"

    progress = ProgressWindow(ui_root)

    try:
        run_format_workflow_with_progress(progress)
    finally:
        progress.close()

    log("[START] Об’єднаний скрипт: патч документів -> (опційно) чистка форматування")
    # Лог у backup
    log_path = BACKUP_ROOT / "FormatBRO.log"
    try:
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("\n".join(LOG_LINES))
    except Exception as e:
        log(f"[ERROR] Не вдалось записати FormatBRO.log: {e}")

    # Переміщення backup у кошик
    try:
        if BACKUP_ROOT.exists() and BACKUP_ROOT.is_dir():
            send2trash(str(BACKUP_ROOT))
            log(f"[CLEANUP] Папку backup переміщено у кошик: {BACKUP_ROOT}")
    except Exception as e:
        log(f"[CLEANUP ERROR] Не вдалося перемістити backup у кошик: {e}")

    # ✅ messagebox поверх усіх вікон
    try:
        ui_root.attributes("-topmost", True)
        ui_root.lift()
        ui_root.update_idletasks()
        messagebox.showinfo("Готово", "Документи оброблено.", parent=ui_root)
        ui_root.destroy()
    except Exception:
        pass

if __name__ == "__main__":
    main()
