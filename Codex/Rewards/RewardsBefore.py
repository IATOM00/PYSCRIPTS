#!python3.11
from openpyxl.utils import column_index_from_string, get_column_letter, range_boundaries
from openpyxl.styles import Alignment, Border, Side, PatternFill, Font, Color
from tkinter import Tk, filedialog, messagebox, StringVar, ttk
import warnings, calendar, shutil, time, json, re, sys, os, queue
from datetime import datetime, date, timedelta
from typing import Dict, Any, List, Tuple
from openpyxl import load_workbook as _openpyxl_load_workbook
from send2trash import send2trash
import tkinter.font as tkfont
from zipfile import BadZipFile, ZipFile
from collections import deque
import multiprocessing as mp
from docx import Document
from pathlib import Path
from copy import copy
import tkinter as tk


class UserFacingError(Exception):
    title = "Помилка"

    def __init__(self, message: str, title: str | None = None):
        super().__init__(message)
        if title:
            self.title = title


class FileAccessError(UserFacingError):
    title = "Файл зайнятий"


_FILE_ACCESS_WINERRORS = {5, 32, 33}


def _is_file_access_error(exc: BaseException) -> bool:
    return isinstance(exc, PermissionError) or getattr(exc, "winerror", None) in _FILE_ACCESS_WINERRORS


def _permission_error_path(exc: BaseException, fallback: Path | None = None) -> Path | None:
    for attr in ("filename2", "filename"):
        value = getattr(exc, attr, None)
        if value:
            try:
                return Path(value)
            except Exception:
                pass
    return Path(fallback) if fallback is not None else None


def _file_access_message(path: Path | None, action: str, exc: BaseException | None = None) -> str:
    path_text = str(path) if path else "невідомий файл"
    detail = f"\n\nТехнічна причина: {exc}" if exc else ""
    return (
        f"Не вдалося {action}.\n\n"
        f"Файл:\n{path_text}\n\n"
        "Причина: файл використовується іншим процесом.\n"
        "Закрийте цей файл, чи процес що його використовує..."
    )


def load_workbook(filename, *args, **kwargs):
    try:
        suffix = Path(filename).suffix.lower()
    except TypeError:
        suffix = ""

    if suffix == ".xlsm":
        kwargs.setdefault("keep_vba", True)
    kwargs.setdefault("keep_links", True)

    try:
        return _openpyxl_load_workbook(filename, *args, **kwargs)
    except OSError as exc:
        if _is_file_access_error(exc):
            try:
                fallback_path = Path(filename)
            except TypeError:
                fallback_path = None
            path = _permission_error_path(exc, fallback_path)
            raise FileAccessError(_file_access_message(path, "відкрити Excel-файл", exc)) from exc
        raise


def install_frozen_executable_icon(root, retry_ms: int = 250) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    _set_frozen_executable_icon(root)
    try:
        root.after(retry_ms, lambda: _set_frozen_executable_icon(root))
    except Exception:
        pass


def _set_frozen_executable_icon(root) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    try:
        import ctypes
        from ctypes import wintypes

        wm_seticon = 0x0080
        icon_small = 0
        icon_big = 1
        gclp_hicon = -14
        gclp_hiconsm = -34

        shell32 = ctypes.WinDLL("shell32", use_last_error=True)
        user32 = ctypes.WinDLL("user32", use_last_error=True)
        long_ptr = (
            ctypes.c_longlong if ctypes.sizeof(ctypes.c_void_p) == 8 else ctypes.c_long
        )

        shell32.ExtractIconExW.argtypes = (
            wintypes.LPCWSTR,
            ctypes.c_int,
            ctypes.POINTER(wintypes.HICON),
            ctypes.POINTER(wintypes.HICON),
            wintypes.UINT,
        )
        shell32.ExtractIconExW.restype = wintypes.UINT
        user32.SendMessageW.argtypes = (
            wintypes.HWND,
            wintypes.UINT,
            wintypes.WPARAM,
            wintypes.LPARAM,
        )
        user32.SendMessageW.restype = long_ptr
        user32.GetParent.argtypes = (wintypes.HWND,)
        user32.GetParent.restype = wintypes.HWND

        try:
            set_class_icon = user32.SetClassLongPtrW
        except AttributeError:
            set_class_icon = user32.SetClassLongW
        set_class_icon.argtypes = (wintypes.HWND, ctypes.c_int, long_ptr)
        set_class_icon.restype = long_ptr

        icon_handles = getattr(root, "_frozen_executable_icon_handles", None)
        if icon_handles is None:
            large_icons = (wintypes.HICON * 1)()
            small_icons = (wintypes.HICON * 1)()
            icon_count = shell32.ExtractIconExW(
                sys.executable,
                0,
                large_icons,
                small_icons,
                1,
            )
            if icon_count <= 0:
                return
            icon_handles = (small_icons[0], large_icons[0])
            root._frozen_executable_icon_handles = icon_handles
        small_icon, large_icon = icon_handles

        root.update_idletasks()
        hwnds = []
        for raw_hwnd in (root.winfo_id(), root.wm_frame()):
            try:
                hwnd = int(str(raw_hwnd), 0)
            except (TypeError, ValueError):
                continue
            while hwnd and hwnd not in hwnds:
                hwnds.append(hwnd)
                hwnd = user32.GetParent(hwnd)

        for hwnd in hwnds:
            if small_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_small, small_icon)
                set_class_icon(hwnd, gclp_hiconsm, small_icon)
            if large_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_big, large_icon)
                set_class_icon(hwnd, gclp_hicon, large_icon)
    except Exception:
        pass


def install_dark_title_bar(window, retry_ms: int = 80) -> None:
    if sys.platform != "win32":
        return

    def _apply() -> None:
        try:
            import ctypes

            hwnd = ctypes.windll.user32.GetParent(window.winfo_id())
            value = ctypes.c_int(1)
            for attr in (20, 19):
                ctypes.windll.dwmapi.DwmSetWindowAttribute(
                    hwnd,
                    attr,
                    ctypes.byref(value),
                    ctypes.sizeof(value),
                )
        except Exception:
            pass

    try:
        window.update_idletasks()
        _apply()
        window.after(retry_ms, _apply)
    except Exception:
        pass

# ===================== LOG + TIMERS =====================

_RUN_T0 = None
_STAGE_T0 = None
_STAGE_NAME = None
_LOG_QUEUE = None
_LOG_PROC = None
_RUN_LOG_LINES = []
_RUN_ENDED = False
_SESSION_LOG_SENT = False
VALIDATE_EXCEL_AFTER_OPENPYXL_SAVE = False
_SAVE_CHECK_EXCEL = None
_SAVE_CHECK_COM_READY = False

STAFF_SHEET_NAME = "таблиця"


def format_path_for_entry_display(path: str) -> str:
    path = path.strip()
    if not path:
        return ""

    normalized = Path(path)
    drive_prefix = ""
    if normalized.drive and len(normalized.drive) == 2 and normalized.drive[1] == ":":
        drive_prefix = f"{normalized.drive}/"

    name = normalized.name
    parent_name = normalized.parent.name
    if parent_name and name:
        return f"{drive_prefix}.../{parent_name}/{name}"
    if name:
        return f"{drive_prefix}{name}" if drive_prefix else f".../{name}"
    if drive_prefix:
        return drive_prefix
    return path.replace("\\", "/")


def fit_path_for_entry_display(path: str, entry) -> str:
    display_path = format_path_for_entry_display(path)
    if not display_path:
        return ""

    entry.update_idletasks()
    font_name = entry.cget("font") or "TkDefaultFont"
    try:
        font = tkfont.nametofont(font_name)
    except Exception:
        return display_path

    available_width = max(120, entry.winfo_width() - 14)
    if font.measure(display_path) <= available_width:
        return display_path

    parts = display_path.split("/")
    if len(parts) >= 3:
        prefix = "/".join(parts[:-1]) + "/..."
        shortened_name = parts[-1]
        while len(shortened_name) > 1 and font.measure(prefix + shortened_name) > available_width:
            shortened_name = shortened_name[1:]
        return prefix + shortened_name

    ellipsis = "..."
    shortened_path = display_path
    while len(shortened_path) > 1 and font.measure(ellipsis + shortened_path) > available_width:
        shortened_path = shortened_path[1:]
    return ellipsis + shortened_path


def refresh_path_entry_display(source_var: StringVar, display_var: StringVar, entry) -> None:
    try:
        exists = bool(entry.winfo_exists())
    except Exception:
        exists = False
    if not exists:
        return
    display_var.set(fit_path_for_entry_display(source_var.get(), entry))


DEFAULT_START_COL_LETTER = "F"
DEFAULT_START_COL_IDX = column_index_from_string(DEFAULT_START_COL_LETTER)
HEADER_PERIOD_RE = re.compile(r"\b(20\d{2})-(0[1-9]|1[0-2])\b")
HEADER_PERIOD_SCAN_SHEETS = ("100 000", "70 000")


def _parse_header_year_month(value) -> tuple[int, int] | None:
    if value is None:
        return None
    match = HEADER_PERIOD_RE.search(str(value))
    if not match:
        return None
    return int(match.group(1)), int(match.group(2))


def detect_template_period(
    target_xlsx: Path,
    preferred_sheets: tuple[str, ...] = HEADER_PERIOD_SCAN_SHEETS,
) -> tuple[int, int]:
    wb = load_workbook(target_xlsx, read_only=True, data_only=True)
    try:
        ordered_sheet_names: list[str] = []
        seen: set[str] = set()
        for sheet_name in (*preferred_sheets, *wb.sheetnames):
            if sheet_name in wb.sheetnames and sheet_name not in seen:
                ordered_sheet_names.append(sheet_name)
                seen.add(sheet_name)

        for sheet_name in ordered_sheet_names:
            ws = wb[sheet_name]
            for row in ws.iter_rows(min_row=1, max_row=3, min_col=1, max_col=5, values_only=True):
                for value in row:
                    parsed = _parse_header_year_month(value)
                    if parsed is not None:
                        return parsed
    finally:
        wb.close()

    raise ValueError("Не знайшов формат YYYY-MM у заголовку шаблону (A1:E3).")


def detect_template_context(
    target_xlsx: Path,
    preferred_sheets: tuple[str, ...] = HEADER_PERIOD_SCAN_SHEETS,
) -> tuple[int, int, int]:
    year, month = detect_template_period(target_xlsx, preferred_sheets=preferred_sheets)
    return DEFAULT_START_COL_IDX, year, month


def configure_launch_styles(root):
    colors = {
        "window": "#F4F7FB",
        "panel": "#FFFFFF",
        "border": "#D6DEE8",
        "text": "#15202B",
        "muted": "#617084",
        "accent": "#0F766E",
        "accent_dark": "#0B5D56",
        "header": "#103C3A",
        "header_badge": "#0F766E",
    }

    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except Exception:
        pass

    try:
        style.configure("LaunchRoot.TFrame", background=colors["window"])
        style.configure("LaunchPanel.TFrame", background=colors["panel"])
        style.configure("LaunchSection.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        style.configure("LaunchField.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10, "bold"))
        style.configure("LaunchMuted.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure(
            "Launch.TEntry",
            fieldbackground="#FFFFFF",
            foreground=colors["text"],
            bordercolor=colors["border"],
            lightcolor=colors["border"],
            darkcolor=colors["border"],
            padding=(8, 6),
        )
        style.map(
            "Launch.TEntry",
            fieldbackground=[("readonly", "#FFFFFF")],
            foreground=[("readonly", colors["text"])],
        )
        style.configure(
            "LaunchBrowse.TButton",
            font=("Segoe UI", 10, "bold"),
            padding=(15, 8),
            foreground="#FFFFFF",
            background="#044743",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchBrowse.TButton",
            background=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
        )
        style.configure(
            "LaunchPrimary.TButton",
            font=("Segoe UI", 10, "bold"),
            padding=(18, 8),
            foreground="#FFFFFF",
            background="#024E49",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchPrimary.TButton",
            background=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
        )
        style.configure(
            "LaunchHero.TButton",
            font=("Segoe UI", 11, "bold"),
            padding=(20, 12),
            foreground="#FFFFFF",
            background="#024E49",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchHero.TButton",
            background=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
        )
    except Exception:
        pass

    return colors


def _find_last_used_row_in_band(ws, min_col: int, max_col: int, start_row: int) -> int:
    last_used = start_row
    max_row = ws.max_row or start_row

    for r in range(start_row, max_row + 1):
        for c in range(min_col, max_col + 1):
            v = ws.cell(row=r, column=c).value
            if v is not None and str(v).strip() != "":
                last_used = r
                break

    return last_used


def _table_totals_row_count(table) -> int:
    count = getattr(table, "totalsRowCount", None)
    if count is not None:
        try:
            return max(0, int(count))
        except Exception:
            return 0

    return 1 if getattr(table, "totalsRowShown", False) else 0


def _table_autofilter_ref(table, min_col: int, min_row: int, max_col: int, max_row: int) -> str:
    totals_count = _table_totals_row_count(table)
    filter_max_row = max(min_row, max_row - totals_count)
    return f"{get_column_letter(min_col)}{min_row}:{get_column_letter(max_col)}{filter_max_row}"


def _sync_excel_table_column_ids(ws, table) -> int:
    table_columns = list(getattr(table, "tableColumns", []) or [])
    if not table_columns:
        return 0

    fixed = 0
    for expected_id, table_column in enumerate(table_columns, start=1):
        if getattr(table_column, "id", None) == expected_id:
            continue
        table_column.id = expected_id
        fixed += 1

    if fixed:
        log(
            f"[TABLE] {ws.title}/{getattr(table, 'displayName', 'Table')} "
            f"renumbered tableColumn ids: fixed={fixed}"
        )

    return fixed


def _sync_excel_table_refs(ws) -> int:
    tables = getattr(ws, "tables", None)
    if not tables:
        return 0

    fixed = 0
    for table in list(tables.values()):
        if _sync_excel_table_column_ids(ws, table):
            fixed += 1

        ref = getattr(table, "ref", "") or ""
        if ":" not in ref:
            continue

        try:
            min_col, min_row, max_col, max_row = range_boundaries(ref)
        except Exception:
            continue

        last_used_row = _find_last_used_row_in_band(ws, min_col, max_col, min_row)
        new_max_row = max(min_row + 1, min(ws.max_row or max_row, max(max_row, last_used_row)))
        new_ref = f"{get_column_letter(min_col)}{min_row}:{get_column_letter(max_col)}{new_max_row}"
        new_filter_ref = _table_autofilter_ref(table, min_col, min_row, max_col, new_max_row)
        changed = False

        if new_ref != ref:
            table.ref = new_ref
            changed = True
            log(f"[TABLE] {ws.title}/{getattr(table, 'displayName', 'Table')}: {ref} -> {new_ref}")

        try:
            auto_filter = getattr(table, "autoFilter", None)
            old_filter_ref = getattr(auto_filter, "ref", "") if auto_filter is not None else ""
            if auto_filter is not None and old_filter_ref != new_filter_ref:
                auto_filter.ref = new_filter_ref
                changed = True
                log(
                    f"[TABLE] {ws.title}/{getattr(table, 'displayName', 'Table')} "
                    f"autoFilter: {old_filter_ref} -> {new_filter_ref}"
                )
        except Exception:
            pass

        if changed:
            fixed += 1

    return fixed


def _assert_valid_xlsx_package(path: Path) -> None:
    try:
        with ZipFile(path) as zf:
            bad_member = zf.testzip()
    except BadZipFile as exc:
        raise ValueError(f"Збережений Excel-файл має пошкоджений ZIP-контейнер: {path}") from exc

    if bad_member:
        raise ValueError(f"Збережений Excel-файл має пошкоджену частину ZIP: {bad_member}")


def _safe_save_workbook(wb, target_xlsx: Path) -> None:
    target_xlsx = Path(target_xlsx)
    tmp_path = target_xlsx.with_name(f".{target_xlsx.stem}.{os.getpid()}.tmp{target_xlsx.suffix}")

    try:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except OSError as exc:
            if _is_file_access_error(exc):
                path = _permission_error_path(exc, tmp_path)
                raise FileAccessError(_file_access_message(path, "прибрати старий тимчасовий Excel-файл", exc)) from exc
            raise

        try:
            wb.save(tmp_path)
        except OSError as exc:
            if _is_file_access_error(exc):
                path = _permission_error_path(exc, tmp_path)
                raise FileAccessError(_file_access_message(path, "створити тимчасовий Excel-файл", exc)) from exc
            raise

        _assert_valid_xlsx_package(tmp_path)
        _validate_saved_workbook_with_excel(tmp_path)

        try:
            tmp_path.replace(target_xlsx)
        except OSError as exc:
            if _is_file_access_error(exc):
                path = _permission_error_path(exc, target_xlsx)
                raise FileAccessError(_file_access_message(path, "перезаписати Excel-файл", exc)) from exc
            raise
    except Exception:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass
        raise


def _get_save_check_excel():
    global _SAVE_CHECK_EXCEL, _SAVE_CHECK_COM_READY
    try:
        import pythoncom
        import win32com.client as win32
    except Exception as exc:
        raise RuntimeError(f"Excel COM validation unavailable: {exc}") from exc

    if _SAVE_CHECK_EXCEL is not None:
        return _SAVE_CHECK_EXCEL

    pythoncom.CoInitialize()
    _SAVE_CHECK_COM_READY = True
    excel = win32.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    excel.AskToUpdateLinks = False
    try:
        excel.AutomationSecurity = 3
    except Exception:
        pass

    _SAVE_CHECK_EXCEL = excel
    return excel


def _shutdown_save_check_excel() -> None:
    global _SAVE_CHECK_EXCEL, _SAVE_CHECK_COM_READY

    if _SAVE_CHECK_EXCEL is not None:
        try:
            _SAVE_CHECK_EXCEL.Quit()
        except Exception:
            pass
        _SAVE_CHECK_EXCEL = None

    if _SAVE_CHECK_COM_READY:
        try:
            import pythoncom

            pythoncom.CoUninitialize()
        except Exception:
            pass
        _SAVE_CHECK_COM_READY = False


def _validate_saved_workbook_with_excel(target_xlsx: Path) -> None:
    if not VALIDATE_EXCEL_AFTER_OPENPYXL_SAVE or sys.platform != "win32":
        return

    try:
        excel = _get_save_check_excel()

        opened = None
        try:
            opened = excel.Workbooks.Open(
                str(target_xlsx),
                UpdateLinks=0,
                ReadOnly=True,
                Notify=False,
                AddToMru=False,
                CorruptLoad=0,
            )
            saved_state = bool(getattr(opened, "Saved", True))
            sheet_count = int(opened.Worksheets.Count)
            log(f"[SAVE/CHECK] Excel open OK: sheets={sheet_count}, saved={saved_state}")
            if not saved_state:
                raise ValueError(
                    "Excel відкрив файл, але одразу позначив його як змінений. "
                    "Це часто означає repair/recovery або перерахунок структури книги."
                )
        finally:
            if opened is not None:
                try:
                    opened.Close(SaveChanges=False)
                except Exception:
                    pass
    except Exception as exc:
        stage = _STAGE_NAME or "unknown stage"
        raise ValueError(f"Excel COM validation failed after openpyxl save during '{stage}': {exc}") from exc


def save_target_workbook(wb, target_xlsx: Path) -> int:
    fixed_total = 0
    for ws in wb.worksheets:
        fixed_total += _sync_excel_table_refs(ws)
    _safe_save_workbook(wb, target_xlsx)
    if fixed_total:
        log(f"[TABLE] synced refs before save: {fixed_total}")
    return fixed_total

def get_staff_sheet(wb):
    for sheet_name in wb.sheetnames:
        if str(sheet_name).strip().casefold() == STAFF_SHEET_NAME:
            return wb[sheet_name]
    raise ValueError(
        f"Аркуш '{STAFF_SHEET_NAME}' не знайдено у штатці. Доступні аркуші: {', '.join(wb.sheetnames)}"
    )

def _ts() -> str:
    return datetime.now().strftime("%H:%M:%S")

def _log_window_worker(log_queue, title: str, main_pid: int):
    import tkinter as tk
    from tkinter.scrolledtext import ScrolledText
    import ctypes
    import os
    import signal

    win = tk.Tk()
    win.withdraw()
    install_frozen_executable_icon(win)
    win.title(title)
    win.geometry("750x450")
    win.attributes("-topmost", True)
    colors = configure_launch_styles(win)
    win.configure(bg=colors["window"])
    install_dark_title_bar(win)

    surface = ttk.Frame(win, style="LaunchRoot.TFrame")
    surface.pack(fill="both", expand=True)
    header = tk.Frame(surface, bg=colors["header"], padx=22, pady=10)
    header.pack(fill="x")
    tk.Label(
        header,
        text=title,
        bg=colors["header"],
        fg="#FFFFFF",
        font=("Segoe UI Semibold", 15),
        anchor="w",
    ).pack(anchor="w")
    tk.Label(
        header,
        text="Журнал виконання оновлюється під час обробки.",
        bg=colors["header"],
        fg="#D7FBF5",
        font=("Segoe UI", 9),
        anchor="w",
    ).pack(anchor="w", pady=(3, 4))
    body = ttk.Frame(surface, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
    body.pack(fill="both", expand=True)
    panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
    panel_shell.pack(fill="both", expand=True)

    text = ScrolledText(
        panel_shell,
        wrap="word",
        bg=colors["panel"],
        fg=colors["text"],
        insertbackground=colors["text"],
        relief="flat",
        bd=0,
        font=("Consolas", 9),
    )
    text.pack(fill="both", expand=True, padx=10, pady=10)
    text.configure(state="disabled")

    bottom = tk.Frame(surface, bg=colors["window"])
    bottom.pack(fill="x", padx=10, pady=(0, 10))

    def copy_selection(event=None):
        try:
            sel = text.get("sel.first", "sel.last")
        except Exception:
            return "break"
        win.clipboard_clear()
        win.clipboard_append(sel)
        return "break"

    text.bind("<Control-c>", copy_selection)
    text.bind("<Control-C>", copy_selection)

    def terminate_main_process():
        # Якщо користувач закрив лог-вікно через X — зупиняємо основний процес повністю.
        try:
            if os.name == "nt":
                PROCESS_TERMINATE = 0x0001
                handle = ctypes.windll.kernel32.OpenProcess(PROCESS_TERMINATE, False, int(main_pid))
                if handle:
                    ctypes.windll.kernel32.TerminateProcess(handle, 1)
                    ctypes.windll.kernel32.CloseHandle(handle)
                    return
            os.kill(int(main_pid), signal.SIGTERM)
        except Exception:
            pass

    def on_close():
        terminate_main_process()
        win.destroy()

    win.protocol("WM_DELETE_WINDOW", on_close)

    def pump_queue():
        try:
            for _ in range(400):
                line = log_queue.get_nowait()
                if line is None:
                    win.destroy()
                    return
                text.configure(state="normal")
                text.insert("end", line)
                text.see("end")
                text.configure(state="disabled")
        except queue.Empty:
            pass
        win.after(80, pump_queue)

    win.after(50, pump_queue)
    win.update_idletasks()
    win.deiconify()
    win.lift()
    win.focus_force()
    win.mainloop()

def _init_log_window(title: str = "Логи: RewardsBefore"):
    global _LOG_QUEUE, _LOG_PROC
    if _LOG_PROC is not None and _LOG_PROC.is_alive():
        return
    ctx = mp.get_context("spawn")
    _LOG_QUEUE = ctx.Queue(maxsize=10000)
    _LOG_PROC = ctx.Process(target=_log_window_worker, args=(_LOG_QUEUE, title, os.getpid()), daemon=True)
    _LOG_PROC.start()

def _close_log_window():
    global _LOG_QUEUE, _LOG_PROC
    try:
        if _LOG_QUEUE is not None:
            _LOG_QUEUE.put_nowait(None)
    except Exception:
        pass
    try:
        if _LOG_PROC is not None:
            _LOG_PROC.join(timeout=1.5)
            if _LOG_PROC.is_alive():
                _LOG_PROC.terminate()
    except Exception:
        pass
    _LOG_QUEUE = None
    _LOG_PROC = None

def log(msg: str):
    global _LOG_QUEUE, _LOG_PROC, _RUN_LOG_LINES
    line = f"[{_ts()}] {msg}"
    print(line)
    _RUN_LOG_LINES.append(line + "\n")
    try:
        if _LOG_QUEUE is not None:
            _LOG_QUEUE.put_nowait(line + "\n")
    except queue.Full:
        # Вікно не встигає читати — просто пропускаємо частину UI-логів.
        pass
    except Exception:
        # Якщо лог-вікно закрили (X) або пайп черги зламався,
        # не ламаємо основний процес.
        _LOG_QUEUE = None
        _LOG_PROC = None


def run_start():
    global _RUN_T0, _RUN_LOG_LINES, _RUN_ENDED, _SESSION_LOG_SENT
    _RUN_LOG_LINES = []
    _RUN_ENDED = False
    _SESSION_LOG_SENT = False
    _RUN_T0 = time.perf_counter()
    log("=== START SCRIPT ===")

def _send_session_log_to_trash(script_label: str = "RewardsBefore"):
    global _SESSION_LOG_SENT
    if _SESSION_LOG_SENT:
        return
    if not _RUN_LOG_LINES:
        return
    try:
        desktop = Path.home() / "Desktop"
        timestamp = datetime.now().strftime("%Y.%m.%d %H.%M")
        log_path = desktop / f"{script_label} - {timestamp}.log"
        log_path.write_text("".join(_RUN_LOG_LINES), encoding="utf-8", errors="replace")
        send2trash(str(log_path))
        _SESSION_LOG_SENT = True
    except Exception:
        pass

def run_end():
    global _RUN_ENDED
    if _RUN_ENDED:
        return
    _RUN_ENDED = True
    if _RUN_T0 is None:
        log("=== END SCRIPT (no timer) ===")
        return
    dt = time.perf_counter() - _RUN_T0
    log(f"=== END SCRIPT === total={dt:.2f}s ({dt/60:.2f} min)")


def show_user_error_window(exc: UserFacingError) -> None:
    root = None
    try:
        root = Tk()
        root.withdraw()
        install_frozen_executable_icon(root)
        install_dark_title_bar(root)
        root.attributes("-topmost", True)
        messagebox.showerror(getattr(exc, "title", "Помилка"), str(exc), parent=root)
    except Exception:
        print(f"{getattr(exc, 'title', 'Помилка')}: {exc}")
    finally:
        if root is not None:
            try:
                root.destroy()
            except Exception:
                pass


def cleanup_failed_run() -> None:
    try:
        run_end()
    except Exception:
        pass
    try:
        _shutdown_save_check_excel()
    except Exception:
        pass
    try:
        _send_session_log_to_trash("RewardsBefore")
    except Exception:
        pass
    try:
        _close_log_window()
    except Exception:
        pass


def stage_start(name: str):
    global _STAGE_T0, _STAGE_NAME
    _STAGE_NAME = name
    _STAGE_T0 = time.perf_counter()
    log(f"--- STAGE START: {name} ---")

def stage_end():
    global _STAGE_T0, _STAGE_NAME
    if _STAGE_T0 is None:
        return
    dt = time.perf_counter() - _STAGE_T0
    log(f"--- STAGE END: {_STAGE_NAME} --- took={dt:.2f}s ({dt/60:.2f} min)")
    _STAGE_T0 = None
    _STAGE_NAME = None

# ===================== GUI: МАЙСТЕР ВИБОРУ (ВСЕ В ОДНОМУ) =====================

def ask_all_inputs_window() -> tuple[Path | None, Path | None, Path, int, int, int, int]:
    root = Tk()
    root.withdraw()
    install_frozen_executable_icon(root)
    root.title("НАЛАШТУВАННЯ: RewardsBefore")
    root.attributes("-topmost", True)
    root.resizable(False, False)
    colors = configure_launch_styles(root)
    root.configure(bg=colors["window"])
    install_dark_title_bar(root)

    bro_var = StringVar(value="")
    staff_var = StringVar(value="")
    xlsx_var = StringVar(value="")
    bro_display_var = StringVar(value="")
    staff_display_var = StringVar(value="")
    xlsx_display_var = StringVar(value="")
    entry_width = 45

    surface = ttk.Frame(root, style="LaunchRoot.TFrame")
    surface.grid(row=0, column=0, sticky="nsew")
    surface.grid_columnconfigure(0, weight=1)

    header = tk.Frame(surface, bg=colors["header"], padx=22, pady=11)
    header.grid(row=0, column=0, sticky="we")
    header.grid_columnconfigure(0, weight=1)
    tk.Label(
        header,
        text="RewardsBefore",
        bg=colors["header"],
        fg="#FFFFFF",
        font=("Segoe UI Semibold", 16),
        anchor="w",
    ).grid(row=0, column=0, sticky="w")
    tk.Label(
        header,
        text="Підготовка джерел та Excel-шаблону. Період визначиться автоматично.",
        bg=colors["header"],
        fg="#D7FBF5",
        font=("Segoe UI", 9),
        anchor="w",
    ).grid(row=1, column=0, sticky="w", pady=(4, 6))
    badge_shell = tk.Frame(header, bg=colors["accent_dark"], padx=1, pady=1)
    badge = tk.Label(
        badge_shell,
        text="НАЛАШТУВАННЯ",
        bg="#053D39",
        fg="#FFFFFF",
        font=("Segoe UI", 12, "bold"),
        padx=10,
        pady=5,
    )
    badge.pack()
    badge_shell.place(relx=1.0, x=6, y=6, anchor="ne")

    body = ttk.Frame(surface, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
    body.grid(row=1, column=0, sticky="nsew")
    panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
    panel_shell.grid(row=0, column=0, sticky="we")
    container = ttk.Frame(panel_shell, style="LaunchPanel.TFrame", padding=(16, 14, 16, 12))
    container.pack(fill="both", expand=True)
    container.grid_columnconfigure(1, weight=1)

    def pick_path(var: StringVar, display_var: StringVar, entry, kind: str, title: str):
        root.lift()
        root.attributes("-topmost", True)
        root.update()
        if kind == "dir":
            p = filedialog.askdirectory(title=title, parent=root)
        else:
            p = filedialog.askopenfilename(
                title=title,
                parent=root,
                filetypes=[("Excel files", "*.xlsx;*.xlsm"), ("All files", "*.*")],
            )
        if p:
            var.set(p)
            refresh_path_entry_display(var, display_var, entry)
        root.lift()
        root.update()

    ttk.Label(container, text="Джерела", style="LaunchSection.TLabel").grid(
        row=0, column=0, columnspan=3, sticky="w"
    )

    ttk.Label(container, text="БРО-рапорти (.docx):", style="LaunchField.TLabel").grid(row=1, column=0, sticky="e", padx=8, pady=6)
    bro_entry = ttk.Entry(container, textvariable=bro_display_var, width=entry_width, state="readonly", style="Launch.TEntry")
    bro_entry.grid(row=1, column=1, sticky="we", padx=8, pady=6)
    ttk.Button(
        container,
        text="ОБРАТИ",
        width=13,
        style="LaunchBrowse.TButton",
        command=lambda: pick_path(bro_var, bro_display_var, bro_entry, "dir", "БРО-рапорти (.docx)"),
    ).grid(row=1, column=2, sticky="w", padx=8, pady=6)

    ttk.Label(container, text="ШТАТКИ (.xlsx):", style="LaunchField.TLabel").grid(row=2, column=0, sticky="e", padx=8, pady=6)
    staff_entry = ttk.Entry(container, textvariable=staff_display_var, width=entry_width, state="readonly", style="Launch.TEntry")
    staff_entry.grid(row=2, column=1, sticky="we", padx=8, pady=6)
    ttk.Button(
        container,
        text="ОБРАТИ",
        width=13,
        style="LaunchBrowse.TButton",
        command=lambda: pick_path(staff_var, staff_display_var, staff_entry, "dir", "ШТАТКИ (.xlsx)"),
    ).grid(row=2, column=2, sticky="w", padx=8, pady=6)

    ttk.Label(container, text="ЦІЛЬ (.xlsx):", style="LaunchField.TLabel").grid(row=3, column=0, sticky="e", padx=8, pady=6)
    xlsx_entry = ttk.Entry(container, textvariable=xlsx_display_var, width=entry_width, state="readonly", style="Launch.TEntry")
    xlsx_entry.grid(row=3, column=1, sticky="we", padx=8, pady=6)
    ttk.Button(
        container,
        text="ОБРАТИ",
        width=13,
        style="LaunchBrowse.TButton",
        command=lambda: pick_path(xlsx_var, xlsx_display_var, xlsx_entry, "file", "ЦІЛЬ (.xlsx)"),
    ).grid(row=3, column=2, sticky="w", padx=8, pady=6)

    bro_var.trace_add("write", lambda *_args: refresh_path_entry_display(bro_var, bro_display_var, bro_entry))
    staff_var.trace_add("write", lambda *_args: refresh_path_entry_display(staff_var, staff_display_var, staff_entry))
    xlsx_var.trace_add("write", lambda *_args: refresh_path_entry_display(xlsx_var, xlsx_display_var, xlsx_entry))
    bro_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(bro_var, bro_display_var, bro_entry))
    staff_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(staff_var, staff_display_var, staff_entry))
    xlsx_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(xlsx_var, xlsx_display_var, xlsx_entry))

    footer_line = tk.Frame(container, bg=colors["border"], height=1)
    footer_line.grid(row=4, column=0, columnspan=3, sticky="we", padx=8, pady=(12, 14))

    result = {"ok": False, "data": None}

    def validate_and_close():
        try:
            bro = bro_var.get().strip()
            staff = staff_var.get().strip()
            xlsx = xlsx_var.get().strip()

            if not xlsx:
                raise ValueError("Не вибрано файл-ціль Excel.")
            if not bro and not staff:
                raise ValueError("Потрібно вибрати хоча б одну папку: БРО або штатки.")

            bro_p = Path(bro) if bro else None
            staff_p = Path(staff) if staff else None
            xlsx_p = Path(xlsx)

            if bro_p is not None and not bro_p.exists():
                raise ValueError("Папка БРО не існує.")
            if staff_p is not None and not staff_p.exists():
                raise ValueError("Папка штаток не існує.")
            if not xlsx_p.exists():
                raise ValueError("Excel-файл не існує.")

            start_col_idx, year, month = detect_template_context(xlsx_p)
            last_day = calendar.monthrange(year, month)[1]

            result["data"] = (bro_p, staff_p, xlsx_p, start_col_idx, last_day, month, year)
            result["ok"] = True
            root.destroy()

        except Exception as ex:
            messagebox.showerror("Помилка", str(ex), parent=root)

    ttk.Button(
        container,
        text="ЗАПУСТИТИ ОБРОБКУ",
        command=validate_and_close,
        style="LaunchHero.TButton",
    ).grid(
        row=5, column=0, columnspan=3, sticky="we", padx=8, pady=(0, 4)
    )

    root.update_idletasks()
    root.geometry(f"{root.winfo_reqwidth()}x{root.winfo_reqheight()}")
    root.deiconify()
    root.lift()
    root.grab_set()
    root.focus_force()
    root.mainloop()

    if not result["ok"]:
        raise SystemExit("Скасовано. Процес зупинено.")

    return result["data"]


def backup_excel_to_trash(xlsx_path: Path):
    timestamp = datetime.now().strftime("%Y.%m.%d %H.%M")
    backup_name = f"{xlsx_path.stem} - {timestamp}{xlsx_path.suffix}"
    backup_path = xlsx_path.with_name(backup_name)

    try:
        shutil.copy2(xlsx_path, backup_path)
        send2trash(str(backup_path))
    except Exception as e:
        if _is_file_access_error(e):
            messagebox.showwarning(
                "Файл зайнятий",
                _file_access_message(_permission_error_path(e, xlsx_path), "зробити backup Excel-файлу", e),
                parent=None,
            )
            return
        messagebox.showwarning(
            "Увага",
            f"Не вдалося зробити backup Excel у кошик:\n{backup_path}\n\n{e}",
            parent=None,
        )


# ===================== НОРМАЛІЗАЦІЯ =====================

_space_re = re.compile(r"\s+")

def _fold(s: str) -> str:
    # casefold() краще за lower() для юнікоду (укр/рос/лат)
    return (s or "").casefold()

def normalize_value(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, datetime):
        if val.time().hour == 0 and val.time().minute == 0 and val.time().second == 0:
            return val.date().isoformat()
        return val.replace(microsecond=0).isoformat(sep=" ")
    if isinstance(val, date):
        return val.isoformat()
    return str(val).strip()

def normalize_name(name: str) -> str:
    if not isinstance(name, str):
        return ""
    name = name.strip()
    name = re.sub(r"\s+", " ", name)
    name = re.sub(r"[.,;]+$", "", name)
    name = re.sub(r"\([^)]*$", "", name).strip()
    return name

def norm_pib(x: Any) -> str:
    if x is None:
        return ""
    s = str(x).strip()
    s = _space_re.sub(" ", s)
    return s.upper()

def format_pib_for_sheet(pib: str) -> str:
    if pib is None:
        return ""
    s = str(pib).strip()
    s = _space_re.sub(" ", s)
    if not s:
        return ""

    def cap_token(tok: str) -> str:
        tok = tok.strip()
        if not tok:
            return tok
        # дефіси: капіталізуємо кожну частину
        parts = tok.split("-")
        out_parts = []
        for p in parts:
            p = p.strip()
            if not p:
                out_parts.append(p)
                continue
            out_parts.append(p[0].upper() + p[1:].lower())
        return "-".join(out_parts)

    return " ".join(cap_token(t) for t in s.split(" "))


# ===================== ПАРС ДАТИ З НАЗВИ ШТАТКИ =====================

_date_in_name = re.compile(r"(\d{2})\.(\d{2})\.(\d{4})")

def extract_date_from_staff_filename(path: Path) -> str:
    m = _date_in_name.search(path.name)
    if not m:
        raise ValueError(f"Не знайшов дату у назві штатки: {path.name}")
    dd, mm, yyyy = m.groups()
    return f"{yyyy}-{mm}-{dd}"

def trash_file(path: Path, label: str = ""):
    try:
        p = Path(path)
        if p.exists() and p.is_file():
            send2trash(str(p))
            log(f"[TRASH] {label} -> {p.name}")
    except Exception as e:
        log(f"[TRASH/WARN] {label} не зміг відправити в кошик: {e}")

def build_ordered_by_date_to_path_map(staff_folder: Path) -> dict[str, Path]:
    staff_folder = Path(staff_folder)

    files = sorted([
        p for p in staff_folder.iterdir()
        if p.is_file()
        and p.suffix.lower() in (".xlsx", ".xlsm")
        and not p.name.startswith("~$")
    ])

    by_date: dict[str, Path] = {}

    for p in files:
        try:
            dk = extract_date_from_staff_filename(p)
        except Exception as e:
            log(f"[STAFF/MAP SKIP] {p.name}: {e}")
            continue

        if dk not in by_date:
            by_date[dk] = p
        else:
            try:
                if p.stat().st_mtime > by_date[dk].stat().st_mtime:
                    by_date[dk] = p
            except Exception:
                pass

    if not by_date:
        raise SystemExit("Не знайшов жодної валідної штатки з датою dd.mm.yyyy у назві.")

    ordered = dict(sorted(by_date.items(), key=lambda kv: kv[0]))
    log(f"[STAFF/MAP] dates={len(ordered)}, first={next(iter(ordered))}, last={list(ordered.keys())[-1]}")
    return ordered


# ===================== (0) ОЧИСТКА A,B + ДНІ, крім жовтих заголовків =====================

CLEAR_START_ROW = 5
NO_FILL = PatternFill()

def _rgb_of_fill(cell) -> str:
    try:
        fill = cell.fill
        if not fill or fill.patternType is None:
            return ""
        fg = getattr(fill, "fgColor", None)
        if not fg:
            return ""
        rgb = getattr(fg, "rgb", None)
        if not rgb:
            return ""
        return str(rgb).upper()
    except Exception:
        return ""

def _is_yellow_header(cell_a) -> bool:
    rgb = _rgb_of_fill(cell_a)
    if not rgb:
        return False
    return rgb.endswith("FFFF00")

def apply_body_formatting_to_sheet(
    ws,
    start_row: int,
    start_col_idx: int,
    last_day: int,
    extra_cols: int = 5,
):
    end_day_col = start_col_idx + last_day - 1 + extra_cols
    font_name = "Times New Roman"
    font_size = 12
    border_all = OUTER_BORDER

    align_left = Alignment(horizontal="left", vertical="center", wrap_text=False)
    align_left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=False)
    align_center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)

    max_row = ws.max_row
    for r in range(start_row, max_row + 1):
        cell_a = ws.cell(row=r, column=1)
        if _is_yellow_header(cell_a):
            continue

        for c in range(1, end_day_col + 1):
            cell = ws.cell(row=r, column=c)
            if cell.value is not None and str(cell.value) != "":
                cur_font = copy(cell.font) if cell.font else Font()
                cur_font.name = font_name
                cur_font.size = font_size
                cell.font = cur_font
            cell.border = border_all

            if c == 2:
                cell.alignment = align_center_wrap
            elif c in (1, 3):
                cell.alignment = align_left_wrap
            elif start_col_idx <= c <= end_day_col:
                cell.alignment = align_center
            else:
                cell.alignment = align_left


def restore_yellow_header_rows(ws, max_col_letter: str = "AO") -> int:
    max_col_idx = column_index_from_string(max_col_letter)
    restored = 0

    for r in range(1, ws.max_row + 1):
        if not _is_yellow_header(ws.cell(row=r, column=1)):
            continue

        for c in range(1, max_col_idx + 1):
            ws.cell(row=r, column=c).fill = FILL_YELLOW
        restored += 1

    return restored


def clear_official_sheets_zero_stage(
    target_xlsx: Path,
    start_col_idx: int,
    last_day: int,
    sheets: list[str] = None,
    start_row: int = 5,
) -> dict:
    if sheets is None:
        sheets = ["100 000", "70 000"]

    wb = load_workbook(target_xlsx)
    last_col_idx = start_col_idx + last_day - 1
    note_col_idx = column_index_from_string("AN")

    cleared_values = 0
    cleared_fills = 0
    skipped_headers = 0

    cols_full_clear = sorted({1, 2, note_col_idx, *range(start_col_idx, last_col_idx + 1)})
    cols_fill_only = [3, 4, 5]

    def _find_last_text_row_local(ws, col_letter: str, start_row_local: int) -> int:
        last = start_row_local - 1
        for r in range(start_row_local, ws.max_row + 1):
            v = ws[f"{col_letter}{r}"].value
            if v is not None and str(v).strip() != "":
                last = r
        return last

    for sh in sheets:
        if sh not in wb.sheetnames:
            continue
        ws = wb[sh]
        max_row = ws.max_row
        last_used = _find_last_text_row_local(ws, "E", start_row)

        for r in range(start_row, max_row + 1):
            cell_a = ws.cell(row=r, column=1)

            if _is_yellow_header(cell_a):
                skipped_headers += 1
                continue

            for c in cols_fill_only:
                cell = ws.cell(row=r, column=c)
                if cell.fill and cell.fill.patternType is not None:
                    cell.fill = NO_FILL
                    cleared_fills += 1

            for c in cols_full_clear:
                cell = ws.cell(row=r, column=c)

                if cell.value is not None and str(cell.value) != "":
                    cell.value = None
                    cleared_values += 1
                else:
                    cell.value = None

                if cell.fill and cell.fill.patternType is not None:
                    cell.fill = NO_FILL
                    cleared_fills += 1

        apply_body_formatting_to_sheet(
            ws=ws,
            start_row=start_row,
            start_col_idx=start_col_idx,
            last_day=last_day,
            extra_cols=5,
        )

        if last_used >= start_row:
            end_row = last_used + 15
            for r in range(start_row, end_row + 1):
                ws.row_dimensions[r].height = 28

    save_target_workbook(wb, target_xlsx)
    wb.close()

    log(f"[CLEAR] skipped_headers={skipped_headers}, cleared_values={cleared_values}, cleared_fills={cleared_fills}")
    return {
        "skipped_headers": skipped_headers,
        "cleared_values": cleared_values,
        "cleared_fills": cleared_fills,
    }

def apply_body_formatting_to_official_sheets(
    target_xlsx: Path,
    start_col_idx: int,
    last_day: int,
    sheets: list[str] = None,
    start_row: int = 5,
    extra_cols: int = 5,
):
    if sheets is None:
        sheets = ["100 000", "70 000"]

    wb = load_workbook(target_xlsx)
    restored_headers: dict[str, int] = {}
    for sh in sheets:
        if sh not in wb.sheetnames:
            continue
        ws = wb[sh]
        apply_body_formatting_to_sheet(
            ws=ws,
            start_row=start_row,
            start_col_idx=start_col_idx,
            last_day=last_day,
            extra_cols=extra_cols,
        )
        restored_headers[sh] = restore_yellow_header_rows(ws, max_col_letter="AO")

    save_target_workbook(wb, target_xlsx)
    wb.close()
    log(f"[FORMAT] restored yellow header rows: {restored_headers}")


# ===================== Колонка B: джерела з БРО =====================

def clean_filename_for_source(filename: str) -> str:
    name, _ = os.path.splitext(filename)
    s = name.casefold()

    # --- прибираємо службові маркери ---
    s = re.sub(r"\bбро\b", " ", s, flags=re.IGNORECASE)
    s = re.sub(r"\b3\s*мб\b", " ", s, flags=re.IGNORECASE)

    # --- прибираємо дати типу 2025-01-12, 01.02.2026 і т.п. ---
    s = re.sub(r"\b\d{4}-\d{2}-\d{2}\b", " ", s)
    s = re.sub(r"\b\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b", " ", s)

    # --- прибираємо роки 20xx будь-де ---
    s = re.sub(r"20\d{2}", " ", s)

    # --- прибираємо всі розділювачі ---
    s = re.sub(r"[+\-_=.,;:(){}\[\]<>\"'`~!@#$%^&*/\\|№]+", " ", s)

    # --- тепер головне: ---
    # прибрати ЦИФРИ, ЯКІ ЙДУТЬ ПІСЛЯ ЛІТЕР
    #   7мр2026 -> 7мр
    #   рбак01  -> рбак
    s = re.sub(r"([а-яіїєґa-z]+)\d+", r"\1", s, flags=re.IGNORECASE)

    # --- прибрати окремі числа, що лишилися ---
    s = re.sub(r"\b\d+\b", " ", s)

    # --- прибрати всі пробіли ---
    s = re.sub(r"\s+", "", s).strip()

    return s

def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    texts = []

    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)

    return "\n".join(texts)

def write_docx_sources_to_column_b(
    br_folder: Path,
    target_xlsx: Path,
    sheets: list[str] = None,
    pib_col: str = "D",
    dst_col: str = "B",
    start_row: int = 2,
) -> int:
    if sheets is None:
        sheets = ["100 000", "70 000"]

    br_folder = Path(br_folder)
    target_xlsx = Path(target_xlsx)

    if not br_folder.exists():
        print(f"[SRC->B] br_folder не існує: {br_folder}")
        return 0
    if not target_xlsx.exists():
        print(f"[SRC->B] target_xlsx не існує: {target_xlsx}")
        return 0

    docx_files = sorted([
        p for p in br_folder.iterdir()
        if p.is_file() and p.suffix.lower() == ".docx" and not p.name.startswith("~$")
    ])
    if not docx_files:
        print(f"[SRC->B] У папці БР нема .docx: {br_folder}")
        return 0

    print(f"[SRC->B] Знайдено .docx: {len(docx_files)} у {br_folder}")

    docs_texts: dict[str, str] = {}
    for p in docx_files:
        try:
            docs_texts[p.name] = read_docx_text(p)
        except Exception as e:
            print(f"[SRC->B] Не прочитав '{p.name}', скіп. Причина: {e}")
            docs_texts[p.name] = ""

    wb = load_workbook(target_xlsx)
    matches_count = 0

    for sh in sheets:
        if sh not in wb.sheetnames:
            print(f"[SRC->B] Нема листа '{sh}' – скіп.")
            continue

        ws = wb[sh]
        max_row = ws.max_row

        last_row = start_row - 1
        for r in range(start_row, max_row + 1):
            v = ws[f"{pib_col}{r}"].value
            if v is not None and str(v).strip() != "":
                last_row = r

        if last_row < start_row:
            print(f"[SRC->B] '{sh}': немає ПІБ у діапазоні з {start_row} (колонка {pib_col}).")
            continue

        print(f"[SRC->B] '{sh}': обробляю рядки {start_row}..{last_row} (ПІБ={pib_col} -> B={dst_col})")

        for row in range(start_row, last_row + 1):
            pib_value = ws[f"{pib_col}{row}"].value
            if pib_value is None or str(pib_value).strip() == "":
                continue

            pib = str(pib_value).strip()
            pib_cf = _fold(pib)

            found_files = []

            for filename, text in docs_texts.items():
                try:
                    if not pib_cf or not text:
                        continue

                    # case-insensitive substring
                    if pib_cf in _fold(text):
                        cleaned = clean_filename_for_source(filename)
                        if cleaned:
                            found_files.append(cleaned)
                except Exception as e:
                    print(f"[SRC->B] Помилка пошуку у '{filename}', скіп. Причина: {e}")
                    continue

            if found_files:
                dst_cell = f"{dst_col}{row}"
                new_unique = []
                for item in found_files:
                    if item not in new_unique:
                        new_unique.append(item)

                if new_unique:
                    ws[dst_cell].value = "; ".join(new_unique).lower()
                    # BRO-джерела мають бути чорним, навіть якщо раніше було dark-blue preload.
                    b_font = copy(ws[dst_cell].font) if ws[dst_cell].font else Font()
                    b_font.color = Color(rgb="FF000000")
                    ws[dst_cell].font = b_font
                    matches_count += len(new_unique)

    save_target_workbook(wb, target_xlsx)
    wb.close()

    print(f"[SRC->B] Додано нових збігів у B: {matches_count}")
    return matches_count


# ===================== СИНХ ЗВАНЬ: тільки колонка C (без посади в B) =====================

FINAL_SOURCE_SHEETS = ["100 000", "70 000"]

FINAL_SOURCE_PIB_COL = "E"
FINAL_SOURCE_START_ROW = 5

UNOFFICIAL_SYNC_SHEETS = ["100 000", "70 000"]
UNOFFICIAL_SYNC_START_ROW = FINAL_SOURCE_START_ROW
UNOFFICIAL_SYNC_COL_POS  = "C"
UNOFFICIAL_SYNC_COL_NAME = "E"
UNOFFICIAL_SYNC_COL_RANK = "D"

# Штатка-джерело:
STAFF_SRC_COL_NAME = "L"  # ПІБ
STAFF_SRC_COL_RANK = "K"  # звання
STAFF_SRC_COL_POS  = "F"  # посада (1-й абзац)

def build_staff_json(staff_path: Path, json_path: Path) -> Path:
    wb = load_workbook(staff_path, data_only=True)
    ws = get_staff_sheet(wb)

    items: list[dict] = []
    for row in range(1, ws.max_row + 1):
        name_val = ws[f"{STAFF_SRC_COL_NAME}{row}"].value
        if name_val is None:
            continue

        norm_name = normalize_name(str(name_val))
        if not norm_name:
            continue

        rank_val = ws[f"{STAFF_SRC_COL_RANK}{row}"].value
        rank = "" if rank_val is None else str(rank_val).strip()

        pos_val = ws[f"{STAFF_SRC_COL_POS}{row}"].value
        pos_raw = "" if pos_val is None else str(pos_val)

        # беремо лише перший рядок/абзац
        pos_raw = pos_raw.replace("\r\n", "\n").replace("\r", "\n")
        position = pos_raw.split("\n")[0].strip()

        items.append({
            "name": str(name_val).strip(),
            "name_norm": norm_name.lower(),
            "rank": rank,
            "position": position,
        })

    wb.close()

    data = {"source_file": str(staff_path), "rows": items}

    json_path = json_path.resolve()
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"[SYNC] staff_rank_pos_map.json: {json_path} (записів: {len(items)})")
    return json_path

def load_staff_rank_mapping_from_json(json_path: Path) -> dict[str, str]:
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    mapping: dict[str, str] = {}
    for row in data.get("rows", []):
        name_norm = (row.get("name_norm", "") or "").strip().lower()
        rank = row.get("rank", "") or ""
        if not name_norm:
            continue
        mapping[name_norm] = rank

    return mapping

def sync_rank_only(
    staff_mapping: dict[str, str],
    target_path: Path,
    sheet_name: str,
    start_row: int,
    col_name_letter: str,
    col_rank_letter: str
):
    wb = load_workbook(target_path)
    if sheet_name not in wb.sheetnames:
        wb.close()
        raise ValueError(f"Аркуш '{sheet_name}' не знайдено в {target_path.name}")

    ws = wb[sheet_name]

    col_name_idx = column_index_from_string(col_name_letter)
    col_rank_idx = column_index_from_string(col_rank_letter)

    processed = 0
    matched = 0
    not_found = 0

    for row in range(start_row, ws.max_row + 1):
        name_val = ws.cell(row=row, column=col_name_idx).value
        if name_val is None:
            continue

        norm_name = normalize_name(str(name_val)).lower()
        if not norm_name:
            continue

        processed += 1
        rank = staff_mapping.get(norm_name)
        if rank is None:
            not_found += 1
            continue

        ws.cell(row=row, column=col_rank_idx).value = rank
        matched += 1

    save_target_workbook(wb, target_path)
    wb.close()

    return processed, matched, not_found

def pick_last_staff_xlsx(staff_file_map: dict[str, Path]) -> Path | None:
    if not staff_file_map:
        return None
    last_date = sorted(staff_file_map.keys())[-1]
    return staff_file_map.get(last_date)

def load_staff_rank_pos_mapping_from_json(json_path: Path) -> dict[str, tuple[str, str]]:
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    mapping: dict[str, tuple[str, str]] = {}
    for row in data.get("rows", []):
        name_norm = (row.get("name_norm", "") or "").strip().lower()
        if not name_norm:
            continue
        rank = (row.get("rank", "") or "").strip()
        pos  = (row.get("position", "") or "").strip()
        mapping[name_norm] = (rank, pos)

    return mapping

def sync_rank_and_pos(
    staff_mapping: dict[str, tuple[str, str]],
    target_path: Path,
    sheet_name: str,
    start_row: int,
    col_name_letter: str,  # E
    col_rank_letter: str,  # D
    col_pos_letter: str,   # C
):
    wb = load_workbook(target_path)
    if sheet_name not in wb.sheetnames:
        wb.close()
        raise ValueError(f"Аркуш '{sheet_name}' не знайдено в {target_path.name}")

    ws = wb[sheet_name]

    col_name_idx = column_index_from_string(col_name_letter)
    col_rank_idx = column_index_from_string(col_rank_letter)
    col_pos_idx  = column_index_from_string(col_pos_letter)

    processed = matched = not_found = 0

    for row in range(start_row, ws.max_row + 1):
        name_val = ws.cell(row=row, column=col_name_idx).value
        if not name_val:
            continue

        norm_name = normalize_name(str(name_val)).lower()
        if not norm_name:
            continue

        processed += 1
        data = staff_mapping.get(norm_name)
        if not data:
            not_found += 1
            continue

        rank, pos = data
        ws.cell(row=row, column=col_rank_idx).value = rank
        ws.cell(row=row, column=col_pos_idx).value = pos
        matched += 1

    save_target_workbook(wb, target_path)
    wb.close()

    return processed, matched, not_found

def sync_rank_pos_after_push(staff_file_map: dict[str, Path], target_xlsx: Path):
    staff_xlsx = pick_last_staff_xlsx(staff_file_map)
    if not staff_xlsx or not staff_xlsx.exists():
        print("[SYNC] Не знайшов штатки для синхронізації звань/посад.")
        return

    tmp_json = staff_xlsx.with_name("_staff_rank_pos_map_tmp.json")

    try:
        build_staff_json(staff_xlsx, tmp_json)  # має писати і rank і position
        staff_mapping = load_staff_rank_pos_mapping_from_json(tmp_json)

        if not staff_mapping:
            print("[SYNC] staff_mapping порожній. Скіп.")
            return

        for sh in UNOFFICIAL_SYNC_SHEETS:
            try:
                p, m, nf = sync_rank_and_pos(
                    staff_mapping,
                    target_xlsx,
                    sheet_name=sh,
                    start_row=UNOFFICIAL_SYNC_START_ROW,
                    col_name_letter=UNOFFICIAL_SYNC_COL_NAME,  # "E"
                    col_rank_letter=UNOFFICIAL_SYNC_COL_RANK,  # "D"
                    col_pos_letter=UNOFFICIAL_SYNC_COL_POS,    # "C"
                )
                print(f"[SYNC RANK+POS] {sh}: processed={p}, matched={m}, not_found={nf}")
            except Exception as e:
                print(f"[SYNC RANK+POS] Помилка на аркуші {sh}: {e}")

    finally:
        try:
            if tmp_json.exists():
                tmp_json.unlink()
        except Exception:
            pass

# ===================== БРО: РАПОРТИ (multi-day in one docx) =====================

BRO_JSON_NAME = "bro_by_date.json"

FILL_GREEN = PatternFill(fill_type="solid", fgColor="FF00B050")
FILL_DARK_GREEN = PatternFill(fill_type="solid", fgColor="FF006100")
FILL_OLIVE = PatternFill(fill_type="solid", fgColor="FFC4D79B")
FILL_YELLOW = PatternFill(fill_type="solid", fgColor="FFFFFF00")
FILL_PALE_GREEN = PatternFill(fill_type="solid", fgColor="FFE2F0D9")


_bro_date_iso = re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b")
_bro_date_ua  = re.compile(r"\b(\d{2})\.(\d{2})\.(\d{4})\b")

_bro_rate70  = re.compile(r"\b70\s*тис", re.IGNORECASE)
_bro_rate100 = re.compile(r"\b100\s*тис", re.IGNORECASE)
_bro_rate30  = re.compile(r"\b30\s*тис", re.IGNORECASE)

# Стоп-рядок кінця списків (перед підписом)
_bro_commander_line = re.compile(r"\bкомандир(а)?\b", re.IGNORECASE)

_bro_name3 = re.compile(
    r"\b([А-ЯІЇЄҐ][А-ЯІЇЄҐа-яіїєґ'’\-]+)\s+([А-ЯІЇЄҐ][А-ЯІЇЄҐа-яіїєґ'’\-]+)\s+([А-ЯІЇЄҐ][А-ЯІЇЄҐа-яіїєґ'’\-]+)\b"
)

def iter_docx_lines(path: Path) -> List[str]:
    doc = Document(str(path))
    lines: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    lines.append(t)

    return lines

def _to_iso_from_ua(dd: str, mm: str, yyyy: str) -> str:
    return f"{yyyy}-{mm}-{dd}"

def _to_iso_from_iso(y: str, mm: str, dd: str) -> str:
    return f"{y}-{mm}-{dd}"

def detect_day_header_date(line: str) -> str | None:
    if not line:
        return None

    s = (line or "").strip()
    if not s:
        return None

    # прибираємо "дрібний шум" на початку (пробіли, №, :, -, # тощо)
    s = re.sub(r"^[\s#:№\-\u2013\u2014]*", "", s)

    # 1) ISO на початку
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        y, mm, dd = m.groups()
        return _to_iso_from_iso(y, mm, dd)

    # 2) UA dd.mm.yyyy на початку (дозволяємо крапку після року та інший хвіст)
    m = re.match(r"^(\d{2})\.(\d{2})\.(\d{4})", s)
    if m:
        dd, mm, y = m.groups()
        return _to_iso_from_ua(dd, mm, y)

    return None

def detect_rate_from_line(line: str) -> str | None:
    s = (line or "").lower()
    if _bro_rate100.search(s):
        return "rate_100"
    if _bro_rate70.search(s):
        return "rate_70"
    if _bro_rate30.search(s):
        return "rate_30"
    return None

def extract_names_from_line(line: str) -> List[str]:
    out = []
    if not line:
        return out

    cleaned = re.sub(r"\([^)]*\)", " ", line)
    cleaned = _space_re.sub(" ", cleaned).strip()

    for s1, s2, s3 in _bro_name3.findall(cleaned):
        pib = f"{s1} {s2} {s3}".strip()
        out.append(pib)
    return out

def safe_call(label: str, path: Path, fn, default=None):
    try:
        return fn()
    except Exception as e:
        log(f"[SKIP/{label}] {path.name}: {type(e).__name__}: {e}")
        return default

def build_bro_by_date_json(bro_folder: Path, out_json: Path) -> Path:
    files = sorted([
        p for p in bro_folder.iterdir()
        if p.is_file() and p.suffix.lower() == ".docx" and not p.name.startswith("~$")
    ])
    if not files:
        raise SystemExit("У папці з БРО не знайшов .docx")

    by_date: Dict[str, Dict[str, List[str]]] = {}

    # тимчасовий буфер поточного "добового" блоку (дата в кінці)
    pending = {"rate_70": [], "rate_100": [], "rate_30": []}
    pending_rate: str | None = None

    def pending_clear():
        pending["rate_70"].clear()
        pending["rate_100"].clear()
        pending["rate_30"].clear()

    def pending_is_empty() -> bool:
        return (not pending["rate_70"]) and (not pending["rate_100"]) and (not pending["rate_30"])

    def commit_pending_to_date(date_key: str):
        """Записує pending у by_date[date_key] (без дублів), потім чистить pending."""
        if not date_key:
            return
        if pending_is_empty():
            return

        day = by_date.setdefault(date_key, {"rate_70": [], "rate_100": [], "rate_30": []})

        for rk in ("rate_70", "rate_100", "rate_30"):
            existing = set(map(norm_pib, day[rk]))
            for name in pending[rk]:
                k = norm_pib(name)
                if k and k not in existing:
                    day[rk].append(name)
                    existing.add(k)

        pending_clear()

    for f in files:
        lines = safe_call("BRO/READ", f, lambda: iter_docx_lines(f), default=None)
        if not lines:
            continue

        pending_clear()
        pending_rate = None

        for ln in lines:
            s = (ln or "").strip()
            if not s:
                continue

            # 1) якщо це "абзац = тільки дата" -> це кінець добового рапорту
            dkey = detect_day_header_date(s)   # твоя функція, але фактично "bare date"
            if dkey:
                # дата закриває добовий рапорт -> комітимо все, що назбирали
                commit_pending_to_date(dkey)
                pending_rate = None
                continue

            # 2) ставка
            r = detect_rate_from_line(s)
            if r:
                pending_rate = r
                continue

            # 3) імена (тільки якщо ми всередині rate-блоку)
            if not pending_rate:
                continue

            names = extract_names_from_line(s)
            if not names:
                continue

            existing = set(map(norm_pib, pending[pending_rate]))
            for n in names:
                k = norm_pib(n)
                if k and k not in existing:
                    pending[pending_rate].append(n)
                    existing.add(k)

        if not pending_is_empty():
            log(f"[BRO/WARN] {f.name}: назбирав names у pending, але не побачив 'абзац=дата' для commit.")

        log(f"[BRO/OK] {f.name}: lines={len(lines)}")

    if not by_date:
        raise SystemExit("Не зібрав жодного валідного дня з БРО (не знайшов абзаців-дати).")

    out_json.parent.mkdir(parents=True, exist_ok=True)
    out_json.write_text(json.dumps({
        "source_folder": str(bro_folder),
        "files_processed": len(files),
        "dates": sorted(by_date.keys()),
        "by_date": by_date,
    }, ensure_ascii=False, indent=2), encoding="utf-8")

    log(f"[BRO] JSON: {out_json} | днів: {len(by_date)}")
    return out_json

# ===================== ЛИСТИ 100 000 / 70 000: нанесення БРО =====================
NO_FILL = PatternFill()
RED_FONT = Font(color="FF8B0000")

DARK_BLUE_FONT_COLOR = "FF00008B"
_THIN = Side(style="thin", color="FF000000")
OUTER_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
LEFT_CENTER_ALIGN = Alignment(horizontal="left", vertical="center", wrap_text=False)
CENTER_CENTER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=False)
ROW_MOVE_LAST_COL_LETTER = "AO"
ROW_MOVE_SKIP_COLS = {
    column_index_from_string("AK"),
    column_index_from_string("AO"),
}

def find_last_text_row(ws, col_letter: str, start_row: int) -> int:
    last = start_row - 1
    for r in range(start_row, ws.max_row + 1):
        v = ws[f"{col_letter}{r}"].value
        if v is not None and str(v).strip() != "":
            last = r
    return last

def _cell_snapshot(cell):
    return (
        cell.value,
        copy(cell.font) if cell.font else None,
        copy(cell.fill) if cell.fill else None,
        copy(cell.border) if cell.border else None,
        copy(cell.alignment) if cell.alignment else None,
        cell.number_format,
        copy(cell.protection) if cell.protection else None,
        copy(cell.comment) if cell.comment else None,
        copy(getattr(cell, "_style", None)) if getattr(cell, "_style", None) else None,
    )

def _cell_restore(cell, snap):
    val, font, fill, border, align, numfmt, prot, comm, _style = snap
    cell.value = val
    if font is not None: cell.font = font
    if fill is not None: cell.fill = fill
    if border is not None: cell.border = border
    if align is not None: cell.alignment = align
    if numfmt is not None: cell.number_format = numfmt
    if prot is not None: cell.protection = prot
    cell.comment = comm
    if _style is not None: cell._style = _style

def _clear_cell_payload(cell) -> None:
    cell.value = None
    cell.comment = None
    cell.fill = NO_FILL

def _row_snapshot(ws, row: int, max_col: int, skip_cols: set[int] | None = None) -> list[tuple | None]:
    skip_cols = skip_cols or ROW_MOVE_SKIP_COLS
    snaps: list[tuple | None] = [None] * max_col
    for c in range(1, max_col + 1):
        if c in skip_cols:
            continue
        snaps[c - 1] = _cell_snapshot(ws.cell(row=row, column=c))
    return snaps

def _restore_row_snapshot(
    ws,
    row: int,
    snaps: list[tuple | None],
    max_col: int,
    skip_cols: set[int] | None = None,
) -> None:
    skip_cols = skip_cols or ROW_MOVE_SKIP_COLS
    for c in range(1, max_col + 1):
        cell = ws.cell(row=row, column=c)
        if c in skip_cols:
            _clear_cell_payload(cell)
            continue
        snap = snaps[c - 1] if c - 1 < len(snaps) else None
        if snap is None:
            _clear_cell_payload(cell)
        else:
            _cell_restore(cell, snap)

def _clear_row_payload(
    ws,
    row: int,
    max_col: int,
    skip_cols: set[int] | None = None,
) -> None:
    for c in range(1, max_col + 1):
        _clear_cell_payload(ws.cell(row=row, column=c))

def _manual_insert_row_payload(
    ws,
    insert_row: int,
    last_used_row: int,
    max_col: int,
    skip_cols: set[int] | None = None,
) -> None:
    skip_cols = skip_cols or ROW_MOVE_SKIP_COLS
    if insert_row <= last_used_row:
        for r in range(last_used_row, insert_row - 1, -1):
            snaps = _row_snapshot(ws, r, max_col, skip_cols)
            _restore_row_snapshot(ws, r + 1, snaps, max_col, skip_cols)
    _clear_row_payload(ws, insert_row, max_col, skip_cols)

def _manual_delete_row_payload(
    ws,
    delete_row: int,
    last_used_row: int,
    max_col: int,
    skip_cols: set[int] | None = None,
) -> int:
    skip_cols = skip_cols or ROW_MOVE_SKIP_COLS
    if delete_row > last_used_row:
        _clear_row_payload(ws, delete_row, max_col, skip_cols)
        return last_used_row

    for r in range(delete_row, last_used_row):
        snaps = _row_snapshot(ws, r + 1, max_col, skip_cols)
        _restore_row_snapshot(ws, r, snaps, max_col, skip_cols)
    _clear_row_payload(ws, last_used_row, max_col, skip_cols)
    return last_used_row - 1

# ---------- "заштатні" cutoff ----------

def find_prikom_cutoff_row(ws, start_row: int, marker_col_letter: str = "D", marker: str = "заштатні") -> int | None:
    m = marker.casefold()
    col_idx = column_index_from_string(marker_col_letter)
    for r in range(start_row, ws.max_row + 1):
        v = ws.cell(row=r, column=col_idx).value
        if v and m in str(v).strip().casefold():
            return r
    return None

def find_named_dead_zone_rows(
    ws,
    start_row: int,
    marker_col_letter: str = "D",
    marker: str = "залучені",
    pib_col_letter: str = FINAL_SOURCE_PIB_COL,
) -> set[int]:
    marker_idx = column_index_from_string(marker_col_letter)
    pib_col_idx = column_index_from_string(pib_col_letter)
    marker_fold = marker.casefold()
    header_row = None

    for r in range(start_row, ws.max_row + 1):
        v = ws.cell(row=r, column=marker_idx).value
        if v and marker_fold in str(v).strip().casefold():
            header_row = r
            break

    if header_row is None:
        return set()

    rows = {header_row}
    for r in range(header_row + 1, ws.max_row + 1):
        pib_val = ws.cell(row=r, column=pib_col_idx).value
        if pib_val is None or str(pib_val).strip() == "":
            break
        rows.add(r)

    return rows

# ---------- read unit/order from LAST staff ----------

def build_last_staff_unit_maps(
    staff_xlsx: Path,
    name_col_letter: str = "L",   # ПІБ у штатці
    unit_col_letter: str = "B",   # підрозділ у штатці (як ти сказав)
    start_row: int = 3,
) -> tuple[dict[str, str], dict[str, int], dict[str, list[str]]]:
    wb = load_workbook(staff_xlsx, data_only=True)
    ws = get_staff_sheet(wb)

    name_col_idx = column_index_from_string(name_col_letter)
    unit_col_idx = column_index_from_string(unit_col_letter)

    unit_by_pib: dict[str, str] = {}
    unit_to_ordered_pibs: dict[str, list[str]] = {}

    current_unit = ""
    for r in range(start_row, ws.max_row + 1):
        unit_val = ws.cell(row=r, column=unit_col_idx).value
        if unit_val is not None and str(unit_val).strip() != "":
            # якщо в штатці unit дублюється на кожному рядку або йде блоками,
            # ми просто оновлюємо current_unit кожного разу коли не пусто.
            current_unit = str(unit_val).strip()

        name_val = ws.cell(row=r, column=name_col_idx).value
        k = norm_pib(name_val)
        if not k:
            continue

        u = (current_unit or "").strip()
        if not u:
            # якщо unit не заданий/порожній — не можемо прив’язати до підрозділу
            continue

        u_norm = re.sub(r"\s+", "", u.casefold())  # прибираємо пробіли, casefold
        unit_by_pib[k] = u_norm
        unit_to_ordered_pibs.setdefault(u_norm, []).append(k)

    wb.close()

    # order_in_unit: позиція в межах unit
    order_in_unit: dict[str, int] = {}
    for u_norm, lst in unit_to_ordered_pibs.items():
        seen = set()
        cleaned = []
        for k in lst:
            if k in seen:
                continue
            seen.add(k)
            cleaned.append(k)
        unit_to_ordered_pibs[u_norm] = cleaned
        for i, k in enumerate(cleaned):
            order_in_unit[k] = i

    return unit_by_pib, order_in_unit, unit_to_ordered_pibs

def build_last_staff_unit_initials_map(
    staff_xlsx: Path,
    name_col_letter: str = "L",   # ПІБ у штатці
    unit_col_letter: str = "B",   # ініціал підрозділу у штатці
    start_row: int = 3,
) -> dict[str, str]:
    wb = load_workbook(staff_xlsx, data_only=True)
    ws = get_staff_sheet(wb)

    name_col_idx = column_index_from_string(name_col_letter)
    unit_col_idx = column_index_from_string(unit_col_letter)

    unit_initial_by_pib: dict[str, str] = {}
    current_unit = ""

    try:
        for r in range(start_row, ws.max_row + 1):
            unit_val = ws.cell(row=r, column=unit_col_idx).value
            if unit_val is not None and str(unit_val).strip() != "":
                current_unit = str(unit_val).strip()

            k = norm_pib(ws.cell(row=r, column=name_col_idx).value)
            if not k or not current_unit:
                continue

            unit_initial_by_pib[k] = current_unit
    finally:
        wb.close()

    return unit_initial_by_pib

def preload_staff_unit_initials_to_column_b(
    target_xlsx: Path,
    staff_xlsx_last: Path,
    sheets: list[str] = None,
    start_row: int = FINAL_SOURCE_START_ROW,
    pib_col_letter: str = FINAL_SOURCE_PIB_COL,  # E
    dst_col_letter: str = "B",
    backup_col_letter: str | None = "AL",
    staff_name_col_letter: str = "L",
    staff_unit_col_letter: str = "B",
    staff_start_row: int = 3,
    font_color_rgb: str = DARK_BLUE_FONT_COLOR,
) -> dict[str, int]:
    if sheets is None:
        sheets = ["100 000", "70 000"]

    unit_initial_by_pib = build_last_staff_unit_initials_map(
        staff_xlsx=staff_xlsx_last,
        name_col_letter=staff_name_col_letter,
        unit_col_letter=staff_unit_col_letter,
        start_row=staff_start_row,
    )
    log(f"[UNIT->B] staff map size={len(unit_initial_by_pib)}")

    pib_col_idx = column_index_from_string(pib_col_letter)
    dst_col_idx = column_index_from_string(dst_col_letter)
    backup_col_idx = column_index_from_string(backup_col_letter) if backup_col_letter else None
    wb = load_workbook(target_xlsx)
    stats: dict[str, int] = {}

    try:
        for sh in sheets:
            if sh not in wb.sheetnames:
                stats[sh] = 0
                continue

            ws = wb[sh]
            if backup_col_idx is not None and backup_col_idx != dst_col_idx:
                for r in range(start_row, ws.max_row + 1):
                    ws.cell(row=r, column=backup_col_idx).value = None
            last_row = find_last_text_row(ws, pib_col_letter, start_row)
            if last_row < start_row:
                stats[sh] = 0
                continue

            written = 0
            for r in range(start_row, last_row + 1):
                pib_key = norm_pib(ws.cell(row=r, column=pib_col_idx).value)
                if not pib_key:
                    continue

                unit_initial = unit_initial_by_pib.get(pib_key, "")
                if not unit_initial:
                    continue

                dst_cell = ws.cell(row=r, column=dst_col_idx)
                dst_cell.value = unit_initial
                if backup_col_idx is not None and backup_col_idx != dst_col_idx:
                    ws.cell(row=r, column=backup_col_idx).value = unit_initial
                f = copy(dst_cell.font) if dst_cell.font else Font()
                f.color = Color(rgb=font_color_rgb)
                dst_cell.font = f
                written += 1

            stats[sh] = written

        save_target_workbook(wb, target_xlsx)
        return stats
    finally:
        wb.close()

# ---------- reorder внутри блоків підрозділів ----------

def reorder_by_unit_blocks_from_last_staff(
    target_xlsx: Path,
    staff_xlsx_last: Path,                 # остання штатку
    sheets: list[str] = None,
    start_row: int = FINAL_SOURCE_START_ROW,
    pib_col_letter: str = FINAL_SOURCE_PIB_COL,  # E
    prikom_col_letter: str = "D",          # де шукаємо "заштатні" на офіційному листі
    staff_name_col_letter: str = "L",      # ПІБ у штатці
    staff_unit_col_letter: str = "B",      # підрозділ у штатці
    staff_start_row: int = 3,
    start_col_idx: int | None = None,
    last_day: int | None = None,
    extra_cols: int = 5,
    force_last_col_letter: str | None = ROW_MOVE_LAST_COL_LETTER,
) -> dict[str, int]:

    if sheets is None:
        sheets = ["100 000", "70 000"]

    pib_col_idx = column_index_from_string(pib_col_letter)

    # межа колонок для snapshot (A..grid+extra або до force_last_col).
    # AK/AO містять формули з Excel metadata, тому перенос рядків їх не копіює.
    force_last_col_idx = column_index_from_string(force_last_col_letter) if force_last_col_letter else None
    last_col_idx = None
    if start_col_idx is not None and last_day is not None:
        calc_last = start_col_idx + last_day - 1 + int(extra_cols)
        last_col_idx = max(pib_col_idx, calc_last)
        if force_last_col_idx is not None:
            last_col_idx = max(last_col_idx, force_last_col_idx)
    else:
        last_col_idx = force_last_col_idx  # може бути None

    # читаємо мапи з останньої штатки
    unit_by_pib, order_in_unit, unit_to_ordered = build_last_staff_unit_maps(
        staff_xlsx_last,
        name_col_letter=staff_name_col_letter,
        unit_col_letter=staff_unit_col_letter,
        start_row=staff_start_row,
    )
    log(f"[REORDER/UNIT] staff_map: pibs={len(unit_by_pib)}, units={len(unit_to_ordered)} (unit_col={staff_unit_col_letter})")

    wb = load_workbook(target_xlsx)
    moved_counts: dict[str, int] = {}

    def _lcs_keys(a: list[str], b: list[str]) -> list[str]:
        if not a or not b:
            return []
        n, m = len(a), len(b)
        dp = [[0] * (m + 1) for _ in range(n + 1)]
        for i in range(n - 1, -1, -1):
            for j in range(m - 1, -1, -1):
                if a[i] == b[j]:
                    dp[i][j] = dp[i + 1][j + 1] + 1
                else:
                    dp[i][j] = dp[i + 1][j] if dp[i + 1][j] >= dp[i][j + 1] else dp[i][j + 1]
        i = 0
        j = 0
        out = []
        while i < n and j < m:
            if a[i] == b[j]:
                out.append(a[i])
                i += 1
                j += 1
            elif dp[i + 1][j] >= dp[i][j + 1]:
                i += 1
            else:
                j += 1
        return out

    try:
        for sh in sheets:
            if sh not in wb.sheetnames:
                moved_counts[sh] = 0
                continue

            ws = wb[sh]
            _last_col = last_col_idx or max(pib_col_idx, ws.max_column)
            dead_zone_rows = find_named_dead_zone_rows(ws, start_row=start_row, marker_col_letter=prikom_col_letter, pib_col_letter=pib_col_letter)

            # межа по ПІБам (повна) та межа редагованої зони (до "заштатні")
            last_used_all = find_last_text_row(ws, pib_col_letter, start_row)
            if last_used_all < start_row:
                moved_counts[sh] = 0
                continue

            # відсічення по "заштатні"
            prikom_row = find_prikom_cutoff_row(ws, start_row=start_row, marker_col_letter=prikom_col_letter, marker="заштатні")
            last_used = last_used_all
            if prikom_row is not None:
                last_used = min(last_used, prikom_row - 1)
            if last_used < start_row:
                moved_counts[sh] = 0
                continue

            # зчитуємо ключі по рядках, але НЕ чіпаємо пусті розділювачі
            row_key: dict[int, str] = {}             # row -> pib_norm ("" якщо пусто/нема)
            row_snap: dict[int, list[tuple | None]] = {}    # row -> snapshots (тільки для рядків з ПІБ)
            for r in range(start_row, last_used + 1):
                if r in dead_zone_rows or _is_yellow_header(ws.cell(row=r, column=1)):
                    row_key[r] = ""
                    continue

                k = norm_pib(ws.cell(row=r, column=pib_col_idx).value)
                if not k:
                    row_key[r] = ""   # пустий рядок-розділювач
                    continue

                row_key[r] = k
                row_snap[r] = _row_snapshot(ws, r, _last_col)

            # донори з "нерухомої" зони після "заштатні":
            # їх порядок не чіпаємо, але окремі ПІБ можемо перетягнути в редагований блок
            donor_snap_by_key: dict[str, deque] = {}
            donor_rows_to_delete: set[int] = set()
            if prikom_row is not None and prikom_row <= last_used_all:
                for r in range(prikom_row, last_used_all + 1):
                    if r in dead_zone_rows or _is_yellow_header(ws.cell(row=r, column=1)):
                        continue
                    k = norm_pib(ws.cell(row=r, column=pib_col_idx).value)
                    if not k:
                        continue
                    snaps = _row_snapshot(ws, r, _last_col)
                    donor_snap_by_key.setdefault(k, deque()).append((snaps, r))

            total_rows = last_used - start_row + 1
            rows_with_pib = sum(1 for r in range(start_row, last_used + 1) if row_key.get(r, ""))
            rows_with_unit = sum(
                1 for r in range(start_row, last_used + 1)
                if row_key.get(r, "") and unit_by_pib.get(row_key.get(r, ""))
            )

            # знайдемо блоки: послідовності рядків з ПІБ, розділені пустими E
            blocks: list[list[int]] = []
            cur: list[int] = []
            for r in range(start_row, last_used + 1):
                k = row_key.get(r, "")
                if k:
                    cur.append(r)
                else:
                    if cur:
                        blocks.append(cur)
                        cur = []
            if cur:
                blocks.append(cur)

            blocks_with_unit = 0
            for block_rows in blocks:
                for r in block_rows:
                    k = row_key.get(r, "")
                    if unit_by_pib.get(k):
                        blocks_with_unit += 1
                        break

            moved_real_total = 0

            for block_rows in blocks:
                # якщо блок замалий, нема сенсу
                if len(block_rows) < 3:
                    continue

                # unit блоку визначаємо за першим ПІБом з мапи
                block_unit = ""
                for r in block_rows:
                    k = row_key.get(r, "")
                    u = unit_by_pib.get(k)
                    if u:
                        block_unit = u
                        break
                if not block_unit:
                    # не знаємо підрозділ блоку — не чіпаємо блок
                    continue

                desired_list = unit_to_ordered.get(block_unit, [])
                if not desired_list:
                    continue

                # у цьому блоці беремо лише тих, хто:
                #  - є в unit_by_pib і належить block_unit
                #  - є в order_in_unit (тобто є в desired_list)
                existing_in_block = []
                for r in block_rows:
                    k = row_key.get(r, "")
                    if not k:
                        continue
                    if unit_by_pib.get(k) == block_unit and k in order_in_unit:
                        existing_in_block.append(k)

                if not existing_in_block:
                    continue

                existing_set = set(existing_in_block)
                donor_available_set = {
                    k for k, q in donor_snap_by_key.items()
                    if q and unit_by_pib.get(k) == block_unit
                }
                available_set = existing_set | donor_available_set

                # якщо anchors знайдені в штатці, беремо тільки відповідний сегмент unit-порядку
                desired_scope = desired_list
                first_key = row_key.get(block_rows[0], "")
                last_key = row_key.get(block_rows[-1], "")
                if first_key in desired_list and last_key in desired_list:
                    i1 = desired_list.index(first_key)
                    i2 = desired_list.index(last_key)
                    if i1 <= i2:
                        desired_scope = desired_list[i1:i2 + 1]

                desired_present = [k for k in desired_scope if k in available_set]

                if len(desired_present) < 2:
                    continue

                # ---------- anchors: перший і останній ПІБ блоку НЕ рухаємо ----------
                # плюс, не рухаємо рядки, які межують з пустим рядком (над/під), але тут блок вже суцільний,
                # тож "межують" це якраз перший/останній. Ми їх вже фіксуємо.
                anchor_rows = {block_rows[0], block_rows[-1]}

                # movable rows = rows in block excluding anchors, і тільки ті що належать цьому unit
                movable_rows = []
                movable_keys = []
                for r in block_rows:
                    if r in anchor_rows:
                        continue
                    k = row_key.get(r, "")
                    if not k:
                        continue
                    if unit_by_pib.get(k) != block_unit:
                        continue
                    if k not in order_in_unit:
                        continue
                    movable_rows.append(r)
                    movable_keys.append(k)

                if not movable_rows:
                    continue

                # цільовий порядок для movable = desired_present але без anchor keys якщо вони належать unit
                anchor_keys = {row_key.get(block_rows[0], ""), row_key.get(block_rows[-1], "")}
                anchor_keys = {k for k in anchor_keys if k and unit_by_pib.get(k) == block_unit}

                desired_for_movable = [k for k in desired_present if k not in anchor_keys]

                # якщо кількість не збігається (через дублікати/змішані підрозділи), підрізаємо по мінімуму
                fill_n = min(len(movable_rows), len(desired_for_movable))
                if fill_n <= 0:
                    continue

                lcs_keys = _lcs_keys(movable_keys, desired_for_movable)
                moved_keys = set(desired_for_movable) - set(lcs_keys)

                # queue snapshots by key (щоб не губити дублікати ключів)
                snaps_by_key: dict[str, deque] = {}
                for r in movable_rows:
                    k = row_key.get(r, "")
                    if not k:
                        continue
                    snaps_by_key.setdefault(k, deque()).append((row_snap[r], None))

                # старі ключі в movable слотах
                old_keys = [row_key.get(r, "") for r in movable_rows[:fill_n]]
                new_keys = desired_for_movable[:fill_n]

                # пишемо тільки там, де реально змінилось
                for i in range(fill_n):
                    dst_r = movable_rows[i]
                    old_k = old_keys[i]
                    new_k = new_keys[i]

                    if new_k == old_k:
                        continue

                    q = snaps_by_key.get(new_k)
                    src_row = None
                    if q:
                        snaps_to_write, src_row = q.popleft()
                    else:
                        dq = donor_snap_by_key.get(new_k)
                        if not dq:
                            continue
                        snaps_to_write, src_row = dq.popleft()

                    _restore_row_snapshot(ws, dst_r, snaps_to_write, _last_col)

                    # якщо взяли ПІБ з "нерухомої" зони — заплануємо видалення рядка-джерела
                    if src_row is not None:
                        donor_rows_to_delete.add(src_row)

                    # темно-синім тільки для ключів, що були реально "зрушені"
                    if new_k != old_k and new_k in moved_keys:
                        pib_cell = ws.cell(row=dst_r, column=pib_col_idx)
                        f = copy(pib_cell.font) if pib_cell.font else Font()
                        f.color = Color(rgb=DARK_BLUE_FONT_COLOR)
                        pib_cell.font = f
                        moved_real_total += 1

            # openpyxl.delete_rows() corrupts these Excel templates. Simulate deletion
            # by shifting payload upward, skipping formula columns AK/AO.
            delete_last_used = find_last_text_row(ws, pib_col_letter, start_row)
            for r_del in sorted(donor_rows_to_delete, reverse=True):
                delete_last_used = _manual_delete_row_payload(ws, r_del, delete_last_used, _last_col)

            moved_counts[sh] = moved_real_total
            log(
                f"[REORDER/UNIT] {sh}: rows={total_rows}, pib={rows_with_pib}, "
                f"mapped={rows_with_unit}, blocks={len(blocks)}, blocks_with_unit={blocks_with_unit}"
            )
            log(f"[REORDER/UNIT] {sh}: moved_real={moved_real_total}")

        save_target_workbook(wb, target_xlsx)
        return moved_counts

    finally:
        wb.close()

def collect_existing_pibs_from_sheet(ws, start_row: int, col_letter: str) -> set[str]:
    out = set()
    for r in range(start_row, ws.max_row + 1):
        v = ws[f"{col_letter}{r}"].value
        if not v:
            continue
        k = norm_pib(v)
        if k:
            out.add(k)
    return out

def _find_or_append_person_row(ws, person: str, start_row: int, pib_col: str) -> tuple[int, bool]:
    target = norm_pib(person)
    if not target:
        return -1, False

    for r in range(start_row, ws.max_row + 1):
        v = ws[f"{pib_col}{r}"].value
        if v and norm_pib(v) == target:
            return r, False

    last_used = find_last_text_row(ws, pib_col, start_row)
    row = last_used + 1
    cell = ws[f"{pib_col}{row}"]
    cell.value = format_pib_for_sheet(person)
    cell.font = RED_FONT
    cell.alignment = LEFT_CENTER_ALIGN
    cell.border = OUTER_BORDER
    return row, True

def _day_to_col(start_col_idx: int, day: int) -> int:
    return start_col_idx + (day - 1)

def _fill_rgb(cell) -> str:
    try:
        fill = cell.fill
        if not fill or fill.patternType is None:
            return ""
        fg = getattr(fill, "fgColor", None)
        rgb = getattr(fg, "rgb", None) if fg else None
        return str(rgb).upper() if rgb else ""
    except Exception:
        return ""

def _is_green_fill(cell) -> bool:
    rgb_u = _fill_rgb(cell)
    return rgb_u.endswith("00B050") or rgb_u.endswith("006100")

def _build_pib_row_map(ws, pib_col: str, start_row: int) -> dict[str, int]:
    pib_col_idx = column_index_from_string(pib_col)
    last_row = find_last_text_row(ws, pib_col, start_row)
    out: dict[str, int] = {}

    for r in range(start_row, last_row + 1):
        pib_key = norm_pib(ws.cell(row=r, column=pib_col_idx).value)
        if pib_key and pib_key not in out:
            out[pib_key] = r

    return out

def _darken_shared_green_days_on_sheet_100(
    ws100,
    ws70,
    pib_col: str,
    start_row: int,
    start_col_idx: int,
    last_day: int,
) -> int:
    row_map_100 = _build_pib_row_map(ws100, pib_col, start_row)
    row_map_70 = _build_pib_row_map(ws70, pib_col, start_row)
    painted = 0

    for pib_key in (set(row_map_100.keys()) & set(row_map_70.keys())):
        row100 = row_map_100[pib_key]
        row70 = row_map_70[pib_key]

        for day in range(1, last_day + 1):
            col = _day_to_col(start_col_idx, day)
            cell100 = ws100.cell(row=row100, column=col)
            cell70 = ws70.cell(row=row70, column=col)

            if not _is_green_fill(cell100) or not _is_green_fill(cell70):
                continue
            if _fill_rgb(cell100).endswith("006100"):
                continue

            cell100.fill = FILL_DARK_GREEN
            painted += 1

    return painted

def build_bro_note_stats_from_json(
    bro_json_path: Path,
    year: int,
    month: int,
    last_day: int,
) -> dict[str, dict[str, int]]:
    data = json.loads(Path(bro_json_path).read_text(encoding="utf-8"))
    by_date = data.get("by_date", {}) or {}
    stats: dict[str, dict[str, int]] = {}

    def ensure_stats(pib_key: str) -> dict[str, int]:
        return stats.setdefault(pib_key, {"olive": 0, "green100": 0, "green70": 0})

    for date_key in sorted(by_date.keys()):
        try:
            d = datetime.strptime(date_key, "%Y-%m-%d").date()
        except Exception:
            continue

        if d.year != year or d.month != month:
            continue
        if not (1 <= d.day <= last_day):
            continue

        day_obj = by_date.get(date_key, {}) or {}
        rate100 = {norm_pib(x) for x in (day_obj.get("rate_100", []) or []) if norm_pib(x)}
        rate70 = {norm_pib(x) for x in (day_obj.get("rate_70", []) or []) if norm_pib(x)}
        rate30 = {norm_pib(x) for x in (day_obj.get("rate_30", []) or []) if norm_pib(x)}

        for pib_key in (rate30 - rate100 - rate70):
            ensure_stats(pib_key)["olive"] += 1

        for pib_key in (rate100 - rate70 - rate30):
            ensure_stats(pib_key)["green100"] += 1

        for pib_key in (rate70 - rate30):
            ensure_stats(pib_key)["green70"] += 1

    return stats

def _format_bro_note_text(olive: int, green100: int, green70: int, last_day: int) -> str:
    total = int(olive) + int(green100) + int(green70)
    other = max(0, int(last_day) - total)
    return f"{total:02d} = {olive:02d}/{green100:02d}/{green70:02d}/{other:02d}"

def write_bro_summary_notes_to_column(
    target_xlsx: Path,
    bro_json_path: Path,
    year: int,
    month: int,
    last_day: int,
    sheets: list[str] = None,
    pib_col_letter: str = FINAL_SOURCE_PIB_COL,
    note_col_letter: str = "AN",
    start_row: int = FINAL_SOURCE_START_ROW,
) -> dict[str, int]:
    if sheets is None:
        sheets = ["100 000", "70 000"]

    note_stats = build_bro_note_stats_from_json(
        bro_json_path=bro_json_path,
        year=year,
        month=month,
        last_day=last_day,
    )

    pib_col_idx = column_index_from_string(pib_col_letter)
    note_col_idx = column_index_from_string(note_col_letter)
    wb = load_workbook(target_xlsx)
    written_stats: dict[str, int] = {}
    restored_headers: dict[str, int] = {}

    try:
        for sh in sheets:
            if sh not in wb.sheetnames:
                written_stats[sh] = 0
                continue

            ws = wb[sh]
            written = 0

            for r in range(start_row, ws.max_row + 1):
                if _is_yellow_header(ws.cell(row=r, column=1)):
                    continue

                note_cell = ws.cell(row=r, column=note_col_idx)
                pib_key = norm_pib(ws.cell(row=r, column=pib_col_idx).value)

                if not pib_key:
                    note_cell.value = None
                    note_cell.fill = NO_FILL
                    continue

                row_stats = note_stats.get(pib_key, {})
                olive = int(row_stats.get("olive", 0))
                green100 = int(row_stats.get("green100", 0))
                green70 = int(row_stats.get("green70", 0))
                total = olive + green100 + green70

                note_cell.value = _format_bro_note_text(olive, green100, green70, last_day)
                note_cell.fill = FILL_PALE_GREEN if total == int(last_day) else NO_FILL
                note_cell.alignment = CENTER_CENTER_ALIGN
                note_cell.border = OUTER_BORDER

                f = copy(note_cell.font) if note_cell.font else Font()
                f.name = "Times New Roman"
                f.size = 12
                f.color = Color(rgb="FF000000")
                note_cell.font = f
                written += 1

            written_stats[sh] = written
            restored_headers[sh] = restore_yellow_header_rows(ws, max_col_letter="AO")

        save_target_workbook(wb, target_xlsx)
        if restored_headers:
            log(f"[BRO/NOTE->AN] restored yellow header rows: {restored_headers}")
        return written_stats
    finally:
        wb.close()

def apply_bro_to_official_sheets(
    bro_json_path: Path,
    target_xlsx: Path,
    start_col_idx: int,
    last_day: int,
    month: int,
    year: int,
    pib_col: str = FINAL_SOURCE_PIB_COL,
    start_row: int = FINAL_SOURCE_START_ROW,
) -> dict[str, int]:
    data = json.loads(Path(bro_json_path).read_text(encoding="utf-8"))
    by_date = data.get("by_date", {}) or {}

    warnings.filterwarnings(
        "ignore",
        message="Conditional Formatting extension is not supported*",
        category=UserWarning,
    )

    wb = load_workbook(target_xlsx)
    if "100 000" not in wb.sheetnames or "70 000" not in wb.sheetnames:
        wb.close()
        raise SystemExit("У файлі-цілі нема листів '100 000' та/або '70 000'.")

    ws100 = wb["100 000"]
    ws70  = wb["70 000"]
    pib_col_idx = column_index_from_string(pib_col)
    painted_green = 0
    painted_dark_green = 0
    painted_olive = 0
    skipped_outside = 0
    added_100 = 0
    added_70 = 0

    for date_key in sorted(by_date.keys()):
        try:
            d = datetime.strptime(date_key, "%Y-%m-%d").date()
        except Exception:
            continue

        if d.year != year or d.month != month:
            skipped_outside += 1
            continue
        if not (1 <= d.day <= last_day):
            continue

        col = _day_to_col(start_col_idx, d.day)

        day_obj = by_date.get(date_key, {}) or {}
        r100 = day_obj.get("rate_100", []) or []
        r70  = day_obj.get("rate_70", []) or []
        r30  = day_obj.get("rate_30", []) or []

        for p in r100:
            row, added = _find_or_append_person_row(ws100, p, start_row, pib_col)
            if row < 0:
                continue
            if added:
                added_100 += 1
            cell = ws100.cell(row=row, column=col)
            cell.value = 1
            rgb = _fill_rgb(cell)
            if str(rgb).upper().endswith("C4D79B"):
                cell.fill = FILL_YELLOW
                pib_cell = ws100.cell(row=row, column=pib_col_idx)
                pib_cell.fill = FILL_YELLOW
            else:
                cell.fill = FILL_GREEN
            painted_green += 1

        for p in r70:
            for ws in (ws70, ws100):
                row, added = _find_or_append_person_row(ws, p, start_row, pib_col)
                if row < 0:
                    continue
                if added:
                    if ws is ws100:
                        added_100 += 1
                    else:
                        added_70 += 1
                cell = ws.cell(row=row, column=col)
                cell.value = 1
                rgb = _fill_rgb(cell)
                if str(rgb).upper().endswith("C4D79B"):
                    cell.fill = FILL_YELLOW
                    pib_cell = ws.cell(row=row, column=pib_col_idx)
                    pib_cell.fill = FILL_YELLOW
                else:
                    cell.fill = FILL_GREEN
                painted_green += 1

        for p in r30:
            for ws in (ws100, ws70):
                row, added = _find_or_append_person_row(ws, p, start_row, pib_col)
                if row < 0:
                    continue
                if added:
                    if ws is ws100:
                        added_100 += 1
                    else:
                        added_70 += 1
                cell = ws.cell(row=row, column=col)

                if _is_green_fill(cell):
                    cell.fill = FILL_YELLOW
                    pib_cell = ws.cell(row=row, column=pib_col_idx)
                    pib_cell.fill = FILL_YELLOW
                    painted_olive += 1
                    continue

                cell.fill = FILL_OLIVE
                painted_olive += 1

    painted_dark_green = _darken_shared_green_days_on_sheet_100(
        ws100=ws100,
        ws70=ws70,
        pib_col=pib_col,
        start_row=start_row,
        start_col_idx=start_col_idx,
        last_day=last_day,
    )

    save_target_workbook(wb, target_xlsx)
    wb.close()

    log(f"[BRO/ADD] added_100={added_100}, added_70={added_70}, total={added_100 + added_70}")
    log(
        f"[BRO->XLSX] green={painted_green}, dark_green_100={painted_dark_green}, "
        f"olive={painted_olive}, skipped_outside={skipped_outside}"
    )
    return {
        "green": painted_green,
        "dark_green_100": painted_dark_green,
        "olive": painted_olive,
        "skipped_outside_month": skipped_outside,
    }


# ===================== ШТАТКА: snapshot для timeline =====================

STAFF_READ_COLS = ["V", "T", "W", "X"]
STAFF_PIB_COL = "L"
STAFF_FILTER_COL = "O"
STAFF_START_ROW = 3

PRIKOM_MARKER = "заштатні"

def row_has_prikom_marker(ws, row: int) -> bool:
    for col_idx in list(range(1, 27)) + list(range(27, 53)):
        v = ws.cell(row=row, column=col_idx).value
        if not v:
            continue
        if PRIKOM_MARKER in str(v).strip().lower():
            return True
    return False

def collect_prikom_pibs_from_official(target_xlsx: Path) -> set[str]:
    wb = load_workbook(target_xlsx, data_only=True)

    out: set[str] = set()
    marker = "заштатні"

    for sh in FINAL_SOURCE_SHEETS:
        if sh not in wb.sheetnames:
            continue
        ws = wb[sh]

        start_row = None
        for r in range(1, ws.max_row + 1):
            v = ws[f"D{r}"].value
            if v and marker in str(v).strip().lower():
                start_row = r + 1
                break

        ...

        for r in range(start_row, ws.max_row + 1):
            pib_val = ws[f"E{r}"].value
            k = norm_pib(pib_val)
            if k:
                out.add(k)

    wb.close()
    log(f"[PRIKOM] з 100 000/70 000 зібрано ПІБ: {len(out)}")
    return out

def extract_staff_day_snapshot(staff_path: Path, forced_pibs: set[str]) -> Tuple[str, Dict[str, Dict[str, str]], List[str]]:
    date_key = extract_date_from_staff_filename(staff_path)

    keep_vba = staff_path.suffix.lower() == ".xlsm"
    wb = load_workbook(staff_path, data_only=False, keep_vba=keep_vba, keep_links=True)
    ws = get_staff_sheet(wb)

    day_map: Dict[str, Dict[str, str]] = {}
    ordered: List[str] = []
    prikom_mode = False

    for r in range(STAFF_START_ROW, ws.max_row + 1):

        if not prikom_mode and row_has_prikom_marker(ws, r):
            prikom_mode = True
            continue

        pib_raw = ws[f"{STAFF_PIB_COL}{r}"].value
        pib_norm = norm_pib(pib_raw)
        if not pib_norm:
            continue

        o_val = normalize_value(ws[f"{STAFF_FILTER_COL}{r}"].value).upper()
        is_3mb = (o_val == "3 МБ")
        is_forced = (pib_norm in forced_pibs)

        if not (is_3mb or prikom_mode or is_forced):
            continue

        rec = {"pib": str(pib_raw).strip()}
        for col in STAFF_READ_COLS:
            rec[col] = normalize_value(ws[f"{col}{r}"].value)

        day_map[pib_norm] = rec
        ordered.append(rec["pib"])

    wb.close()

    seen = set()
    ordered_unique = []
    for p in ordered:
        k = norm_pib(p)
        if not k or k in seen:
            continue
        seen.add(k)
        ordered_unique.append(p)

    return date_key, day_map, ordered_unique


# ===================== ТАЙМЛАЙН (STAFF) =====================

def build_staff_timeline_json(staff_folder: Path, out_json: Path, forced_pibs: set[str]) -> Tuple[Path, List[str], Dict[str, List[str]]]:
    files = sorted([p for p in staff_folder.iterdir()
                    if p.is_file() and p.suffix.lower() in [".xlsx", ".xlsm"] and not p.name.startswith("~$")])
    if not files:
        raise SystemExit("У папці зі штатками не знайшов .xlsx/.xlsm")

    days: List[Tuple[str, Dict[str, Dict[str, str]], List[str]]] = []
    ordered_by_date: Dict[str, List[str]] = {}

    for f in files:
        res = safe_call("STAFF/SNAPSHOT", f, lambda: extract_staff_day_snapshot(f, forced_pibs), default=None)
        if not res:
            continue
        date_key, day_map, ordered = res
        log(f"[STAFF/OK] {f.name}: PIB={len(day_map)}")
        days.append((date_key, day_map, ordered))
        ordered_by_date[date_key] = ordered

    if not days:
        raise SystemExit("Нема жодної валідної штатки з датою у назві.")

    days.sort(key=lambda x: x[0])

    month_to_dates: Dict[str, List[str]] = {}
    for dk, _m, _o in days:
        mk = dk[:7]
        month_to_dates.setdefault(mk, []).append(dk)

    month_first_last = set()
    for mk, lst in month_to_dates.items():
        lst_sorted = sorted(lst)
        month_first_last.add(lst_sorted[0])
        month_first_last.add(lst_sorted[-1])

    timeline: Dict[str, Dict[str, Any]] = {}

    prev_present: set[str] = set()
    last_seen: Dict[str, Dict[str, str]] = {}

    def ensure_person(pib_norm: str, pib_original: str):
        obj = timeline.setdefault(pib_norm, {"pib": pib_original, "points": {}})
        if not obj.get("pib") and pib_original:
            obj["pib"] = pib_original
        return obj

    def payload_from_rec(rec: Dict[str, str]) -> Dict[str, str]:
        return {
            "T": (rec.get("T", "") or "").strip(),
            "W": (rec.get("W", "") or "").strip(),
            "X": (rec.get("X", "") or "").strip(),
        }

    def pick_key_for_regular(rec: Dict[str, str], fallback_date: str) -> str:
        v = (rec.get("V", "") or "").strip()
        return v if v else fallback_date

    def put_point(pib_norm: str, pib_original: str, key: str, rec: Dict[str, str], kind: str, x_override: str | None = None):
        if not key:
            return
        obj = ensure_person(pib_norm, pib_original)

        payload = payload_from_rec(rec)
        if x_override is not None:
            payload["X"] = x_override

        existing = obj["points"].get(key)
        if existing and existing.get("_kind") == "file" and kind == "v":
            # зберігаємо "прибув/вибув", бо це системний маркер і він важливіший за порожній X зі штатки
            ex_x = (existing.get("X", "") or "").strip().casefold()
            merged = {"_kind": "file", **payload}
            if ex_x in ("прибув", "вибув"):
                merged["X"] = existing.get("X", "")
            obj["points"][key] = merged
            return

        obj["points"][key] = {"_kind": kind, **payload}

    for date_key, day_map, _ordered in days:
        cur_present = set(day_map.keys())

        # ====== (2) Появився у штаті -> X = "прибув" ======
        appeared = cur_present - prev_present
        for pib_norm in appeared:
            rec = day_map[pib_norm]
            put_point(pib_norm, rec.get("pib", ""), key=date_key, rec=rec, kind="file", x_override="прибув")

        # ====== (2) Зник зі штатки -> X = "вибув" ======
        disappeared = prev_present - cur_present
        for pib_norm in disappeared:
            last_rec = last_seen.get(pib_norm, {}) or {}
            if not last_rec:
                continue
            put_point(pib_norm, last_rec.get("pib", ""), key=date_key, rec=last_rec, kind="file", x_override="вибув")

        # regular points by V (як було)
        for pib_norm, rec in day_map.items():
            key_v = pick_key_for_regular(rec, fallback_date=date_key)
            put_point(pib_norm, rec.get("pib", ""), key=key_v, rec=rec, kind="v")
            last_seen[pib_norm] = rec

            if date_key in month_first_last:
                put_point(pib_norm, rec.get("pib", ""), key=date_key, rec=rec, kind="file")

        prev_present = cur_present

    iso_re = re.compile(r"^\d{4}-\d{2}-\d{2}$")

    def sort_key(k: str):
        if iso_re.match(k):
            return (0, k)
        return (1, k)

    out_timeline = {}
    for pib_norm, obj in timeline.items():
        pts = obj.get("points", {}) or {}
        keys_sorted = sorted(pts.keys(), key=sort_key)
        events = [{"V": k, **pts[k]} for k in keys_sorted]
        out_timeline[pib_norm] = {"pib": obj.get("pib", ""), "events": events}

    out_json.parent.mkdir(parents=True, exist_ok=True)
    out_json.write_text(json.dumps({
        "source_folder": str(staff_folder),
        "files_processed": len(days),
        "dates": [d[0] for d in days],
        "month_first_last_files": sorted(month_first_last),
        "timeline": out_timeline,
    }, ensure_ascii=False, indent=2), encoding="utf-8")

    all_pibs_original = sorted({v.get("pib", "").strip() for v in out_timeline.values() if v.get("pib", "").strip()})
    log(f"[TIMELINE] JSON: {out_json} | людей: {len(all_pibs_original)}")
    return out_json, all_pibs_original, ordered_by_date


# ===================== ДОДАТИ ВІДСУТНІХ (червоний) =====================

def try_insert_near_neighbors(
    ws,
    person: str,
    staff_order: List[str],
    start_row: int,
    pib_col: str,
    excluded_rows: set[int] | None = None,
    max_col_letter: str = ROW_MOVE_LAST_COL_LETTER,
) -> int:
    target = norm_pib(person)
    if not target:
        return -1

    if excluded_rows is None:
        excluded_rows = set()

    existing_row = {}
    for r in range(start_row, ws.max_row + 1):
        if r in excluded_rows:
            continue
        v = ws[f"{pib_col}{r}"].value
        if not v:
            continue
        existing_row[norm_pib(v)] = r

    idx = None
    staff_norm = [norm_pib(x) for x in staff_order]
    for i, k in enumerate(staff_norm):
        if k == target:
            idx = i
            break
    if idx is None:
        return -1

    max_col = column_index_from_string(max_col_letter)

    for j in range(idx - 1, -1, -1):
        nk = staff_norm[j]
        if nk in existing_row:
            insert_after = existing_row[nk]
            candidate_row = insert_after + 1
            if candidate_row in excluded_rows:
                continue
            last_used = find_last_text_row(ws, pib_col, start_row)
            _manual_insert_row_payload(ws, candidate_row, last_used, max_col)
            return candidate_row

    for j in range(idx + 1, len(staff_norm)):
        nk = staff_norm[j]
        if nk in existing_row:
            insert_before = existing_row[nk]
            if insert_before in excluded_rows:
                continue
            last_used = find_last_text_row(ws, pib_col, start_row)
            _manual_insert_row_payload(ws, insert_before, last_used, max_col)
            return insert_before

    return -1

def add_missing_people_to_official_sheets(
    target_xlsx: Path,
    all_pibs: List[str],
    staff_order_hint: List[str],
    start_row: int = FINAL_SOURCE_START_ROW,
    pib_col: str = FINAL_SOURCE_PIB_COL,
) -> Dict[str, int]:
    wb = load_workbook(target_xlsx)
    added = {s: 0 for s in FINAL_SOURCE_SHEETS}
    skipped_due_to_dead_zone = {s: 0 for s in FINAL_SOURCE_SHEETS}
    dead_zone_present = {s: False for s in FINAL_SOURCE_SHEETS}

    existing_all = set()
    for sh in FINAL_SOURCE_SHEETS:
        if sh not in wb.sheetnames:
            continue
        existing_all |= collect_existing_pibs_from_sheet(wb[sh], start_row, pib_col)

    to_add = []
    for p in all_pibs:
        if norm_pib(p) not in existing_all:
            to_add.append(p)

    if not to_add:
        wb.close()
        log("[ADD] Усі ПІБ зі штаток вже є в 100 000/70 000.")
        return {**added, "total": 0}

    for sh in FINAL_SOURCE_SHEETS:
        if sh not in wb.sheetnames:
            continue
        dead_zone_present[sh] = bool(
            find_named_dead_zone_rows(
                wb[sh],
                start_row=start_row,
                pib_col_letter=pib_col,
            )
        )

    for person in to_add:
        for sh in FINAL_SOURCE_SHEETS:
            if sh not in wb.sheetnames:
                continue
            ws = wb[sh]
            if dead_zone_present.get(sh):
                skipped_due_to_dead_zone[sh] += 1
                continue

            row = try_insert_near_neighbors(
                ws,
                person,
                staff_order_hint,
                start_row,
                pib_col,
                excluded_rows=set(),
            )
            if row < 0:
                last_used = find_last_text_row(ws, pib_col, start_row)
                row = last_used + 1

            cell = ws[f"{pib_col}{row}"]
            cell.value = format_pib_for_sheet(person)
            cell.font = RED_FONT
            cell.alignment = LEFT_CENTER_ALIGN
            cell.border = OUTER_BORDER
            added[sh] += 1

    save_target_workbook(wb, target_xlsx)
    wb.close()

    total = sum(added.values())
    log(
        f"[ADD] Додано (червоним): {added}, dead_zone_present={dead_zone_present}, "
        f"skipped_dead_zone={skipped_due_to_dead_zone}, total={total}"
    )
    return {
        **added,
        "dead_zone_present": dead_zone_present,
        "skipped_dead_zone": skipped_due_to_dead_zone,
        "total": total,
    }


# ===================== JSON timeline -> колонка A =====================

def timeline_to_text(events: list[dict]) -> str:
    if not events:
        return ""

    iso_re = re.compile(r"^\d{4}-\d{2}-\d{2}$")

    def sort_key_ev(ev):
        v = (ev.get("V") or "").strip()
        if iso_re.match(v):
            return (0, v)
        return (1, v)

    events_sorted = sorted(events, key=sort_key_ev)

    lines = []
    seen_exact_lines = set()
    prev_payload_for_v = None

    for ev in events_sorted:
        V = (ev.get("V") or "").strip()
        T = (ev.get("T") or "").strip()
        W = (ev.get("W") or "").strip()
        X = (ev.get("X") or "").strip()
        kind = (ev.get("_kind") or "v").strip().lower()

        payload = (T, W, X)

        if kind == "v":
            if prev_payload_for_v is not None and payload == prev_payload_for_v:
                continue
            prev_payload_for_v = payload

        s = f"{V}, {T}, {W}, {X};".strip()
        if not s:
            continue

        if s in seen_exact_lines:
            continue
        seen_exact_lines.add(s)

        lines.append(s)

    return "\n".join(lines).strip()

def build_prepared_from_timeline_json(timeline_json_path: Path) -> dict[str, str]:
    data = json.loads(Path(timeline_json_path).read_text(encoding="utf-8"))
    tl = data.get("timeline", {}) or {}

    prepared = {}
    for pib_norm_key, obj in tl.items():
        key = norm_pib(pib_norm_key)
        events = obj.get("events", []) or []
        text = timeline_to_text(events)
        if key and text:
            prepared[key] = text
    return prepared

def write_column_a_via_openpyxl(target_xlsx: Path, prepared: dict[str, str]) -> tuple[int, int]:
    wb = load_workbook(target_xlsx)
    written = 0
    skipped = 0

    try:
        for sheet_name in FINAL_SOURCE_SHEETS:
            if sheet_name not in wb.sheetnames:
                continue

            ws = wb[sheet_name]
            last_row = find_last_text_row(ws, "E", FINAL_SOURCE_START_ROW)
            if last_row < FINAL_SOURCE_START_ROW:
                continue

            for r in range(FINAL_SOURCE_START_ROW, last_row + 1):
                pib_val = ws.cell(row=r, column=5).value
                key = norm_pib(pib_val)
                if not key:
                    continue

                txt = prepared.get(key)
                if not txt:
                    skipped += 1
                    continue

                ws.cell(row=r, column=1).value = txt
                written += 1

        save_target_workbook(wb, target_xlsx)
    finally:
        wb.close()

    return written, skipped

def push_timeline_to_column_a(target_xlsx: Path, timeline_json_path: Path):
    stage_start("JSON timeline -> write Column A (100 000/70 000)")
    prepared = build_prepared_from_timeline_json(timeline_json_path)
    log(f"[A-PUSH] prepared entries: {len(prepared)}")

    if not prepared:
        stage_end()
        raise SystemExit("У staff_timeline.json не знайшов даних для запису у колонку A.")

    written, skipped = write_column_a_via_openpyxl(target_xlsx, prepared)
    log(f"[A-PUSH] written={written}, skipped(no json)={skipped}")
    stage_end()


# ===================== РОЗКРАСКА (як у тебе) =====================

def final_xl_rgb(r: int, g: int, b: int) -> int:
    return (b << 16) + (g << 8) + r

_final_space_re = re.compile(r"\s+")
_iso_date_re = re.compile(r"^\d{4}-\d{2}-\d{2}$")

def _try_parse_iso(s: str) -> date | None:
    s = (s or "").strip()
    if not _iso_date_re.match(s):
        return None
    try:
        return datetime.strptime(s, "%Y-%m-%d").date()
    except Exception:
        return None

def _parse_a_text_to_events(a_text: str) -> list[dict]:
    if not a_text:
        return []

    text = str(a_text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00A0", " ")
    text = _final_space_re.sub(" ", text).strip()

    if not text:
        return []

    chunks = [c.strip() for c in text.split(";")]
    out = []
    for c in chunks:
        if not c:
            continue
        c = c.strip().strip(",").strip()
        if not c:
            continue
        parts = [p.strip() for p in c.split(",")]
        while len(parts) < 4:
            parts.append("")
        V, T, W, X = parts[0], parts[1], parts[2], parts[3]
        out.append({"V": V, "T": T, "W": W, "X": X})
    return out

# ===== priority for reasons (higher = stronger) =====
_REASON_PRIORITY = {
    final_xl_rgb(255, 0, 0):     5,  # RED
    final_xl_rgb(255, 165, 0):   4,  # ORANGE
    final_xl_rgb(255, 153, 102): 3,  # ARREST
    final_xl_rgb(0, 112, 192):   2,  # BLUE
    final_xl_rgb(96, 96, 96):    2,  # DARK_GREY
    final_xl_rgb(160, 160, 160): 1,  # GREY
}

def _event_reason_color(ev: dict) -> int | None:
    V = (ev.get("V") or "").strip()
    T = (ev.get("T") or "").strip()
    W = (ev.get("W") or "").strip()
    X = (ev.get("X") or "").strip()

    DARK_GREY = final_xl_rgb(96, 96, 96)
    GREY      = final_xl_rgb(160, 160, 160)

    x_clean = X.casefold()
    blob = " ".join([V, T, W, X]).lower()

    RED    = final_xl_rgb(255, 0, 0)
    ORANGE = final_xl_rgb(255, 165, 0)
    BLUE   = final_xl_rgb(0, 112, 192)
    ARREST = final_xl_rgb(255, 153, 102)

    # 1) сильні стани
    if "арешт" in blob:
        return ARREST
    if "сзч" in blob:
        return RED
    if any(k in blob for k in ["смерт", "загиб", "безвісти", "зник"]):
        return ORANGE
    if any(k in blob for k in ["шпит", "реаб"]):
        return BLUE
    if any(k in blob for k in ["відпуст", "звільн", "декрет", "пологов"]):
        return GREY

    # 2) рух складу (fallback, але після сильних)
    if x_clean in ("прибув", "вибув"):
        return DARK_GREY

    # 3) порожня причина під дату -> GREY
    if V and not T and not W and not X:
        return GREY

    # 4) відрядження: як і було
    if "відряджен" in blob:
        if W.strip().casefold() != "уос":
            return GREY

    return None

def _event_has_vybuv_any(ev: dict) -> bool:
    blob = " ".join([
        str(ev.get("V") or ""),
        str(ev.get("T") or ""),
        str(ev.get("W") or ""),
        str(ev.get("X") or ""),
    ]).casefold()
    return "вибув" in blob

def _final_color_rgb(color_int: int) -> str:
    color_int = int(color_int)
    r = color_int & 0xFF
    g = (color_int >> 8) & 0xFF
    b = (color_int >> 16) & 0xFF
    return f"{r:02X}{g:02X}{b:02X}"

def _final_color_argb(color_int: int) -> str:
    return f"FF{_final_color_rgb(color_int)}"

def is_effectively_blank_fill(cell) -> bool:
    rgb = _fill_rgb(cell)
    return not rgb or rgb.endswith("FFFFFF")

def set_fill(cell, color_int: int):
    cell.fill = PatternFill(fill_type="solid", fgColor=_final_color_argb(color_int))

def apply_fill_with_collision(cell, new_color: int, yellow_color: int) -> bool:
    if is_effectively_blank_fill(cell):
        set_fill(cell, new_color)
        return False

    existing_rgb = _fill_rgb(cell)
    new_rgb = _final_color_rgb(new_color)
    if existing_rgb.endswith(new_rgb):
        return False

    strong_existing = {
        _final_color_rgb(final_xl_rgb(255, 0, 0)),     # RED
        _final_color_rgb(final_xl_rgb(255, 165, 0)),   # ORANGE
        _final_color_rgb(final_xl_rgb(255, 153, 102)), # ARREST
        _final_color_rgb(final_xl_rgb(0, 112, 192)),   # BLUE
    }
    if int(new_color) == int(final_xl_rgb(96, 96, 96)) and any(
        existing_rgb.endswith(rgb) for rgb in strong_existing
    ):
        return False

    set_fill(cell, yellow_color)
    return True

def _pick_top_color(colors: set[int]) -> int:
    GREY = final_xl_rgb(160, 160, 160)

    if not colors:
        return int(GREY)

    best = None
    best_pr = -10**9

    for c in {int(x) for x in colors}:
        pr = _REASON_PRIORITY.get(int(c), 0)  # невідомі = 0
        if pr > best_pr:
            best_pr = pr
            best = int(c)

    return int(best if best is not None else GREY)

def final_apply_coloring_in_source_via_openpyxl(
    source_xlsx: Path,
    current_year: int,
    current_month: int,
    start_col_idx: int,
    last_day: int,
    final_source_sheets: list[str] = ("100 000", "70 000"),
    pib_col_idx: int = 5,   # E
    a_col_idx: int = 1,     # A
    grid_from_col_idx: int | None = None,  # якщо None -> start_col_idx (F)
):
    YELLOW = final_xl_rgb(255, 255, 0)

    if grid_from_col_idx is None:
        grid_from_col_idx = start_col_idx  # F

    wb = load_workbook(source_xlsx)
    try:
        month_first = date(current_year, current_month, 1)
        month_last  = date(current_year, current_month, last_day)
        pib_col_letter = get_column_letter(pib_col_idx)

        for sheet_name in final_source_sheets:
            if sheet_name not in wb.sheetnames:
                continue

            ws = wb[sheet_name]
            last_row = find_last_text_row(ws, pib_col_letter, 1)

            for r in range(1, last_row + 1):
                val_a = ws.cell(row=r, column=a_col_idx).value
                if not val_a:
                    continue

                events = _parse_a_text_to_events(val_a)
                if not events:
                    continue

                # беремо тільки події, де V ISO-дата
                events2 = []
                last_iso_idx = -1
                for idx, ev in enumerate(events):
                    dV = _try_parse_iso(ev.get("V", ""))
                    if not dV:
                        continue
                    last_iso_idx = idx
                    events2.append((dV, ev, idx))

                if not events2:
                    continue

                # сорт по V
                events2.sort(key=lambda x: x[0])

                # precompute "last vybuv" force + effective colors
                events2_eff = []
                for dV, ev, src_idx in events2:
                    force_last_vybuv = (src_idx == last_iso_idx and _event_has_vybuv_any(ev))
                    if force_last_vybuv:
                        color = int(final_xl_rgb(96, 96, 96))  # DARK_GREY (force "вибув")
                    else:
                        color = _event_reason_color(ev)
                    events2_eff.append((dV, ev, src_idx, force_last_vybuv, color))

                # готуємо періоди (start,end,color)
                periods: list[tuple[date, date, int]] = []
                last_idx = len(events2_eff) - 1
                for i, (dV, ev, src_idx, force_last_vybuv, color) in enumerate(events2_eff):
                    if color is None:
                        continue

                    # 1) planned_end
                    planned_end = _try_parse_iso(ev.get("X", ""))
                    if planned_end is None:
                        planned_end = month_last

                    # 2) SPECIAL: "прибув" та "вибув"
                    x_clean = (ev.get("X") or "").strip().casefold()
                    if force_last_vybuv:
                        x_clean = "вибув"

                    if x_clean == "прибув":
                        # pre-period: ДО дня прибуття
                        s0 = month_first
                        e0 = dV - timedelta(days=1)
                        if e0 >= s0:
                            periods.append((s0, e0, int(final_xl_rgb(96, 96, 96))))  # DARK_GREY
                        # "прибув" = межа стану, після дати прибуття не фарбуємо
                        continue
                    if x_clean == "вибув":
                        # post-period: ВІД дня вибуття і до кінця періоду
                        s0 = max(dV, month_first)
                        e0 = month_last
                        if e0 >= s0:
                            periods.append((s0, e0, int(color)))
                        # "вибув" = межа стану, до дати вибуття не фарбуємо
                        continue

                    # 3) FIX: "відпустка/реаб" з кривим X (X < V)
                    if planned_end < dV:
                        blob_low = " ".join([
                            (ev.get("T") or ""),
                            (ev.get("W") or ""),
                            (ev.get("X") or ""),
                        ]).lower()
                        if "відпуст" in blob_low or "реаб" in blob_low:
                            planned_end = month_last

                    # 4) обрізання по nextV, але НЕ обрізаємо якщо наступна подія має той самий колір (продовження стану)
                    j = i + 1
                    while j < len(events2_eff):
                        nextV, nextEv = events2_eff[j][0], events2_eff[j][1]
                        nextColor = events2_eff[j][4]
                        # якщо наступна подія того ж "типу" (той же колір) — ігноруємо її як межу періоду
                        if nextColor is not None and nextColor == color:
                            j += 1
                            continue

                        # перша "інша" подія — вона і є межею
                        if nextV <= planned_end:
                            planned_end = nextV - timedelta(days=1)
                        break

                    # 5) перетин з місяцем
                    if planned_end < month_first or dV > month_last:
                        continue

                    s = max(dV, month_first)
                    e = min(planned_end, month_last)
                    periods.append((s, e, int(color)))

                if not periods:
                    continue

                # ===== A..D = top_color, НІЯКОГО YELLOW ТУТ =====
                unique_colors = {int(c) for _, _, c in periods}
                top_color = _pick_top_color(unique_colors)

                for c in range(1, 5):
                    set_fill(ws.cell(row=r, column=c), int(top_color))

                # E (ПІБ) теж top_color, але може стати жовтим якщо були колізії в днях
                pib_cell = ws.cell(row=r, column=pib_col_idx)  # E
                set_fill(pib_cell, int(top_color))

                had_day_collision = False

                # ===== дні (grid) =====
                for s, e, color in periods:
                    cur = s
                    while cur <= e:
                        col = _day_to_col(grid_from_col_idx, cur.day)
                        cell = ws.cell(row=r, column=col)

                        if apply_fill_with_collision(cell, int(color), int(YELLOW)):
                            had_day_collision = True

                        cur += timedelta(days=1)

                # ===== якщо була колізія у днях -> жовтим тільки E =====
                if had_day_collision:
                    set_fill(pib_cell, int(YELLOW))

        save_target_workbook(wb, source_xlsx)
    finally:
        wb.close()

# ===================== MAIN =====================

def main():
    run_start()

    bro_folder, staff_folder, target_xlsx, start_col_idx, last_day, month, year = ask_all_inputs_window()
    has_bro = bro_folder is not None
    has_staff = staff_folder is not None

    _init_log_window()
    backup_excel_to_trash(target_xlsx)
    log(f"[TIME] year={year}, month={month}, last_day={last_day}, start_col_idx={start_col_idx}")
    log(f"[INPUT] BRO folder   : {bro_folder}")
    log(f"[INPUT] STAFF folder : {staff_folder}")
    log(f"[INPUT] TARGET file  : {target_xlsx}")
    log(f"[MODE] has_bro={has_bro}, has_staff={has_staff}")

    base_dir = target_xlsx.parent
    out_staff_json = base_dir / "staff_timeline.json"
    out_bro_json = base_dir / BRO_JSON_NAME
    trash_file(out_staff_json, "STALE STAFF JSON")
    trash_file(out_bro_json, "STALE BRO JSON")

    stage_start("CLEAR: wipe A,B + days grid; unfill C,D (keep headers)")
    clear_stats = clear_official_sheets_zero_stage(
        target_xlsx=target_xlsx,
        start_col_idx=start_col_idx,
        last_day=last_day,
    )
    stage_end()

    ordered_by_date_to_path_map: dict[str, Path] = {}
    if has_staff:
        stage_start("STAFF MAP: build ordered_by_date_to_path_map")
        ordered_by_date_to_path_map = build_ordered_by_date_to_path_map(staff_folder)
        stage_end()
    else:
        log("[STAFF] Папку штаток не вибрано, staff-етапи буде скіпнуто.")

    success = False
    try:
        if has_bro:
            stage_start("BRO: parse docx -> bro_by_date.json")
            bro_json_path = build_bro_by_date_json(bro_folder, out_bro_json)
            stage_end()

            stage_start("BRO: apply rates to 100 000 / 70 000 (days grid)")
            bro_stats = apply_bro_to_official_sheets(
                bro_json_path=bro_json_path,
                target_xlsx=target_xlsx,
                start_col_idx=start_col_idx,
                last_day=last_day,
                month=month,
                year=year,
                pib_col=FINAL_SOURCE_PIB_COL,
                start_row=FINAL_SOURCE_START_ROW,
            )
            stage_end()
        else:
            log("[BRO] Папку БРО не вибрано, BRO-етапи скіпнуто.")

        if has_staff:
            stage_start("Read prikom PIBs from 100 000/70 000 (col C marker)")
            forced_pibs = collect_prikom_pibs_from_official(target_xlsx)
            stage_end()

            stage_start("Build staff_timeline.json from staff files")
            json_path, all_pibs, ordered_by_date = build_staff_timeline_json(staff_folder, out_staff_json, forced_pibs)
            stage_end()

            stage_start("Add missing PIBs to 100 000 / 70 000 (red font)")
            last_date = sorted(ordered_by_date.keys())[-1]
            staff_order_hint = ordered_by_date.get(last_date, [])
            stats = add_missing_people_to_official_sheets(
                target_xlsx=target_xlsx,
                all_pibs=all_pibs,
                staff_order_hint=staff_order_hint,
                start_row=FINAL_SOURCE_START_ROW,
                pib_col=FINAL_SOURCE_PIB_COL,
            )
            stage_end()

            stage_start("UNIT->B: preload last staff unit initials by PIB (dark-blue)")
            last_staff_xlsx = pick_last_staff_xlsx(ordered_by_date_to_path_map)
            if last_staff_xlsx and last_staff_xlsx.exists():
                unit_b_stats = preload_staff_unit_initials_to_column_b(
                    target_xlsx=target_xlsx,
                    staff_xlsx_last=last_staff_xlsx,
                    sheets=["100 000", "70 000"],
                    start_row=FINAL_SOURCE_START_ROW,
                    pib_col_letter="E",
                    dst_col_letter="B",
                    backup_col_letter="AL",
                    staff_name_col_letter="L",
                    staff_unit_col_letter="B",
                    staff_start_row=3,
                )
                log(f"[UNIT->B] stats={unit_b_stats}")
            else:
                log("[UNIT->B/WARN] Не знайшов останню штатку, етап скіпнуто.")
            stage_end()

        if has_bro:
            stage_start("SRC->B: write BRO docx sources to column B (100 000/70 000)")
            b_added = write_docx_sources_to_column_b(
                br_folder=bro_folder,
                target_xlsx=target_xlsx,
                sheets=["100 000", "70 000"],
                pib_col="E",
                dst_col="B",
                start_row=FINAL_SOURCE_START_ROW
            )
            log(f"[SRC->B] added={b_added}")
            stage_end()
        else:
            log("[SRC->B] Папку БРО не вибрано, заповнення колонки B з DOCX скіпнуто.")

        if has_staff:
            stage_start("SYNC: last staff -> ranks/pos to columns (100 000/70 000)")
            sync_rank_pos_after_push(staff_file_map=ordered_by_date_to_path_map, target_xlsx=target_xlsx)
            stage_end()

            push_timeline_to_column_a(target_xlsx=target_xlsx, timeline_json_path=out_staff_json)
            trash_file(out_staff_json, "STAFF JSON")

            stage_start("REORDER by UNIT blocks (last staff)")
            last_staff_xlsx = pick_last_staff_xlsx(ordered_by_date_to_path_map)
            reorder_stats = reorder_by_unit_blocks_from_last_staff(
                target_xlsx=target_xlsx,
                staff_xlsx_last=last_staff_xlsx,
                sheets=["100 000", "70 000"],
                start_row=8,
                pib_col_letter="E",
                prikom_col_letter="D",
                staff_name_col_letter="L",
                staff_unit_col_letter="B",
                staff_start_row=3,
                start_col_idx=start_col_idx,
                last_day=last_day,
                extra_cols=5,
                force_last_col_letter=ROW_MOVE_LAST_COL_LETTER,
            )
            log(f"[REORDER/UNIT] stats={reorder_stats}")
            stage_end()

            stage_start("FINAL: coloring A..D + days grid (from Column A periods)")
            final_apply_coloring_in_source_via_openpyxl(
                source_xlsx=target_xlsx,
                current_year=year,
                current_month=month,
                start_col_idx=start_col_idx,
                last_day=last_day
            )
            stage_end()
        else:
            log("[STAFF] Синхронізацію звань/посад, timeline та reorder скіпнуто: штатки не вибрано.")

        stage_start("FINAL: text formatting (A..C wrap, B center, days center)")
        apply_body_formatting_to_official_sheets(
            target_xlsx=target_xlsx,
            start_col_idx=start_col_idx,
            last_day=last_day,
            start_row=CLEAR_START_ROW,
            extra_cols=5,
        )
        stage_end()

        if has_bro:
            stage_start("FINAL: BRO summary notes -> column AN (100 000/70 000)")
            bro_note_stats = write_bro_summary_notes_to_column(
                target_xlsx=target_xlsx,
                bro_json_path=out_bro_json,
                year=year,
                month=month,
                last_day=last_day,
                sheets=["100 000", "70 000"],
                pib_col_letter="E",
                note_col_letter="AN",
                start_row=FINAL_SOURCE_START_ROW,
            )
            log(f"[BRO/NOTE->AN] stats={bro_note_stats}")
            stage_end()
            trash_file(out_bro_json, "BRO JSON")
        else:
            log("[BRO/NOTE->AN] Папку БРО не вибрано, колонку AN скіпнуто.")

        success = True

    except UserFacingError as exc:
        log(f"[ERROR] {getattr(exc, 'title', 'Помилка')}: {exc}")
        raise
    except OSError as exc:
        if _is_file_access_error(exc):
            log(f"[ERROR] Файл зайнятий: {exc}")
        raise
    finally:
        run_end()
        _shutdown_save_check_excel()
        if not success:
            _send_session_log_to_trash("RewardsBefore")
            _close_log_window()

    root = Tk()
    root.withdraw()
    install_frozen_executable_icon(root)
    install_dark_title_bar(root)
    root.attributes("-topmost", True)

    msg = (
        f"RewardsBefore: листи '100 000'/'70 000' заповнені."
    )

    messagebox.showinfo("Готово!", msg, parent=root)
    _send_session_log_to_trash("RewardsBefore")
    _close_log_window()
    root.destroy()

if __name__ == "__main__":
    mp.freeze_support()
    try:
        main()
    except UserFacingError as exc:
        try:
            log(f"[ERROR] {getattr(exc, 'title', 'Помилка')}: {exc}")
        except Exception:
            pass
        cleanup_failed_run()
        show_user_error_window(exc)
        sys.exit(1)
    except OSError as exc:
        if _is_file_access_error(exc):
            user_exc = FileAccessError(
                _file_access_message(_permission_error_path(exc), "доступитися до файлу", exc)
            )
            try:
                log(f"[ERROR] {user_exc.title}: {user_exc}")
            except Exception:
                pass
            cleanup_failed_run()
            show_user_error_window(user_exc)
            sys.exit(1)
        cleanup_failed_run()
        raise
    except BaseException:
        cleanup_failed_run()
        raise
