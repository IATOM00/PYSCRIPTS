from __future__ import annotations

from typing import Any, Dict, Iterable, List, Optional, Tuple
import sys, re, time, shutil, queue, logging, threading
from dataclasses import dataclass, field
from docx.table import _Cell, Table
from datetime import date, datetime
from send2trash import send2trash
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from pathlib import Path
import openpyxl

# ============================================================
# НАЛАШТУВАННЯ
# ============================================================

SHEET_NAME = "таблиця"

# Фіксовані колонки у штатці
XLSX_COLS = {
    "unit": "O",           # Підрозділ
    "department": "B",     # Відділення
    "position": "F",       # Посада
    "full_position": "G",  # Повна посада
    "rank": "K",           # Звання
    "full_name": "L",      # ПІБ
    "reason": "T",         # Причина
    "location": "W",       # Локація
    "date_out": "V",       # Дата убуття
    "date_in": "X",        # Дата прибуття
}

# Заголовки у DOCX -> поля зі штатки
DOCX_HEADER_ALIASES = {
    "department": ["відділення"],
    "position": ["посада"],
    "full_position": ["повна назва посади"],
    "rank": ["звання"],
    "full_name": ["піб", "прізвище ім'я по батькові", "прізвище, ім'я, по батькові"],
}

QUESTIONABLE_MARKERS = [
    "арешт",
    "зникн",
    "звільн",
    "полон",
    "сзч",
    "смерт",
    "загиб",
]

HEADER_SEARCH_ROW_LIMIT = 5
EXCEL_HEADER_SCAN_MAX_ROWS = 50
EXCEL_HEADER_GROUP_WINDOW = 5
TARGET_FILE_EXTENSIONS = {".docx", ".xlsx", ".xlsm"}
UPDATABLE_HEADER_FIELDS = {
    "department": "department",
    "position": "position",
    "full_position": "full_position",
    "rank": "rank",
}
DOCX_CHANGED_FILL = "FFFF00"
EXCEL_CHANGED_FILL_COLOR = 65535
EXCEL_SOLID_FILL_PATTERN = 1
OUR_UNIT_NAMES = ("3 МБ", "УПР", "зш")
OUR_UNIT_PRIORITY = {unit_name: idx for idx, unit_name in enumerate(OUR_UNIT_NAMES)}
SKIPPED_FULL_NAME_MARKERS = ("черговий",)
REPORT_BASENAME = "Report"
LOG_BASENAME = "OS_debug"
CURRENT_LOG_PATH: Optional[Path] = None


def install_frozen_executable_icon(root, retry_ms: int = 250) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    _set_frozen_executable_icon(root)
    try:
        root.after(retry_ms, lambda: _set_frozen_executable_icon(root))
    except Exception:
        pass


def _set_frozen_executable_icon(root) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    try:
        import ctypes
        from ctypes import wintypes

        wm_seticon = 0x0080
        icon_small = 0
        icon_big = 1
        gclp_hicon = -14
        gclp_hiconsm = -34

        shell32 = ctypes.WinDLL("shell32", use_last_error=True)
        user32 = ctypes.WinDLL("user32", use_last_error=True)
        long_ptr = (
            ctypes.c_longlong if ctypes.sizeof(ctypes.c_void_p) == 8 else ctypes.c_long
        )

        shell32.ExtractIconExW.argtypes = (
            wintypes.LPCWSTR,
            ctypes.c_int,
            ctypes.POINTER(wintypes.HICON),
            ctypes.POINTER(wintypes.HICON),
            wintypes.UINT,
        )
        shell32.ExtractIconExW.restype = wintypes.UINT
        user32.SendMessageW.argtypes = (
            wintypes.HWND,
            wintypes.UINT,
            wintypes.WPARAM,
            wintypes.LPARAM,
        )
        user32.SendMessageW.restype = long_ptr
        user32.GetParent.argtypes = (wintypes.HWND,)
        user32.GetParent.restype = wintypes.HWND

        try:
            set_class_icon = user32.SetClassLongPtrW
        except AttributeError:
            set_class_icon = user32.SetClassLongW
        set_class_icon.argtypes = (wintypes.HWND, ctypes.c_int, long_ptr)
        set_class_icon.restype = long_ptr

        icon_handles = getattr(root, "_frozen_executable_icon_handles", None)
        if icon_handles is None:
            large_icons = (wintypes.HICON * 1)()
            small_icons = (wintypes.HICON * 1)()
            icon_count = shell32.ExtractIconExW(
                sys.executable,
                0,
                large_icons,
                small_icons,
                1,
            )
            if icon_count <= 0:
                return
            icon_handles = (small_icons[0], large_icons[0])
            root._frozen_executable_icon_handles = icon_handles
        small_icon, large_icon = icon_handles

        root.update_idletasks()
        hwnds = []
        for raw_hwnd in (root.winfo_id(), root.wm_frame()):
            try:
                hwnd = int(str(raw_hwnd), 0)
            except (TypeError, ValueError):
                continue
            while hwnd and hwnd not in hwnds:
                hwnds.append(hwnd)
                hwnd = user32.GetParent(hwnd)

        for hwnd in hwnds:
            if small_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_small, small_icon)
                set_class_icon(hwnd, gclp_hiconsm, small_icon)
            if large_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_big, large_icon)
                set_class_icon(hwnd, gclp_hicon, large_icon)
    except Exception:
        pass


def install_dark_title_bar(window, retry_ms: int = 80) -> None:
    if sys.platform != "win32":
        return

    def _apply() -> None:
        try:
            import ctypes

            hwnd = ctypes.windll.user32.GetParent(window.winfo_id())
            value = ctypes.c_int(1)
            for attr in (20, 19):
                ctypes.windll.dwmapi.DwmSetWindowAttribute(
                    hwnd,
                    attr,
                    ctypes.byref(value),
                    ctypes.sizeof(value),
                )
        except Exception:
            pass

    try:
        window.update_idletasks()
        _apply()
        window.after(retry_ms, _apply)
    except Exception:
        pass


try:
    import tkinter as tk
    import tkinter.font as tkfont
    from tkinter import filedialog, messagebox, ttk
except ImportError:
    tk = None
    tkfont = None
    filedialog = None
    messagebox = None
    ttk = None

try:
    import win32com.client as win32
except ImportError:
    win32 = None

try:
    import pythoncom
except ImportError:
    pythoncom = None

# ============================================================
# МОДЕЛІ ДАНИХ
# ============================================================

@dataclass
class StaffRecord:
    full_name_raw: str
    full_name_key: str
    department: str = ""
    position: str = ""
    full_position: str = ""
    rank: str = ""
    unit: str = ""
    reason: str = ""
    location: str = ""
    date_out: str = ""
    date_in: str = ""


@dataclass
class QuestionableEntry:
    full_name: str
    reason: str


@dataclass
class MissingEntry:
    full_name: str
    comment: str = "відсутній у штатці (джерелі)"


@dataclass
class RegularEntry:
    full_name: str
    rank: str = ""
    department: str = ""
    position: str = ""
    reason: str = ""
    location: str = ""
    date_out: str = ""
    date_in: str = ""


@dataclass
class FileReport:
    file_name: str
    has_matching_headers: bool = False
    was_changed: bool = False
    questionable: List[QuestionableEntry] = field(default_factory=list)
    missing: List[MissingEntry] = field(default_factory=list)
    regular: List[RegularEntry] = field(default_factory=list)


@dataclass
class ExcelHeaderMatch:
    logical_name: str
    matched_alias: str
    row: int
    col: int


# ============================================================
# НОРМАЛІЗАЦІЯ ТЕКСТУ
# ============================================================


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = str(value).strip()
    text = re.sub(r"\s+", " ", text)
    return text



def normalize_name(value: object) -> str:
    text = normalize_text(value).lower()
    text = text.replace("'", "").replace("`", "").replace("’", "")
    text = re.sub(r"[^\w\sіїєґІЇЄҐ-]", " ", text, flags=re.UNICODE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def should_skip_full_name(value: object) -> bool:
    haystack = normalize_text(value).lower()
    return any(marker in haystack for marker in SKIPPED_FULL_NAME_MARKERS)



def normalize_header(value: object) -> str:
    text = normalize_text(value).lower()
    text = text.replace("'", "").replace("`", "").replace("’", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()



def matches_marker(reason_text: str, markers: Iterable[str]) -> Optional[str]:
    haystack = normalize_text(reason_text).lower()
    for marker in markers:
        if marker.lower() in haystack:
            return marker
    return None


def normalize_for_change_detection(value: object) -> str:
    text = normalize_text(value).lower().replace("\xa0", " ")
    return re.sub(r"\s+", "", text)


def value_changed_for_highlight(old_value: object, new_value: object) -> bool:
    return normalize_for_change_detection(old_value) != normalize_for_change_detection(new_value)


def make_run_timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H.%M")


def build_report_filename(timestamp: str) -> str:
    return f"{REPORT_BASENAME} - {timestamp}.log"


def build_debug_log_filename(timestamp: str) -> str:
    return f"{LOG_BASENAME} - {timestamp}.log"


def send_debug_log_to_trash() -> None:
    global CURRENT_LOG_PATH
    if CURRENT_LOG_PATH is None or not CURRENT_LOG_PATH.exists():
        return
    logging.shutdown()
    try:
        send2trash(str(CURRENT_LOG_PATH))
    except Exception:
        pass
    CURRENT_LOG_PATH = None


# ============================================================
# EXCEL / ШТАТКА
# ============================================================


def cell_to_string(ws, row_idx: int, col_letter: str) -> str:
    return normalize_text(ws[f"{col_letter}{row_idx}"].value)


def get_unit_priority(unit_name: str) -> Optional[int]:
    return OUR_UNIT_PRIORITY.get(unit_name)


def first_paragraph_text(value: object) -> str:
    if value is None:
        return ""
    first_line = str(value).replace("\r\n", "\n").replace("\r", "\n").split("\n", 1)[0]
    return normalize_text(first_line)



def load_staff_records(xlsx_path: Path) -> Dict[str, StaffRecord]:
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    if SHEET_NAME not in wb.sheetnames:
        raise ValueError(f"У файлі немає аркуша '{SHEET_NAME}'")

    ws = wb[SHEET_NAME]
    records: Dict[str, StaffRecord] = {}

    for row_idx in range(2, ws.max_row + 1):
        unit_name = cell_to_string(ws, row_idx, XLSX_COLS["unit"])
        unit_priority = get_unit_priority(unit_name)
        if unit_priority is None:
            continue

        full_name = cell_to_string(ws, row_idx, XLSX_COLS["full_name"])
        if not full_name:
            continue
        if should_skip_full_name(full_name):
            continue

        key = normalize_name(full_name)
        if not key:
            continue

        record = StaffRecord(
            full_name_raw=full_name,
            full_name_key=key,
            unit=unit_name,
            department=cell_to_string(ws, row_idx, XLSX_COLS["department"]),
            position=first_paragraph_text(ws[f"{XLSX_COLS['position']}{row_idx}"].value),
            full_position=cell_to_string(ws, row_idx, XLSX_COLS["full_position"]),
            rank=cell_to_string(ws, row_idx, XLSX_COLS["rank"]),
            reason=cell_to_string(ws, row_idx, XLSX_COLS["reason"]),
            location=cell_to_string(ws, row_idx, XLSX_COLS["location"]),
            date_out=cell_to_string(ws, row_idx, XLSX_COLS["date_out"]),
            date_in=cell_to_string(ws, row_idx, XLSX_COLS["date_in"]),
        )

        current_record = records.get(key)
        if current_record is None:
            records[key] = record
            continue

        current_priority = get_unit_priority(current_record.unit)
        if current_priority is None:
            records[key] = record
        elif unit_priority < current_priority:
            logging.info(
                "Дублікат ПІБ '%s': запис із підрозділу '%s' має вищий пріоритет за '%s'.",
                full_name,
                unit_name,
                current_record.unit,
            )
            records[key] = record
        elif unit_priority == current_priority:
            records[key] = record
        else:
            logging.info(
                "Дублікат ПІБ '%s': запис із підрозділу '%s' проігноровано, бо '%s' має вищий пріоритет.",
                full_name,
                unit_name,
                current_record.unit,
            )

    return records


# ============================================================
# DOCX / TARGET FILES
# ============================================================


def is_supported_target_file(path: Path) -> bool:
    return (
        path.is_file()
        and not path.name.startswith(".")
        and not path.name.startswith("~$")
        and path.suffix.lower() in TARGET_FILE_EXTENSIONS
    )


def find_target_files(folder: Path, excluded_paths: Optional[Iterable[Path]] = None) -> List[Path]:
    excluded = {path.expanduser().resolve() for path in (excluded_paths or [])}
    files: List[Path] = []

    for path in folder.rglob("*"):
        if not is_supported_target_file(path):
            continue
        resolved = path.resolve()
        if resolved in excluded:
            continue
        files.append(resolved)

    return sorted(files, key=lambda path: str(path).lower())



def iter_tables_recursive(document: Document) -> Iterable[Table]:
    for table in document.tables:
        yield table
        yield from iter_nested_tables(table)



def iter_nested_tables(table: Table) -> Iterable[Table]:
    for row in table.rows:
        for cell in row.cells:
            for nested in cell.tables:
                yield nested
                yield from iter_nested_tables(nested)



def detect_header_map(table: Table) -> Optional[Tuple[Dict[str, int], int]]:
    rows = list(table.rows)
    if not rows:
        return None

    best_match: Optional[Tuple[Dict[str, int], int]] = None
    best_score = -1

    for row_idx, row in enumerate(rows[:HEADER_SEARCH_ROW_LIMIT]):
        resolved: Dict[str, int] = {}

        for cell_idx, cell in enumerate(row.cells):
            header_text = normalize_header(cell.text)
            if not header_text:
                continue
            for target_field, aliases in DOCX_HEADER_ALIASES.items():
                if any(alias == header_text for alias in aliases):
                    resolved[target_field] = cell_idx

        if "full_name" not in resolved:
            continue

        score = len(resolved)
        if score > best_score:
            best_score = score
            best_match = (resolved, row_idx)

    return best_match



def clear_cell_but_keep_style(cell: _Cell) -> None:
    # Залишаємо перший параграф, очищаємо його рани. Це грубувато, але працює як старт.
    if not cell.paragraphs:
        cell.text = ""
        return

    first_paragraph = cell.paragraphs[0]
    for paragraph in cell.paragraphs[1:]:
        p = paragraph._element
        p.getparent().remove(p)

    for run in first_paragraph.runs:
        run.text = ""

    if not first_paragraph.runs:
        first_paragraph.add_run("")



def write_text_preserve_basic_style(cell: _Cell, value: str) -> None:
    clear_cell_but_keep_style(cell)
    paragraph = cell.paragraphs[0]

    if paragraph.runs:
        paragraph.runs[0].text = value
    else:
        paragraph.add_run(value)


def apply_docx_cell_fill(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)



def update_table_row(
    row,
    header_map: Dict[str, int],
    record: StaffRecord,
) -> bool:
    has_changes = False
    for docx_field, record_attr in UPDATABLE_HEADER_FIELDS.items():
        col_idx = header_map.get(docx_field)
        if col_idx is None:
            continue
        if col_idx >= len(row.cells):
            continue

        target_cell = row.cells[col_idx]
        new_value = getattr(record, record_attr, "")
        is_changed = value_changed_for_highlight(target_cell.text, new_value)
        if not is_changed:
            continue

        write_text_preserve_basic_style(target_cell, new_value)
        apply_docx_cell_fill(target_cell, DOCX_CHANGED_FILL)
        has_changes = True

    return has_changes


def register_staff_record_in_report(
    full_name_from_file: str,
    staff_map: Dict[str, StaffRecord],
    report: FileReport,
    seen_regular: Dict[str, RegularEntry],
    seen_questionable: set[Tuple[str, str]],
    seen_missing: set[str],
) -> Optional[StaffRecord]:
    if should_skip_full_name(full_name_from_file):
        return None

    full_name_key = normalize_name(full_name_from_file)
    if not full_name_key:
        return None

    record = staff_map.get(full_name_key)
    if not record:
        if full_name_from_file not in seen_missing:
            seen_missing.add(full_name_from_file)
            report.missing.append(MissingEntry(full_name=full_name_from_file))
        return None

    marker = matches_marker(record.reason, QUESTIONABLE_MARKERS)
    if marker:
        pair = (record.full_name_raw, record.reason)
        if pair not in seen_questionable:
            seen_questionable.add(pair)
            report.questionable.append(
                QuestionableEntry(
                    full_name=record.full_name_raw,
                    reason=record.reason,
                )
            )
    else:
        seen_regular[record.full_name_raw] = RegularEntry(
            full_name=record.full_name_raw,
            rank=record.rank,
            department=record.department,
            position=record.position,
            reason=record.reason,
            location=record.location,
            date_out=record.date_out,
            date_in=record.date_in,
        )

    return record



def process_docx_file(
    docx_path: Path,
    staff_map: Dict[str, StaffRecord],
    update_data: bool = True,
) -> FileReport:
    document = Document(docx_path)
    report = FileReport(file_name=docx_path.name)

    seen_regular: Dict[str, RegularEntry] = {}
    seen_questionable: set[Tuple[str, str]] = set()
    seen_missing: set[str] = set()
    has_changes = False

    for table in iter_tables_recursive(document):
        header_match = detect_header_map(table)
        if not header_match:
            continue
        report.has_matching_headers = True
        header_map, header_row_idx = header_match
        rows = list(table.rows)

        full_name_col = header_map["full_name"]

        for row in rows[header_row_idx + 1:]:
            if full_name_col >= len(row.cells):
                continue

            full_name_docx = normalize_text(row.cells[full_name_col].text)
            if not full_name_docx:
                continue

            record = register_staff_record_in_report(
                full_name_docx,
                staff_map,
                report,
                seen_regular,
                seen_questionable,
                seen_missing,
            )
            if not record:
                continue

            if update_data:
                if update_table_row(row, header_map, record):
                    has_changes = True

    report.regular = sorted(seen_regular.values(), key=lambda x: normalize_name(x.full_name))
    if update_data and has_changes:
        document.save(docx_path)
        report.was_changed = True
    elif update_data:
        logging.info("DOCX без змін, файл не перезаписано: %s", docx_path)
    return report


def is_exact_header_match(cell_value: object, aliases: Iterable[str]) -> Optional[str]:
    normalized = normalize_header(cell_value)
    if not normalized:
        return None

    for alias in aliases:
        if normalized == normalize_header(alias):
            return alias
    return None


def load_target_workbook(workbook_path: Path, update_data: bool):
    is_macro_book = workbook_path.suffix.lower() == ".xlsm"
    return openpyxl.load_workbook(
        workbook_path,
        data_only=True,
        keep_vba=is_macro_book,
    )


def find_excel_header_matches(worksheet) -> List[ExcelHeaderMatch]:
    matches: List[ExcelHeaderMatch] = []
    max_scan_row = min(worksheet.max_row, EXCEL_HEADER_SCAN_MAX_ROWS)

    for row in worksheet.iter_rows(min_row=1, max_row=max_scan_row):
        for cell in row:
            for logical_name, aliases in DOCX_HEADER_ALIASES.items():
                matched_alias = is_exact_header_match(cell.value, aliases)
                if matched_alias is not None:
                    matches.append(
                        ExcelHeaderMatch(
                            logical_name=logical_name,
                            matched_alias=matched_alias,
                            row=cell.row,
                            col=cell.column,
                        )
                    )

    return matches


def detect_excel_header_context(worksheet) -> Optional[Tuple[Dict[str, int], int]]:
    matches = find_excel_header_matches(worksheet)
    full_name_matches = sorted(
        [match for match in matches if match.logical_name == "full_name"],
        key=lambda match: (match.row, match.col),
    )
    best_match: Optional[Tuple[Dict[str, int], int]] = None
    best_signature = (-1, -1)

    for anchor in full_name_matches:
        grouped = [match for match in matches if abs(match.row - anchor.row) <= EXCEL_HEADER_GROUP_WINDOW]
        if not grouped:
            continue

        grouped.sort(key=lambda match: (abs(match.row - anchor.row), match.row, match.col))

        resolved: Dict[str, int] = {}
        for match in grouped:
            if match.logical_name not in resolved:
                resolved[match.logical_name] = match.col

        if "full_name" not in resolved:
            continue

        bottom_header_row = max(match.row for match in grouped)
        signature = (len(resolved), bottom_header_row)
        if signature > best_signature:
            best_signature = signature
            best_match = (resolved, bottom_header_row + 1)

    return best_match


def apply_excel_updates_via_com(
    workbook_path: Path,
    header_map: Dict[str, int],
    updates: List[Tuple[int, StaffRecord]],
) -> bool:
    if not updates:
        return False

    if win32 is None:
        raise RuntimeError(
            "Для безпечного оновлення XLSX/XLSM потрібен встановлений pywin32 "
            "(win32com.client), інакше Excel-книги краще не перезаписувати."
        )

    excel = None
    workbook = None

    try:
        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        try:
            excel.AskToUpdateLinks = False
        except Exception:
            pass

        workbook = excel.Workbooks.Open(str(workbook_path), UpdateLinks=0, ReadOnly=False)
        worksheet = workbook.Worksheets(1)
        has_changes = False

        for row_idx, record in updates:
            for logical_name, record_attr in UPDATABLE_HEADER_FIELDS.items():
                col_idx = header_map.get(logical_name)
                if col_idx is None:
                    continue

                cell = worksheet.Cells(row_idx, col_idx)
                target_range = cell
                try:
                    if bool(cell.MergeCells):
                        target_range = cell.MergeArea
                        cell = target_range.Cells(1, 1)
                except Exception:
                    pass

                new_value = getattr(record, record_attr, "")
                is_changed = value_changed_for_highlight(cell.Value, new_value)
                if not is_changed:
                    continue

                cell.Value = new_value
                target_range.Interior.Pattern = EXCEL_SOLID_FILL_PATTERN
                target_range.Interior.Color = EXCEL_CHANGED_FILL_COLOR
                has_changes = True

        if has_changes:
            workbook.Save()
        return has_changes
    finally:
        if workbook is not None:
            try:
                workbook.Close(SaveChanges=False)
            except Exception:
                pass
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass


def process_excel_file(
    workbook_path: Path,
    staff_map: Dict[str, StaffRecord],
    update_data: bool = True,
) -> FileReport:
    workbook = load_target_workbook(workbook_path, update_data=update_data)
    report = FileReport(file_name=workbook_path.name)
    pending_updates: List[Tuple[int, StaffRecord]] = []
    header_map: Optional[Dict[str, int]] = None

    try:
        worksheets = workbook.worksheets
        if not worksheets:
            raise ValueError(f"У файлі '{workbook_path.name}' немає доступних аркушів.")

        worksheet = worksheets[0]
        header_context = detect_excel_header_context(worksheet)
        if not header_context:
            return report

        report.has_matching_headers = True
        header_map, data_start_row = header_context
        logging.info(
            "Excel-книга '%s': аркуш '%s', старт даних з рядка %s",
            workbook_path.name,
            worksheet.title,
            data_start_row,
        )

        seen_regular: Dict[str, RegularEntry] = {}
        seen_questionable: set[Tuple[str, str]] = set()
        seen_missing: set[str] = set()

        full_name_col = header_map["full_name"]

        for row_idx in range(data_start_row, worksheet.max_row + 1):
            full_name_excel = normalize_text(worksheet.cell(row=row_idx, column=full_name_col).value)
            if not full_name_excel:
                continue

            record = register_staff_record_in_report(
                full_name_excel,
                staff_map,
                report,
                seen_regular,
                seen_questionable,
                seen_missing,
            )
            if not record:
                continue

            if update_data:
                pending_updates.append((row_idx, record))

        report.regular = sorted(seen_regular.values(), key=lambda x: normalize_name(x.full_name))
    finally:
        try:
            workbook.close()
        except Exception:
            pass

    if update_data and header_map is not None:
        has_changes = apply_excel_updates_via_com(workbook_path, header_map, pending_updates)
        if has_changes:
            report.was_changed = True
        else:
            logging.info("Excel без змін, файл не перезаписано: %s", workbook_path)

    return report


def process_target_file(
    file_path: Path,
    staff_map: Dict[str, StaffRecord],
    update_data: bool = True,
) -> FileReport:
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return process_docx_file(file_path, staff_map, update_data=update_data)
    if suffix in {".xlsx", ".xlsm"}:
        return process_excel_file(file_path, staff_map, update_data=update_data)
    raise ValueError(f"Непідтримуваний тип файлу: {file_path.suffix}")


# ============================================================
# BACKUP
# ============================================================

def create_folder_backup_in_trash(
    folder: Path,
    timestamp: str,
    progress: ProgressWindow | None = None,
    stage_header: str = "Етап 1/3: Резервна копія",
) -> str:
    """
    Створює повну копію папки поруч із оригіналом, додає суфікс Backup + дата/час,
    після чого відправляє цю копію в кошик.

    Повертає ім'я backup-папки для логів / звіту.
    """
    backup_name = f"{folder.name} - {timestamp}"
    backup_path = folder.parent / backup_name

    if backup_path.exists():
        raise FileExistsError(f"Backup-папка вже існує: {backup_path}")

    if progress:
        progress.update(
            header=stage_header,
            detail="Створення резервної копії та переміщення у кошик...",
            current=0,
            total=1,
            file_name=folder.name,
        )
    shutil.copytree(folder, backup_path)
    send2trash(str(backup_path))
    if progress:
        progress.update(
            header=stage_header,
            detail="Резервну копію створено",
            current=1,
            total=1,
            file_name=backup_name,
        )

    return backup_name


# ============================================================
# ЗВІТ
# ============================================================


def format_regular_entry(entry: RegularEntry) -> str:
    payload = [
        entry.rank,
        entry.department,
        entry.position,
        entry.reason,
        entry.location,
        entry.date_out,
        entry.date_in,
    ]
    return f"{entry.full_name}: " + "; ".join(value for value in payload if value)


def format_report_file_name(file_report: FileReport) -> str:
    changed_marker = " (файл змінено)" if file_report.was_changed else ""
    return f"'{file_report.file_name}'{changed_marker}"


def build_questionable_lines(file_report: FileReport) -> List[str]:
    lines: List[str] = []
    for item in sorted(file_report.missing, key=lambda x: normalize_name(x.full_name)):
        lines.append(f"{item.full_name}: {item.comment}")
    for item in sorted(file_report.questionable, key=lambda x: normalize_name(x.full_name)):
        lines.append(f"{item.full_name}: {item.reason}")
    return lines


def report_issue_count(file_report: FileReport) -> int:
    return len(file_report.missing) + len(file_report.questionable)


def build_report_text(
    file_reports: List[FileReport],
    run_timestamp: str,
    staff_file_name: str,
    work_folder: Path,
) -> str:
    lines: List[str] = []
    visible_reports = [file_report for file_report in file_reports if file_report.has_matching_headers]

    lines.append(f'Папка: "{work_folder}"')
    lines.append(f'Штатка: "{staff_file_name}"')
    lines.append(f"ДЗ: {run_timestamp}")
    lines.append("")

    lines.append("---- ПІД ПИТАННЯМ:")
    lines.append("")

    no_issue_reports = sorted(
        [file_report for file_report in visible_reports if report_issue_count(file_report) == 0],
        key=lambda x: x.file_name.lower(),
    )
    issue_reports = sorted(
        [file_report for file_report in visible_reports if report_issue_count(file_report) > 0],
        key=lambda x: (report_issue_count(x), x.file_name.lower()),
    )

    if no_issue_reports:
        lines.append("НЕМАЄ ЗАПИСІВ:")
        for file_report in no_issue_reports:
            lines.append(format_report_file_name(file_report))
        lines.append("")

    if not visible_reports:
        lines.append("Немає файлів із заголовками для звірки.")
        lines.append("")

    for file_report in issue_reports:
        lines.append(format_report_file_name(file_report))
        lines.extend(build_questionable_lines(file_report))
        lines.append("")

    lines.append("--- УСЕ ІНШЕ:")
    lines.append("")
    aggregated_regular: Dict[str, RegularEntry] = {}

    for file_report in visible_reports:
        for entry in file_report.regular:
            aggregated_regular[entry.full_name] = entry

    for entry in sorted(aggregated_regular.values(), key=lambda x: normalize_name(x.full_name)):
        lines.append(format_regular_entry(entry))

    return "\n".join(lines).strip() + "\n"



def write_report(
    root_folder: Path,
    file_reports: List[FileReport],
    timestamp: str,
    staff_file_name: str,
    work_folder: Path,
) -> Path:
    report_path = root_folder / build_report_filename(timestamp)
    report_text = build_report_text(file_reports, timestamp, staff_file_name, work_folder)
    report_path.write_text(report_text, encoding="utf-8")
    return report_path


# ============================================================
# GUI
# ============================================================


def configure_launch_styles(root: Any) -> Dict[str, str]:
    colors = {
        "window": "#F4F7FB",
        "panel": "#FFFFFF",
        "border": "#D6DEE8",
        "text": "#15202B",
        "muted": "#617084",
        "accent": "#0F766E",
        "accent_dark": "#0B5D56",
        "header": "#103C3A",
        "header_badge": "#0F766E",
    }

    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except Exception:
        pass

    try:
        style.configure("LaunchRoot.TFrame", background=colors["window"])
        style.configure("LaunchPanel.TFrame", background=colors["panel"])
        style.configure("LaunchSection.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        style.configure("LaunchField.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10, "bold"))
        style.configure("LaunchMuted.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("LaunchTiny.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 8))
        style.configure("LaunchStatus.TLabel", background=colors["window"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("ProgressHeader.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        style.configure("ProgressBody.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10))
        style.configure("ProgressFile.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("ProgressCount.TLabel", background=colors["panel"], foreground=colors["accent_dark"], font=("Segoe UI", 10, "bold"))
        style.configure(
            "Launch.Horizontal.TProgressbar",
            background=colors["header_badge"],
            troughcolor="#E8EEF5",
            bordercolor=colors["border"],
            lightcolor=colors["header_badge"],
            darkcolor=colors["header_badge"],
            thickness=12,
        )
        style.configure(
            "Launch.TEntry",
            fieldbackground="#FFFFFF",
            foreground=colors["text"],
            bordercolor=colors["border"],
            lightcolor=colors["border"],
            darkcolor=colors["border"],
            padding=(8, 6),
        )
        style.map(
            "Launch.TEntry",
            fieldbackground=[("readonly", "#FFFFFF")],
            foreground=[("readonly", colors["text"])],
        )
        style.configure(
            "LaunchBrowse.TButton",
            font=("Segoe UI", 10, "bold"),
            padding=(18, 10),
            foreground="#FFFFFF",
            background="#053D39",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchBrowse.TButton",
            background=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
            bordercolor=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
        )
        style.configure(
            "LaunchPrimary.TButton",
            font=("Segoe UI", 11, "bold"),
            padding=(20, 14),
            foreground="#FFFFFF",
            background="#053D39",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.configure(
            "LaunchPrimaryHover.TButton",
            font=("Segoe UI", 11, "bold"),
            padding=(20, 14),
            foreground="#FFFFFF",
            background=colors["accent_dark"],
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchPrimary.TButton",
            background=[("pressed", "#042F2C"), ("active", colors["accent_dark"]), ("disabled", "#B8C2CC")],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF"), ("disabled", "#EEF2F7")],
            bordercolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
        )
        style.map(
            "LaunchPrimaryHover.TButton",
            background=[("pressed", "#042F2C"), ("active", colors["accent_dark"]), ("disabled", "#B8C2CC")],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF"), ("disabled", "#EEF2F7")],
            bordercolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
        )
    except Exception:
        pass

    return colors


class ProgressWindow:
    def __init__(self, owner: Any):
        if tk is None or ttk is None:
            raise RuntimeError("Tkinter недоступний.")

        self.close_requested = False
        self.dialog = tk.Toplevel(owner)
        self.dialog.withdraw()
        install_frozen_executable_icon(self.dialog)
        self.dialog.title("AuditOS - Обробка")
        self.dialog.resizable(False, False)
        self.dialog.attributes("-topmost", True)
        self.dialog.protocol("WM_DELETE_WINDOW", self.request_close)
        colors = configure_launch_styles(self.dialog)
        self.dialog.configure(bg=colors["window"])
        install_dark_title_bar(self.dialog)

        container = ttk.Frame(self.dialog, style="LaunchRoot.TFrame")
        container.pack(fill="both", expand=True)

        header = tk.Frame(container, bg=colors["header"], padx=22, pady=10)
        header.pack(fill="x")
        header.grid_columnconfigure(0, weight=1)
        tk.Label(
            header,
            text="AuditOS",
            bg=colors["header"],
            fg="#FFFFFF",
            font=("Segoe UI Semibold", 15),
            anchor="w",
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            header,
            text="Аналіз документів виконується. Будь ласка, дочекайтеся завершення.",
            bg=colors["header"],
            fg="#D7FBF5",
            font=("Segoe UI", 9),
            anchor="w",
        ).grid(row=1, column=0, sticky="w", pady=(0, 4))
        badge_shell = tk.Frame(header, bg=colors["accent_dark"], padx=1, pady=1)
        badge = tk.Label(
            badge_shell,
            text="ОБРОБКА",
            bg="#053D39",
            fg="#FFFFFF",
            font=("Segoe UI", 12, "bold"),
            padx=10,
            pady=5,
        )
        badge.pack()
        badge_shell.place(relx=1.0, x=5, y=5, anchor="ne")

        body = ttk.Frame(container, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
        body.pack(fill="both", expand=True)
        panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
        panel_shell.pack(fill="both", expand=True)
        frame = ttk.Frame(panel_shell, style="LaunchPanel.TFrame", padding=(16, 14, 16, 12))
        frame.pack(fill="both", expand=True)
        frame.grid_columnconfigure(0, weight=1)

        self.header_var = tk.StringVar(value="Підготовка...")
        self.detail_var = tk.StringVar(value="Будь ласка, зачекайте")
        self.file_var = tk.StringVar(value="")
        self.count_var = tk.StringVar(value="0 / 0")

        ttk.Label(frame, textvariable=self.header_var, style="ProgressHeader.TLabel").grid(
            row=0, column=0, padx=(0, 12), sticky="w"
        )
        ttk.Label(frame, textvariable=self.count_var, style="ProgressCount.TLabel").grid(
            row=0, column=1, sticky="e"
        )
        ttk.Label(frame, textvariable=self.detail_var, style="ProgressBody.TLabel", wraplength=500).grid(
            row=1, column=0, columnspan=2, pady=(8, 6), sticky="w"
        )
        ttk.Label(frame, textvariable=self.file_var, style="ProgressFile.TLabel", wraplength=500).grid(
            row=2, column=0, columnspan=2, pady=(0, 12), sticky="w"
        )
        self.progress = ttk.Progressbar(
            frame,
            orient="horizontal",
            mode="determinate",
            maximum=100,
            style="Launch.Horizontal.TProgressbar",
        )
        self.progress.grid(row=3, column=0, columnspan=2, sticky="we")
        self.dialog.update_idletasks()
        width = max(560, self.dialog.winfo_reqwidth())
        height = max(248, self.dialog.winfo_reqheight())
        self.dialog.geometry(f"{width}x{height}")
        self.dialog.deiconify()
        self.dialog.lift()
        self.refresh()

    def update(self, header=None, detail=None, current=None, total=None, file_name=None):
        try:
            if not self.dialog.winfo_exists():
                return
            if header is not None:
                self.header_var.set(header)
            if detail is not None:
                self.detail_var.set(detail)
            if file_name is not None:
                self.file_var.set(file_name)
            if current is not None and total is not None:
                safe_total = total if total > 0 else 1
                pct = max(0.0, min(100.0, (current / safe_total) * 100.0))
                self.progress["value"] = pct
                self.count_var.set(f"{current} / {total}")
            self.refresh()
        except tk.TclError:
            return

    def refresh(self):
        try:
            if self.close_requested or not _widget_exists(self.dialog):
                return
            self.dialog.update_idletasks()
            self.dialog.update()
        except Exception:
            pass

    def close(self):
        self.close_requested = True
        try:
            if _widget_exists(self.dialog):
                self.dialog.destroy()
        except Exception:
            pass

    def request_close(self):
        self.close()

    def show_success_then_close(self, detail: str, file_name: str = "", delay_ms: int = 1500):
        if self.close_requested or not _widget_exists(self.dialog):
            return
        self.update(
            header="Готово",
            detail=detail,
            current=1,
            total=1,
            file_name=file_name,
        )
        deadline = time.monotonic() + delay_ms / 1000.0
        while time.monotonic() < deadline:
            if not _widget_exists(self.dialog):
                return
            self.refresh()
            time.sleep(0.05)
        self.close()


class ProgressUpdateProxy:
    def __init__(self, updates: "queue.Queue[Dict[str, Any]]"):
        self.updates = updates

    def update(self, header=None, detail=None, current=None, total=None, file_name=None):
        self.updates.put(
            {
                "header": header,
                "detail": detail,
                "current": current,
                "total": total,
                "file_name": file_name,
            }
        )


def apply_pending_progress_updates(
    progress: ProgressWindow,
    updates: "queue.Queue[Dict[str, Any]]",
) -> None:
    while True:
        try:
            payload = updates.get_nowait()
        except queue.Empty:
            return
        progress.update(**payload)


def run_process_all_with_progress(
    xlsx_path: Path,
    docx_folder: Path,
    update_data: bool,
    progress: ProgressWindow,
) -> Tuple[List[FileReport], Path]:
    updates: "queue.Queue[Dict[str, Any]]" = queue.Queue()
    result: Dict[str, Any] = {"payload": None, "error": None}
    progress_proxy = ProgressUpdateProxy(updates)

    def worker() -> None:
        com_initialized = False
        try:
            if pythoncom is not None:
                pythoncom.CoInitialize()
                com_initialized = True
            result["payload"] = process_all(
                xlsx_path,
                docx_folder,
                update_data=update_data,
                progress=progress_proxy,
            )
        except Exception as exc:
            result["error"] = exc
        finally:
            if com_initialized:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()

    while thread.is_alive():
        apply_pending_progress_updates(progress, updates)
        progress.refresh()
        time.sleep(0.03)

    thread.join()
    apply_pending_progress_updates(progress, updates)
    progress.refresh()

    if result["error"] is not None:
        raise result["error"]
    if result["payload"] is None:
        raise RuntimeError("Обробка завершилася без результату.")
    return result["payload"]


ToggleSwitchBase = ttk.Frame if ttk is not None else object


class ToggleSwitch(ToggleSwitchBase):
    def __init__(self, owner: Any, variable: Any):
        super().__init__(owner)
        self.variable = variable
        self.state_var = tk.StringVar()

        style = ttk.Style(self)
        try:
            parent_style = owner.cget("style")
            if parent_style:
                self.configure(style=parent_style)
                canvas_bg = style.lookup(parent_style, "background")
            else:
                canvas_bg = ""
        except Exception:
            canvas_bg = ""
        canvas_bg = canvas_bg or style.lookup("LaunchPanel.TFrame", "background") or self.winfo_toplevel().cget("bg")
        state_label_style = "ToggleState.TLabel"
        try:
            style.configure(
                state_label_style,
                background="#F4F4F8",
                foreground="#15202B",
                font=("Segoe UI", 9),
            )
        except Exception:
            state_label_style = ""

        self.canvas = tk.Canvas(
            self,
            width=52,
            height=28,
            bd=0,
            highlightthickness=0,
            bg=canvas_bg,
            cursor="hand2",
        )
        self.canvas.grid(row=0, column=0, padx=(0, 8), sticky="w")

        self.state_label = ttk.Label(self, textvariable=self.state_var, style=state_label_style)
        self.state_label.grid(row=0, column=1, sticky="w")

        for widget in (self, self.canvas, self.state_label):
            widget.bind("<Button-1>", self.toggle)

        self.variable.trace_add("write", self._redraw)
        self._redraw()

    def toggle(self, _event=None) -> None:
        self.variable.set(not self.variable.get())

    def _redraw(self, *_args) -> None:
        is_enabled = bool(self.variable.get())
        track_color = "#0F766E" if is_enabled else "#CBD5E1"
        knob_color = "#FFFFFF"
        self.state_var.set("Увімкнено" if is_enabled else "Вимкнено")

        self.canvas.delete("all")

        left = 2
        top = 2
        right = 50
        bottom = 26
        radius = (bottom - top) // 2

        self.canvas.create_oval(left, top, left + 2 * radius, bottom, fill=track_color, outline=track_color)
        self.canvas.create_oval(right - 2 * radius, top, right, bottom, fill=track_color, outline=track_color)
        self.canvas.create_rectangle(left + radius, top, right - radius, bottom, fill=track_color, outline=track_color)

        knob_diameter = (2 * radius) - 2
        knob_left = right - knob_diameter - 1 if is_enabled else left + 1
        self.canvas.create_oval(
            knob_left,
            top + 1,
            knob_left + knob_diameter,
            bottom - 1,
            fill=knob_color,
            outline=knob_color,
        )


def build_stage_header(stage_idx: int, total_stages: int, title: str) -> str:
    return f"Етап {stage_idx}/{total_stages}: {title}"


def _widget_exists(widget: Any | None) -> bool:
    if tk is None or widget is None:
        return False
    try:
        return bool(widget.winfo_exists())
    except (tk.TclError, RuntimeError):
        return False


def _destroy_widget(widget: Any | None) -> None:
    if not _widget_exists(widget):
        return
    try:
        widget.destroy()
    except Exception:
        pass


def _prepare_dialog_parent(root: Any) -> None:
    if not _widget_exists(root):
        return
    try:
        try:
            current_grab = root.grab_current()
            if current_grab is not None and current_grab is not root:
                current_grab.grab_release()
        except Exception:
            pass
        if hasattr(root, "deiconify") and not root.winfo_ismapped():
            root.deiconify()
        root.lift()
        root.attributes("-topmost", True)
        root.update()
    except Exception:
        pass


def _show_topmost_message(kind: str, title: str, text: str, parent: Any | None = None) -> None:
    if tk is None or messagebox is None:
        stream = None
        try:
            import sys

            stream = sys.stderr if kind == "error" else sys.stdout
        except Exception:
            pass
        if stream is not None:
            print(f"{title}: {text}", file=stream)
        return

    if _widget_exists(parent):
        _prepare_dialog_parent(parent)
        if kind == "info":
            messagebox.showinfo(title, text, parent=parent)
        elif kind == "warning":
            messagebox.showwarning(title, text, parent=parent)
        else:
            messagebox.showerror(title, text, parent=parent)
        return

    root = tk.Tk()
    root.withdraw()
    install_frozen_executable_icon(root)
    install_dark_title_bar(root)
    root.attributes("-topmost", True)
    try:
        if kind == "info":
            messagebox.showinfo(title, text, parent=root)
        elif kind == "warning":
            messagebox.showwarning(title, text, parent=root)
        else:
            messagebox.showerror(title, text, parent=root)
    finally:
        _destroy_widget(root)


def format_path_for_entry_display(path: str) -> str:
    path = path.strip()
    if not path:
        return ""

    normalized = Path(path)
    drive_prefix = ""
    if normalized.drive and len(normalized.drive) == 2 and normalized.drive[1] == ":":
        drive_prefix = f"{normalized.drive}/"

    name = normalized.name
    parent_name = normalized.parent.name
    if parent_name and name:
        return f"{drive_prefix}.../{parent_name}/{name}"
    if name:
        return f"{drive_prefix}{name}" if drive_prefix else f".../{name}"
    if drive_prefix:
        return drive_prefix
    return path.replace("\\", "/")


def fit_path_for_entry_display(path: str, entry: Any) -> str:
    display_path = format_path_for_entry_display(path)
    if not display_path:
        return ""

    entry.update_idletasks()
    font_name = entry.cget("font") or "TkDefaultFont"
    try:
        font = tkfont.nametofont(font_name)
    except tk.TclError:
        try:
            font = tkfont.nametofont("TkDefaultFont")
        except tk.TclError:
            return display_path

    available_width = max(120, entry.winfo_width() - 14)
    if font.measure(display_path) <= available_width:
        return display_path

    parts = display_path.split("/")
    if len(parts) >= 3:
        prefix = "/".join(parts[:-1]) + "/..."
        shortened_name = parts[-1]
        while len(shortened_name) > 1 and font.measure(prefix + shortened_name) > available_width:
            shortened_name = shortened_name[1:]
        return prefix + shortened_name

    ellipsis = "..."
    shortened_path = display_path
    while len(shortened_path) > 1 and font.measure(ellipsis + shortened_path) > available_width:
        shortened_path = shortened_path[1:]
    return ellipsis + shortened_path


def refresh_path_entry_display(source_var: Any, display_var: Any, entry: Any) -> None:
    if not _widget_exists(entry):
        return
    display_var.set(fit_path_for_entry_display(source_var.get(), entry))


def choose_file(
    root: Any,
    source_var: Any,
    display_var: Any,
    entry: Any,
    file_types: List[Tuple[str, str]],
) -> None:
    _prepare_dialog_parent(root)
    path = filedialog.askopenfilename(filetypes=file_types, parent=root)
    if path:
        source_var.set(path)
        refresh_path_entry_display(source_var, display_var, entry)
    _prepare_dialog_parent(root)



def choose_folder(root: Any, source_var: Any, display_var: Any, entry: Any) -> None:
    _prepare_dialog_parent(root)
    path = filedialog.askdirectory(parent=root)
    if path:
        source_var.set(path)
        refresh_path_entry_display(source_var, display_var, entry)
    _prepare_dialog_parent(root)


def bind_primary_button_hover(button: Any) -> None:
    def on_enter(_event) -> None:
        button.configure(style="LaunchPrimaryHover.TButton")
        try:
            button.focus_force()
        except Exception:
            pass

    def on_leave(_event) -> None:
        button.configure(style="LaunchPrimary.TButton")
        try:
            button.winfo_toplevel().focus_force()
        except Exception:
            pass

    button.bind("<Enter>", on_enter, add="+")
    button.bind("<Leave>", on_leave, add="+")


def run_gui() -> Optional[Tuple[Path, Path, bool]]:
    if tk is None or ttk is None or filedialog is None:
        raise RuntimeError("Tkinter недоступний: неможливо показати вікно налаштувань.")

    selected: Dict[str, object] = {"xlsx": "", "folder": "", "update_data": False, "confirmed": False}

    def on_start() -> None:
        selected["xlsx"] = xlsx_var.get().strip()
        selected["folder"] = folder_var.get().strip()
        selected["update_data"] = bool(update_var.get())
        if not selected["xlsx"] or not selected["folder"]:
            _show_topmost_message("error", "Помилка", "Потрібно вказати шлях до штатки та папки з файлами.", parent=root)
            return
        selected["confirmed"] = True
        root.destroy()

    def on_cancel() -> None:
        selected["confirmed"] = False
        root.destroy()

    root = tk.Tk()
    root.withdraw()
    install_frozen_executable_icon(root)
    root.title("НАЛАШТУВАННЯ - SyncOS")
    root.attributes("-topmost", True)
    root.resizable(False, False)
    root.protocol("WM_DELETE_WINDOW", on_cancel)
    colors = configure_launch_styles(root)
    root.configure(bg=colors["window"])
    install_dark_title_bar(root)

    xlsx_var = tk.StringVar(value="")
    folder_var = tk.StringVar(value="")
    xlsx_display_var = tk.StringVar(value="")
    folder_display_var = tk.StringVar(value="")
    update_var = tk.BooleanVar(value=False)

    container = ttk.Frame(root, style="LaunchRoot.TFrame")
    container.grid(row=0, column=0, sticky="nsew")
    container.grid_columnconfigure(0, weight=1)

    header = tk.Frame(container, bg=colors["header"], padx=22, pady=12)
    header.grid(row=0, column=0, sticky="we")
    header.grid_columnconfigure(0, weight=1)
    tk.Label(
        header,
        text="AuditOS",
        bg=colors["header"],
        fg="#FFFFFF",
        font=("Segoe UI Semibold", 16),
        anchor="w",
    ).grid(row=0, column=0, sticky="w")
    tk.Label(
        header,
        text="Підготовка аналізу DOCX / XLSX / XLSM за даними зі штатки...",
        bg=colors["header"],
        fg="#D7FBF5",
        font=("Segoe UI", 9),
        anchor="w",
    ).grid(row=1, column=0, sticky="w", pady=(2, 6))
    badge_shell = tk.Frame(header, bg=colors["accent_dark"], padx=1, pady=1)
    badge = tk.Label(
        badge_shell,
        text="НАЛАШТУВАННЯ",
        bg="#053D39",
        fg="#FFFFFF",
        font=("Segoe UI", 12, "bold"),
        padx=10,
        pady=5,
    )
    badge.pack()
    badge_shell.place(relx=1.0, x=5, y=5, anchor="ne")

    body = ttk.Frame(container, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
    body.grid(row=1, column=0, sticky="nsew")
    body.grid_columnconfigure(0, weight=1)
    panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
    panel_shell.grid(row=0, column=0, sticky="we")
    settings = ttk.Frame(panel_shell, style="LaunchPanel.TFrame", padding=(16, 14, 16, 12))
    settings.pack(fill="both", expand=True)
    settings.grid_columnconfigure(1, weight=1)

    ttk.Label(settings, text="ДЖЕРЕЛА", style="LaunchSection.TLabel").grid(
        row=0, column=0, columnspan=3, sticky="w"
    )
    ttk.Label(
        settings,
        text="Вкажіть штатний файл, папку для обробки та режим роботи.",
        style="LaunchMuted.TLabel",
        wraplength=560,
    ).grid(row=1, column=0, columnspan=3, pady=(4, 10), sticky="w")

    ttk.Label(settings, text="ШТАТКА:", style="LaunchField.TLabel").grid(
        row=2, column=0, padx=(0, 8), pady=7, sticky="e"
    )
    xlsx_entry = ttk.Entry(settings, textvariable=xlsx_display_var, width=48, state="readonly", style="Launch.TEntry")
    xlsx_entry.grid(row=2, column=1, padx=(0, 15), pady=7, sticky="we")
    ttk.Button(
        settings,
        text="ОБРАТИ",
        width=13,
        style="LaunchBrowse.TButton",
        command=lambda: choose_file(root, xlsx_var, xlsx_display_var, xlsx_entry, [("Excel", "*.xlsx *.xlsm"), ("All files", "*.*")]),
    ).grid(row=2, column=2, pady=7, sticky="nsew")

    ttk.Label(settings, text="ПАПКА:", style="LaunchField.TLabel").grid(
        row=3, column=0, padx=(0, 8), pady=7, sticky="e"
    )
    folder_entry = ttk.Entry(settings, textvariable=folder_display_var, width=48, state="readonly", style="Launch.TEntry")
    folder_entry.grid(row=3, column=1, padx=(0, 15), pady=7, sticky="we")
    ttk.Button(
        settings,
        text="ОБРАТИ",
        width=13,
        style="LaunchBrowse.TButton",
        command=lambda: choose_folder(root, folder_var, folder_display_var, folder_entry),
    ).grid(row=3, column=2, pady=7, sticky="nsew")

    ttk.Separator(settings, orient="horizontal").grid(
        row=4, column=0, columnspan=3, sticky="we", pady=(10, 12)
    )
    update_row = ttk.Frame(settings, style="LaunchPanel.TFrame")
    update_row.grid(row=5, column=0, columnspan=3, sticky="w")
    ttk.Label(update_row, text="Оновлення даних", style="LaunchField.TLabel").grid(
        row=0, column=0, sticky="w"
    )
    ToggleSwitch(update_row, update_var).grid(row=0, column=1, padx=(10, 0), sticky="w")

    ttk.Label(
        settings,
        text="Коли вимкнено, скрипт лише читає файли й формує стандартний LOG-звіт...",
        style="LaunchMuted.TLabel",
        wraplength=560,
        justify="left",
    ).grid(row=6, column=0, columnspan=3, pady=(10, 0), sticky="w")

    footer = ttk.Frame(body, style="LaunchRoot.TFrame")
    footer.grid(row=1, column=0, pady=(10, 0), sticky="we")
    footer.grid_columnconfigure(0, weight=1)
    start_button = ttk.Button(
        footer,
        text="ЗАПУСТИТИ",
        style="LaunchPrimary.TButton",
        command=on_start,
        padding=(18, 18),
    )
    start_button.grid(row=0, column=0, sticky="we")
    bind_primary_button_hover(start_button)

    xlsx_var.trace_add("write", lambda *_args: refresh_path_entry_display(xlsx_var, xlsx_display_var, xlsx_entry))
    folder_var.trace_add("write", lambda *_args: refresh_path_entry_display(folder_var, folder_display_var, folder_entry))
    xlsx_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(xlsx_var, xlsx_display_var, xlsx_entry))
    folder_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(folder_var, folder_display_var, folder_entry))

    root.update_idletasks()
    root.geometry(f"{root.winfo_reqwidth()}x{root.winfo_reqheight()}")
    root.deiconify()
    root.lift()
    root.grab_set()
    root.focus_force()
    root.mainloop()

    if not bool(selected["confirmed"]):
        return None

    xlsx_path = Path(selected["xlsx"]).expanduser().resolve()
    docx_folder = Path(selected["folder"]).expanduser().resolve()
    update_data = bool(selected["update_data"])

    if not xlsx_path.exists() or not docx_folder.exists():
        raise FileNotFoundError("Один або обидва шляхи не існують.")

    return xlsx_path, docx_folder, update_data


# ============================================================
# ОСНОВНИЙ ПРОЦЕС
# ============================================================


def setup_logging(work_folder: Path, timestamp: str) -> Path:
    global CURRENT_LOG_PATH
    log_path = work_folder / build_debug_log_filename(timestamp)
    logging.basicConfig(
        level=logging.INFO,
        format="[%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(),
        ],
        force=True,
    )
    CURRENT_LOG_PATH = log_path
    return log_path



def process_all(
    xlsx_path: Path,
    docx_folder: Path,
    update_data: bool = True,
    progress: ProgressWindow | None = None,
) -> Tuple[List[FileReport], Path]:
    run_timestamp = make_run_timestamp()
    setup_logging(docx_folder, run_timestamp)
    total_stages = 3 if update_data else 2
    prepare_stage = 2 if update_data else 1
    process_stage = 3 if update_data else 2
    prepare_header = build_stage_header(prepare_stage, total_stages, "Підготовка")
    process_header = build_stage_header(process_stage, total_stages, "Оновлення файлів" if update_data else "Аналіз файлів")

    logging.info("Режим оновлення даних: %s", "увімкнено" if update_data else "вимкнено (лише звіт)")
    if progress:
        progress.update(
            header=prepare_header,
            detail="Завантаження штатки...",
            current=0,
            total=1,
            file_name=xlsx_path.name,
        )
    logging.info("Завантаження штатки: %s", xlsx_path)
    staff_map = load_staff_records(xlsx_path)
    logging.info("Записів у штатці: %s", len(staff_map))
    if progress:
        progress.update(
            header=prepare_header,
            detail="Штатку завантажено",
            current=1,
            total=1,
            file_name=xlsx_path.name,
        )

    target_files = find_target_files(docx_folder, excluded_paths=[xlsx_path])
    if not target_files:
        raise FileNotFoundError("У вказаній папці не знайдено файлів .docx, .xlsx або .xlsm для обробки")

    logging.info("Файлів для обробки знайдено: %s", len(target_files))

    excel_targets = [path for path in target_files if path.suffix.lower() in {".xlsx", ".xlsm"}]
    if update_data and excel_targets and win32 is None:
        raise RuntimeError(
            "Для безпечного оновлення XLSX/XLSM потрібен встановлений pywin32 (win32com.client). "
            "Зараз він недоступний, тому можна або вимкнути 'Оновлення даних' для режиму лише читання, "
            "або спершу встановити pywin32."
        )

    if update_data:
        backup_name = create_folder_backup_in_trash(
            docx_folder,
            run_timestamp,
            progress=progress,
            stage_header=build_stage_header(1, total_stages, "Резервна копія"),
        )
        logging.info("Створено backup-копію папки та переміщено в кошик: %s", backup_name)
    else:
        logging.info("Резервну копію пропущено: режим лише читання без запису в цільові файли")

    file_reports: List[FileReport] = []
    if progress:
        progress.update(
            header=process_header,
            detail="Оновлення файлів..." if update_data else "Аналіз файлів без запису...",
            current=0,
            total=len(target_files),
            file_name="",
        )
    for idx, target_file in enumerate(target_files, start=1):
        logging.info("%s: %s", "Оновлення" if update_data else "Аналіз", target_file)
        try:
            report = process_target_file(target_file, staff_map, update_data=update_data)
            file_reports.append(report)
        except Exception as exc:
            logging.exception("Помилка при обробці '%s': %s", target_file.name, exc)
        if progress:
            progress.update(
                header=process_header,
                detail="Оновлення файлів..." if update_data else "Аналіз файлів без запису...",
                current=idx,
                total=len(target_files),
                file_name=target_file.name,
            )

    if progress:
        progress.update(
            header="Фіналізація",
            detail="Формування звіту...",
            current=len(target_files),
            total=len(target_files),
            file_name=build_report_filename(run_timestamp),
        )

    report_path = write_report(docx_folder, file_reports, run_timestamp, xlsx_path.name, docx_folder)
    logging.info("Звіт сформовано: %s", report_path)

    send_debug_log_to_trash()

    return file_reports, report_path



def main() -> None:
    ui_root: Any | None = None
    progress: Optional[ProgressWindow] = None
    try:
        selection = run_gui()
        if selection is None:
            return
        xlsx_path, docx_folder, update_data = selection
        ui_root = tk.Tk()
        ui_root.withdraw()
        install_frozen_executable_icon(ui_root)
        install_dark_title_bar(ui_root)
        ui_root.attributes("-topmost", True)
        progress = ProgressWindow(ui_root)
        _, report_path = run_process_all_with_progress(
            xlsx_path,
            docx_folder,
            update_data,
            progress,
        )
        done_detail = (
            "Оновлення даних завершено."
            if update_data
            else "Звіт сформовано без запису даних у файли."
        )
        progress.show_success_then_close(done_detail, file_name=f"Звіт: {report_path.name}")
        progress = None
        print(f"{done_detail}\nЗвіт: {report_path}")
    except Exception as exc:
        logging.exception("Критична помилка: %s", exc)
        try:
            message_parent = ui_root
            if progress is not None and _widget_exists(progress.dialog):
                message_parent = progress.dialog
            _show_topmost_message("error", "Помилка", str(exc), parent=message_parent)
        except Exception:
            print(f"Помилка: {exc}")
    finally:
        if progress is not None:
            progress.close()
        if ui_root is not None:
            _destroy_widget(ui_root)
        send_debug_log_to_trash()


if __name__ == "__main__":
    main()
