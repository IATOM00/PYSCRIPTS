from __future__ import annotations

from typing import Any, Dict, Iterable, List, Optional, Tuple
import sys, re, time, shutil, queue, threading
from dataclasses import dataclass, field
from openpyxl.styles import PatternFill
from docx.table import _Cell, Table
from datetime import date, datetime
from send2trash import send2trash
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from pathlib import Path
import openpyxl

try:
    import tkinter as tk
    import tkinter.font as tkfont
    from tkinter import filedialog, messagebox, ttk
except ImportError:
    tk = None
    tkfont = None
    filedialog = None
    messagebox = None
    ttk = None

try:
    import win32com.client as win32
except ImportError:
    win32 = None

try:
    import pythoncom
except ImportError:
    pythoncom = None


KEY_HEADER = "Ключ"
KEY_HEADER_NORM = "ключ"
SUPPORTED_EXTENSIONS = {".docx", ".xlsx", ".xlsm"}
EXCEL_EXTENSIONS = {".xlsx", ".xlsm"}
ALL_EXCEL_SHEETS_DISPLAY = "Усі аркуші"
DOCX_CHANGED_FILL = "FFFF00"
DOCX_TARGET_FONT = "Times New Roman"
EXCEL_CHANGED_FILL = "FFFF00"
EXCEL_COM_CHANGED_FILL = 65535
EXCEL_COM_SOLID_FILL_PATTERN = 1


def install_frozen_executable_icon(root: Any, retry_ms: int = 250) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    _set_frozen_executable_icon(root)
    try:
        root.after(retry_ms, lambda: _set_frozen_executable_icon(root))
    except Exception:
        pass


def _set_frozen_executable_icon(root: Any) -> None:
    if sys.platform != "win32" or not getattr(sys, "frozen", False):
        return

    try:
        import ctypes
        from ctypes import wintypes

        wm_seticon = 0x0080
        icon_small = 0
        icon_big = 1
        gclp_hicon = -14
        gclp_hiconsm = -34

        shell32 = ctypes.WinDLL("shell32", use_last_error=True)
        user32 = ctypes.WinDLL("user32", use_last_error=True)
        long_ptr = ctypes.c_longlong if ctypes.sizeof(ctypes.c_void_p) == 8 else ctypes.c_long

        shell32.ExtractIconExW.argtypes = (
            wintypes.LPCWSTR,
            ctypes.c_int,
            ctypes.POINTER(wintypes.HICON),
            ctypes.POINTER(wintypes.HICON),
            wintypes.UINT,
        )
        shell32.ExtractIconExW.restype = wintypes.UINT
        user32.SendMessageW.argtypes = (
            wintypes.HWND,
            wintypes.UINT,
            wintypes.WPARAM,
            wintypes.LPARAM,
        )
        user32.SendMessageW.restype = long_ptr
        user32.GetParent.argtypes = (wintypes.HWND,)
        user32.GetParent.restype = wintypes.HWND

        try:
            set_class_icon = user32.SetClassLongPtrW
        except AttributeError:
            set_class_icon = user32.SetClassLongW
        set_class_icon.argtypes = (wintypes.HWND, ctypes.c_int, long_ptr)
        set_class_icon.restype = long_ptr

        icon_handles = getattr(root, "_frozen_executable_icon_handles", None)
        if icon_handles is None:
            large_icons = (wintypes.HICON * 1)()
            small_icons = (wintypes.HICON * 1)()
            icon_count = shell32.ExtractIconExW(sys.executable, 0, large_icons, small_icons, 1)
            if icon_count <= 0:
                return
            icon_handles = (small_icons[0], large_icons[0])
            root._frozen_executable_icon_handles = icon_handles
        small_icon, large_icon = icon_handles

        root.update_idletasks()
        hwnds = []
        for raw_hwnd in (root.winfo_id(), root.wm_frame()):
            try:
                hwnd = int(str(raw_hwnd), 0)
            except (TypeError, ValueError):
                continue
            while hwnd and hwnd not in hwnds:
                hwnds.append(hwnd)
                hwnd = user32.GetParent(hwnd)

        for hwnd in hwnds:
            if small_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_small, small_icon)
                set_class_icon(hwnd, gclp_hiconsm, small_icon)
            if large_icon:
                user32.SendMessageW(hwnd, wm_seticon, icon_big, large_icon)
                set_class_icon(hwnd, gclp_hicon, large_icon)
    except Exception:
        pass


def install_dark_title_bar(window: Any, retry_ms: int = 80) -> None:
    if sys.platform != "win32":
        return

    def _apply() -> None:
        try:
            import ctypes

            hwnd = ctypes.windll.user32.GetParent(window.winfo_id())
            value = ctypes.c_int(1)
            for attr in (20, 19):
                ctypes.windll.dwmapi.DwmSetWindowAttribute(
                    hwnd,
                    attr,
                    ctypes.byref(value),
                    ctypes.sizeof(value),
                )
        except Exception:
            pass

    try:
        window.update_idletasks()
        _apply()
        window.after(retry_ms, _apply)
    except Exception:
        pass


@dataclass
class SourceRecord:
    display_key: str
    fields: Dict[str, Any] = field(default_factory=dict)
    field_labels: Dict[str, str] = field(default_factory=dict)


@dataclass
class SourceIndex:
    blocks_count: int = 0
    records: Dict[str, SourceRecord] = field(default_factory=dict)
    headers: Dict[str, str] = field(default_factory=dict)
    duplicate_key_rows: int = 0
    conflicting_values: int = 0


@dataclass
class DocxBlock:
    table: Table
    header_row_idx: int
    key_col: int
    headers: Dict[str, int]
    labels: Dict[str, str]
    name: str


@dataclass
class ExcelBlock:
    sheet_name: str
    header_row: int
    key_col: int
    max_row: int
    headers: Dict[str, int]
    labels: Dict[str, str]
    name: str


@dataclass(frozen=True)
class ExcelUpdate:
    sheet_name: str
    row: int
    col: int
    header: str
    new_value: Any


@dataclass
class SyncSummary:
    source_path: Path
    target_path: Path
    highlight_changes: bool
    source_sheet_name: Optional[str] = None
    target_sheet_name: Optional[str] = None
    source_blocks: int = 0
    target_blocks: int = 0
    source_records: int = 0
    matched_rows: int = 0
    missing_keys: int = 0
    updated_cells: int = 0
    unchanged_cells: int = 0
    target_saved: bool = False
    backup_name: str = ""
    duplicate_source_rows: int = 0
    source_conflicts: int = 0
    synced_headers: set[str] = field(default_factory=set)


class SyncError(RuntimeError):
    pass


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, float) and value.is_integer():
        text = str(int(value))
    else:
        text = str(value)
    text = text.replace("\xa0", " ").strip()
    return re.sub(r"\s+", " ", text)


def normalize_header(value: object) -> str:
    text = normalize_text(value).lower()
    text = text.replace("'", "").replace("`", "").replace("’", "")
    return re.sub(r"\s+", " ", text).strip()


def normalize_key(value: object) -> str:
    text = normalize_text(value).lower()
    text = text.replace("`", "'").replace("’", "'")
    return re.sub(r"\s+", " ", text).strip()


def normalize_for_change_detection(value: object) -> str:
    return normalize_key(value)


def value_changed(old_value: object, new_value: object) -> bool:
    return normalize_for_change_detection(old_value) != normalize_for_change_detection(new_value)


def value_to_docx_text(value: object) -> str:
    return normalize_text(value)


def ensure_supported_file(path: Path, role_name: str) -> None:
    if not path.exists():
        raise FileNotFoundError(f"{role_name}: файл не знайдено: {path}")
    if not path.is_file():
        raise SyncError(f"{role_name}: потрібно вказати файл, а не папку: {path}")
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise SyncError(
            f"{role_name}: підтримуються лише файли .docx, .xlsx, .xlsm. "
            f"Отримано: {path.suffix or '(без розширення)'}"
        )


def is_excel_file(path: Path) -> bool:
    return path.suffix.lower() in EXCEL_EXTENSIONS


def normalize_optional_sheet_name(sheet_name: Optional[str]) -> Optional[str]:
    if sheet_name is None:
        return None
    sheet_name = str(sheet_name).strip()
    return sheet_name or None


def same_excel_sheet_name(first: str, second: str) -> bool:
    return first.casefold() == second.casefold()


def make_backup_timestamp() -> str:
    return datetime.now().strftime("%Y-%m-%d %H.%M")


def build_unique_backup_path(path: Path) -> Path:
    timestamp = make_backup_timestamp()
    base_name = f"{path.stem} - {timestamp}"
    backup_path = path.with_name(f"{base_name}{path.suffix}")
    counter = 2

    while backup_path.exists():
        backup_path = path.with_name(f"{base_name} ({counter}){path.suffix}")
        counter += 1

    return backup_path


def create_file_backup_in_trash(path: Path) -> str:
    backup_path = build_unique_backup_path(path)
    shutil.copy2(path, backup_path)
    send2trash(str(backup_path))
    return backup_path.name


def build_header_map(pairs: Iterable[Tuple[int, object]]) -> Tuple[Dict[str, int], Dict[str, str]]:
    headers: Dict[str, int] = {}
    labels: Dict[str, str] = {}

    for index, raw_value in pairs:
        label = normalize_text(raw_value)
        normalized = normalize_header(label)
        if not normalized:
            continue
        if normalized not in headers:
            headers[normalized] = index
            labels[normalized] = label

    return headers, labels


def add_source_field(
    source_index: SourceIndex,
    record: SourceRecord,
    header: str,
    label: str,
    value: Any,
) -> None:
    if header == KEY_HEADER_NORM:
        return

    source_index.headers.setdefault(header, label)
    record.field_labels.setdefault(header, label)

    if header not in record.fields:
        record.fields[header] = value
        return

    current_value = record.fields[header]
    if not normalize_text(current_value) and normalize_text(value):
        record.fields[header] = value
        return

    if value_changed(current_value, value):
        source_index.conflicting_values += 1


def add_source_row(
    source_index: SourceIndex,
    key_value: Any,
    row_values: Dict[str, Tuple[str, Any]],
) -> None:
    display_key = normalize_text(key_value)
    normalized_key = normalize_key(key_value)
    if not normalized_key:
        return

    record = source_index.records.get(normalized_key)
    if record is None:
        record = SourceRecord(display_key=display_key)
        source_index.records[normalized_key] = record
    else:
        source_index.duplicate_key_rows += 1

    for header, (label, value) in row_values.items():
        add_source_field(source_index, record, header, label, value)


def iter_tables_recursive(document: Any) -> Iterable[Table]:
    for table in document.tables:
        yield table
        yield from iter_nested_tables(table)


def iter_nested_tables(table: Table) -> Iterable[Table]:
    for row in table.rows:
        for cell in row.cells:
            for nested in cell.tables:
                yield nested
                yield from iter_nested_tables(nested)


def find_docx_blocks(document: Any) -> List[DocxBlock]:
    blocks: List[DocxBlock] = []

    for table_idx, table in enumerate(iter_tables_recursive(document), start=1):
        rows = list(table.rows)
        for row_idx, row in enumerate(rows):
            headers, labels = build_header_map(
                (col_idx, cell.text) for col_idx, cell in enumerate(row.cells)
            )
            if KEY_HEADER_NORM not in headers:
                continue

            seen_cells: set[int] = set()
            for col_idx, cell in enumerate(row.cells):
                if normalize_header(cell.text) != KEY_HEADER_NORM:
                    continue
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                blocks.append(
                    DocxBlock(
                        table=table,
                        header_row_idx=row_idx,
                        key_col=col_idx,
                        headers=headers,
                        labels=labels,
                        name=f"таблиця {table_idx}, рядок {row_idx + 1}",
                    )
                )

    return blocks


def read_docx_source(path: Path, source_index: SourceIndex) -> None:
    document = Document(path)
    blocks = find_docx_blocks(document)
    source_index.blocks_count += len(blocks)

    for block in blocks:
        rows = list(block.table.rows)
        for row in rows[block.header_row_idx + 1 :]:
            if block.key_col >= len(row.cells):
                continue

            key_value = row.cells[block.key_col].text
            row_values: Dict[str, Tuple[str, Any]] = {}

            for header, col_idx in block.headers.items():
                if header == KEY_HEADER_NORM or col_idx >= len(row.cells):
                    continue
                row_values[header] = (block.labels.get(header, header), normalize_text(row.cells[col_idx].text))

            add_source_row(source_index, key_value, row_values)


def find_merged_range(worksheet: Any, row: int, col: int) -> Any | None:
    for merged_range in worksheet.merged_cells.ranges:
        if (
            merged_range.min_row <= row <= merged_range.max_row
            and merged_range.min_col <= col <= merged_range.max_col
        ):
            return merged_range
    return None


def get_excel_cell_value(worksheet: Any, row: int, col: int) -> Any:
    cell = worksheet.cell(row=row, column=col)
    if cell.value is not None:
        return cell.value

    merged_range = find_merged_range(worksheet, row, col)
    if merged_range is None:
        return None
    return worksheet.cell(row=merged_range.min_row, column=merged_range.min_col).value


def build_excel_header_map(worksheet: Any, row_idx: int) -> Tuple[Dict[str, int], Dict[str, str]]:
    return build_header_map(
        (col_idx, get_excel_cell_value(worksheet, row_idx, col_idx))
        for col_idx in range(1, worksheet.max_column + 1)
    )


def iter_excel_worksheets(workbook: Any, sheet_name: Optional[str] = None) -> Iterable[Any]:
    if sheet_name is None:
        yield from workbook.worksheets
        return

    if sheet_name not in workbook.sheetnames:
        raise SyncError(f"В Excel-книзі не знайдено аркуш: {sheet_name}")

    yield workbook[sheet_name]


def find_excel_blocks(workbook: Any, sheet_name: Optional[str] = None) -> List[ExcelBlock]:
    blocks: List[ExcelBlock] = []

    for worksheet in iter_excel_worksheets(workbook, sheet_name):
        seen_key_cells: set[Tuple[int, int]] = set()
        for row_idx in range(1, worksheet.max_row + 1):
            for col_idx in range(1, worksheet.max_column + 1):
                if normalize_header(get_excel_cell_value(worksheet, row_idx, col_idx)) != KEY_HEADER_NORM:
                    continue

                merged_range = find_merged_range(worksheet, row_idx, col_idx)
                key_row = merged_range.min_row if merged_range is not None else row_idx
                key_col = merged_range.min_col if merged_range is not None else col_idx
                if (key_row, key_col) in seen_key_cells:
                    continue
                seen_key_cells.add((key_row, key_col))

                headers, labels = build_excel_header_map(worksheet, key_row)
                blocks.append(
                    ExcelBlock(
                        sheet_name=worksheet.title,
                        header_row=key_row,
                        key_col=key_col,
                        max_row=worksheet.max_row,
                        headers=headers,
                        labels=labels,
                        name=f"{worksheet.title}!R{key_row}C{key_col}",
                    )
                )

    return blocks


def load_excel_workbook_for_read(path: Path) -> Any:
    return openpyxl.load_workbook(
        path,
        data_only=True,
        keep_vba=path.suffix.lower() == ".xlsm",
    )


def get_excel_sheet_names(path: Path) -> List[str]:
    workbook = load_excel_workbook_for_read(path)
    try:
        return list(workbook.sheetnames)
    finally:
        try:
            workbook.close()
        except Exception:
            pass


def read_excel_source(
    path: Path,
    source_index: SourceIndex,
    sheet_name: Optional[str] = None,
) -> None:
    workbook = load_excel_workbook_for_read(path)
    try:
        blocks = find_excel_blocks(workbook, sheet_name)
        source_index.blocks_count += len(blocks)

        for block in blocks:
            worksheet = workbook[block.sheet_name]
            for row_idx in range(block.header_row + 1, block.max_row + 1):
                key_value = get_excel_cell_value(worksheet, row_idx, block.key_col)
                if not normalize_key(key_value):
                    continue

                row_values: Dict[str, Tuple[str, Any]] = {}
                for header, col_idx in block.headers.items():
                    if header == KEY_HEADER_NORM:
                        continue
                    row_values[header] = (
                        block.labels.get(header, header),
                        get_excel_cell_value(worksheet, row_idx, col_idx),
                    )

                add_source_row(source_index, key_value, row_values)
    finally:
        try:
            workbook.close()
        except Exception:
            pass


def build_source_index(
    source_path: Path,
    source_sheet_name: Optional[str] = None,
) -> SourceIndex:
    source_index = SourceIndex()
    suffix = source_path.suffix.lower()

    if suffix == ".docx":
        read_docx_source(source_path, source_index)
    elif suffix in EXCEL_EXTENSIONS:
        read_excel_source(source_path, source_index, source_sheet_name)
    else:
        raise SyncError(f"Непідтримуваний тип джерела: {source_path.suffix}")

    return source_index


def common_headers(target_headers: Dict[str, int], source_index: SourceIndex) -> List[str]:
    return [
        header
        for header in target_headers
        if header != KEY_HEADER_NORM and header in source_index.headers
    ]


def remember_synced_headers(
    summary: SyncSummary,
    headers: Iterable[str],
    source_index: SourceIndex,
    target_labels: Dict[str, str],
) -> None:
    for header in headers:
        label = target_labels.get(header) or source_index.headers.get(header) or header
        summary.synced_headers.add(label)


def write_docx_plain_text_only(cell: _Cell, value: str) -> None:
    text_nodes = list(cell._tc.iter(qn("w:t")))
    if text_nodes:
        text_nodes[0].text = value
        for text_node in text_nodes[1:]:
            text_node.text = ""
        apply_docx_font_to_cell_runs(cell, DOCX_TARGET_FONT)
        return

    if cell.paragraphs:
        run = cell.paragraphs[0].add_run(value)
        apply_docx_font_to_run(run._r, DOCX_TARGET_FONT)
    else:
        cell.text = value
        apply_docx_font_to_cell_runs(cell, DOCX_TARGET_FONT)


def apply_docx_font_to_cell_runs(cell: _Cell, font_name: str) -> None:
    for run in cell._tc.iter(qn("w:r")):
        apply_docx_font_to_run(run, font_name)


def apply_docx_font_to_run(run_element: Any, font_name: str) -> None:
    run_properties = run_element.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_element.insert(0, run_properties)

    run_fonts = run_properties.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)

    for attr_name in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{attr_name}"), font_name)


def apply_docx_cell_fill(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def sync_docx_target(
    target_path: Path,
    source_index: SourceIndex,
    summary: SyncSummary,
    highlight_changes: bool,
) -> None:
    document = Document(target_path)
    blocks = find_docx_blocks(document)
    summary.target_blocks = len(blocks)
    has_changes = False

    for block in blocks:
        headers_to_sync = common_headers(block.headers, source_index)
        if not headers_to_sync:
            continue
        remember_synced_headers(summary, headers_to_sync, source_index, block.labels)

        rows = list(block.table.rows)
        for row in rows[block.header_row_idx + 1 :]:
            if block.key_col >= len(row.cells):
                continue

            key = normalize_key(row.cells[block.key_col].text)
            if not key:
                continue

            record = source_index.records.get(key)
            if record is None:
                summary.missing_keys += 1
                continue

            summary.matched_rows += 1
            for header in headers_to_sync:
                if header not in record.fields:
                    continue
                col_idx = block.headers[header]
                if col_idx >= len(row.cells):
                    continue

                target_cell = row.cells[col_idx]
                new_value = value_to_docx_text(record.fields[header])
                if value_changed(target_cell.text, new_value):
                    write_docx_plain_text_only(target_cell, new_value)
                    if highlight_changes:
                        apply_docx_cell_fill(target_cell, DOCX_CHANGED_FILL)
                    summary.updated_cells += 1
                    has_changes = True
                else:
                    summary.unchanged_cells += 1

    if has_changes:
        document.save(target_path)
        summary.target_saved = True


def collect_excel_updates(
    target_path: Path,
    source_index: SourceIndex,
    summary: SyncSummary,
    target_sheet_name: Optional[str] = None,
) -> List[ExcelUpdate]:
    workbook = load_excel_workbook_for_read(target_path)
    updates: List[ExcelUpdate] = []
    seen_cells: set[Tuple[str, int, int]] = set()

    try:
        blocks = find_excel_blocks(workbook, target_sheet_name)
        summary.target_blocks = len(blocks)

        for block in blocks:
            headers_to_sync = common_headers(block.headers, source_index)
            if not headers_to_sync:
                continue
            remember_synced_headers(summary, headers_to_sync, source_index, block.labels)

            worksheet = workbook[block.sheet_name]
            for row_idx in range(block.header_row + 1, block.max_row + 1):
                key = normalize_key(get_excel_cell_value(worksheet, row_idx, block.key_col))
                if not key:
                    continue

                record = source_index.records.get(key)
                if record is None:
                    summary.missing_keys += 1
                    continue

                summary.matched_rows += 1
                for header in headers_to_sync:
                    if header not in record.fields:
                        continue

                    col_idx = block.headers[header]
                    cell_key = (block.sheet_name, row_idx, col_idx)
                    if cell_key in seen_cells:
                        continue
                    seen_cells.add(cell_key)

                    new_value = record.fields[header]
                    old_value = get_excel_cell_value(worksheet, row_idx, col_idx)
                    if not value_changed(old_value, new_value):
                        summary.unchanged_cells += 1
                        continue

                    updates.append(
                        ExcelUpdate(
                            sheet_name=block.sheet_name,
                            row=row_idx,
                            col=col_idx,
                            header=header,
                            new_value=new_value,
                        )
                    )
    finally:
        try:
            workbook.close()
        except Exception:
            pass

    return updates


def apply_excel_updates(
    target_path: Path,
    updates: List[ExcelUpdate],
    highlight_changes: bool,
) -> Tuple[int, int]:
    if not updates:
        return 0, 0

    if win32 is not None:
        try:
            return apply_excel_updates_via_com(target_path, updates, highlight_changes)
        except Exception:
            pass

    return apply_excel_updates_via_openpyxl(target_path, updates, highlight_changes)


def apply_excel_updates_via_com(
    target_path: Path,
    updates: List[ExcelUpdate],
    highlight_changes: bool,
) -> Tuple[int, int]:
    excel = None
    workbook = None
    changed_count = 0
    unchanged_count = 0
    com_initialized = False

    try:
        if pythoncom is not None:
            pythoncom.CoInitialize()
            com_initialized = True

        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        try:
            excel.AskToUpdateLinks = False
        except Exception:
            pass

        workbook = excel.Workbooks.Open(str(target_path), UpdateLinks=0, ReadOnly=False)
        worksheets: Dict[str, Any] = {}

        for update in updates:
            worksheet = worksheets.get(update.sheet_name)
            if worksheet is None:
                worksheet = workbook.Worksheets(update.sheet_name)
                worksheets[update.sheet_name] = worksheet

            cell = worksheet.Cells(update.row, update.col)
            target_range = cell
            target_cell = cell
            try:
                if bool(cell.MergeCells):
                    target_range = cell.MergeArea
                    target_cell = target_range.Cells(1, 1)
            except Exception:
                pass

            old_value = target_cell.Value
            if not value_changed(old_value, update.new_value):
                unchanged_count += 1
                continue

            target_cell.Value = update.new_value
            if highlight_changes:
                target_range.Interior.Pattern = EXCEL_COM_SOLID_FILL_PATTERN
                target_range.Interior.Color = EXCEL_COM_CHANGED_FILL
            changed_count += 1

        if changed_count:
            workbook.Save()

        return changed_count, unchanged_count
    finally:
        if workbook is not None:
            try:
                workbook.Close(SaveChanges=False)
            except Exception:
                pass
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        if com_initialized:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass


def excel_write_anchor(worksheet: Any, row: int, col: int) -> Tuple[int, int, Any | None]:
    merged_range = find_merged_range(worksheet, row, col)
    if merged_range is None:
        return row, col, None
    return merged_range.min_row, merged_range.min_col, merged_range


def apply_openpyxl_fill(worksheet: Any, row: int, col: int, merged_range: Any | None) -> None:
    fill = PatternFill(fill_type="solid", fgColor=EXCEL_CHANGED_FILL)
    if merged_range is None:
        worksheet.cell(row=row, column=col).fill = fill
        return

    for row_cells in worksheet.iter_rows(
        min_row=merged_range.min_row,
        max_row=merged_range.max_row,
        min_col=merged_range.min_col,
        max_col=merged_range.max_col,
    ):
        for cell in row_cells:
            try:
                cell.fill = fill
            except Exception:
                pass


def apply_excel_updates_via_openpyxl(
    target_path: Path,
    updates: List[ExcelUpdate],
    highlight_changes: bool,
) -> Tuple[int, int]:
    workbook = openpyxl.load_workbook(
        target_path,
        keep_vba=target_path.suffix.lower() == ".xlsm",
    )
    changed_count = 0
    unchanged_count = 0

    try:
        for update in updates:
            worksheet = workbook[update.sheet_name]
            row, col, merged_range = excel_write_anchor(worksheet, update.row, update.col)
            cell = worksheet.cell(row=row, column=col)

            if not value_changed(cell.value, update.new_value):
                unchanged_count += 1
                continue

            cell.value = update.new_value
            if highlight_changes:
                apply_openpyxl_fill(worksheet, row, col, merged_range)
            changed_count += 1

        if changed_count:
            workbook.save(target_path)

        return changed_count, unchanged_count
    finally:
        try:
            workbook.close()
        except Exception:
            pass


def sync_excel_target(
    target_path: Path,
    source_index: SourceIndex,
    summary: SyncSummary,
    highlight_changes: bool,
    target_sheet_name: Optional[str] = None,
) -> None:
    updates = collect_excel_updates(target_path, source_index, summary, target_sheet_name)
    changed_count, unchanged_count = apply_excel_updates(target_path, updates, highlight_changes)
    summary.updated_cells += changed_count
    summary.unchanged_cells += unchanged_count
    summary.target_saved = changed_count > 0


def validate_sheet_selection(
    source_path: Path,
    target_path: Path,
    source_sheet_name: Optional[str],
    target_sheet_name: Optional[str],
) -> None:
    if source_sheet_name and not is_excel_file(source_path):
        raise SyncError("Аркуш джерела можна вибирати лише для Excel-файлу.")
    if target_sheet_name and not is_excel_file(target_path):
        raise SyncError("Аркуш цілі можна вибирати лише для Excel-файлу.")

    if source_path != target_path:
        return

    if not is_excel_file(source_path):
        raise SyncError("Джерело і ціль можуть збігатися лише для синхронізації аркушів однієї Excel-книги.")
    if not source_sheet_name or not target_sheet_name:
        raise SyncError(
            "Для синхронізації в межах однієї Excel-книги потрібно вибрати конкретний аркуш джерела та аркуш цілі."
        )
    if same_excel_sheet_name(source_sheet_name, target_sheet_name):
        raise SyncError("Для однієї Excel-книги аркуш джерела і аркуш цілі мають бути різними.")


def format_excel_endpoint_name(path: Path, sheet_name: Optional[str]) -> str:
    if sheet_name:
        return f"{path.name} / {sheet_name}"
    return path.name


def sync_files(
    source_path: Path,
    target_path: Path,
    highlight_changes: bool = False,
    source_sheet_name: Optional[str] = None,
    target_sheet_name: Optional[str] = None,
    progress: Optional["ProgressWindow"] = None,
) -> SyncSummary:
    source_path = source_path.expanduser().resolve()
    target_path = target_path.expanduser().resolve()
    source_sheet_name = normalize_optional_sheet_name(source_sheet_name)
    target_sheet_name = normalize_optional_sheet_name(target_sheet_name)

    ensure_supported_file(source_path, "ДЖЕРЕЛО")
    ensure_supported_file(target_path, "ЦІЛЬ")
    validate_sheet_selection(source_path, target_path, source_sheet_name, target_sheet_name)

    summary = SyncSummary(
        source_path=source_path,
        target_path=target_path,
        highlight_changes=highlight_changes,
        source_sheet_name=source_sheet_name,
        target_sheet_name=target_sheet_name,
    )

    if progress:
        progress.update(
            header="Етап 1/3: Резервна копія",
            detail="Створення копії цілі та переміщення у кошик...",
            current=0,
            total=1,
            file_name=target_path.name,
        )

    summary.backup_name = create_file_backup_in_trash(target_path)

    if progress:
        progress.update(
            header="Етап 1/3: Резервна копія",
            detail="Резервну копію створено",
            current=1,
            total=1,
            file_name=summary.backup_name,
        )
        progress.update(
            header="Етап 2/3: Джерело",
            detail="Пошук заголовка 'Ключ' і читання даних...",
            current=0,
            total=1,
            file_name=format_excel_endpoint_name(source_path, source_sheet_name),
        )

    source_index = build_source_index(source_path, source_sheet_name)
    summary.source_blocks = source_index.blocks_count
    summary.source_records = len(source_index.records)
    summary.duplicate_source_rows = source_index.duplicate_key_rows
    summary.source_conflicts = source_index.conflicting_values

    if source_index.blocks_count == 0:
        if source_sheet_name:
            raise SyncError(f"У джерелі на аркуші '{source_sheet_name}' не знайдено заголовок 'Ключ'.")
        raise SyncError("У джерелі не знайдено заголовок 'Ключ'.")
    if not source_index.records:
        raise SyncError("У джерелі знайдено 'Ключ', але під ним немає даних для синхронізації.")

    if progress:
        progress.update(
            header="Етап 2/3: Джерело",
            detail="Дані джерела прочитано",
            current=1,
            total=1,
            file_name=f"Ключів: {summary.source_records}",
        )
        progress.update(
            header="Етап 3/3: Ціль",
            detail="Пошук спільних заголовків і оновлення цілі...",
            current=0,
            total=1,
            file_name=format_excel_endpoint_name(target_path, target_sheet_name),
        )

    suffix = target_path.suffix.lower()
    if suffix == ".docx":
        sync_docx_target(target_path, source_index, summary, highlight_changes)
    elif suffix in EXCEL_EXTENSIONS:
        sync_excel_target(target_path, source_index, summary, highlight_changes, target_sheet_name)
    else:
        raise SyncError(f"Непідтримуваний тип цілі: {target_path.suffix}")

    if summary.target_blocks == 0:
        if target_sheet_name:
            raise SyncError(f"У цілі на аркуші '{target_sheet_name}' не знайдено заголовок 'Ключ'.")
        raise SyncError("У цілі не знайдено заголовок 'Ключ'.")
    if not summary.synced_headers:
        raise SyncError(
            "Не знайдено спільних заголовків для оновлення. "
            "Назви колонок у джерелі та цілі мають збігатися."
        )

    if progress:
        progress.update(
            header="Етап 3/3: Ціль",
            detail="Синхронізацію завершено",
            current=1,
            total=1,
            file_name=f"Оновлено клітинок: {summary.updated_cells}",
        )

    return summary


def format_header_list(headers: Iterable[str], limit: int = 8) -> str:
    ordered = sorted(headers, key=lambda value: normalize_header(value))
    if not ordered:
        return "-"
    visible = ordered[:limit]
    text = ", ".join(visible)
    if len(ordered) > limit:
        text += f" та ще {len(ordered) - limit}"
    return text


def build_summary_message(summary: SyncSummary) -> str:
    lines = [
        "Синхронізацію завершено.",
        f"Оновлено клітинок: {summary.updated_cells}",
        f"Рядків цілі з ключем у джерелі: {summary.matched_rows}",
        f"Ключів цілі без збігу в джерелі: {summary.missing_keys}",
        f"Спільні заголовки: {format_header_list(summary.synced_headers)}",
    ]

    if summary.source_sheet_name:
        lines.append(f"Аркуш джерела: {summary.source_sheet_name}")
    if summary.target_sheet_name:
        lines.append(f"Аркуш цілі: {summary.target_sheet_name}")
    if summary.highlight_changes and summary.updated_cells:
        lines.append("Змінені дані замальовано жовтим.")
    if summary.backup_name:
        lines.append(f"Backup-копію цілі переміщено в кошик: {summary.backup_name}")
    if not summary.target_saved:
        lines.append("Файл цілі не перезаписувався, бо фактичних змін не було.")
    if summary.duplicate_source_rows:
        lines.append(f"Дублі рядків у джерелі за ключем: {summary.duplicate_source_rows}")
    if summary.source_conflicts:
        lines.append(f"Конфліктних повторів значень у джерелі: {summary.source_conflicts}")

    return "\n".join(lines)


def configure_launch_styles(root: Any) -> Dict[str, str]:
    colors = {
        "window": "#F4F7FB",
        "panel": "#FFFFFF",
        "border": "#D6DEE8",
        "text": "#15202B",
        "muted": "#617084",
        "accent": "#0F766E",
        "accent_dark": "#0B5D56",
        "header": "#103C3A",
        "header_badge": "#0F766E",
    }

    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except Exception:
        pass

    try:
        style.configure("LaunchRoot.TFrame", background=colors["window"])
        style.configure("LaunchPanel.TFrame", background=colors["panel"])
        style.configure("LaunchSection.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        style.configure("LaunchField.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10, "bold"))
        style.configure("LaunchMuted.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("LaunchTiny.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 8))
        style.configure("LaunchStatus.TLabel", background=colors["window"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("ProgressHeader.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 12, "bold"))
        style.configure("ProgressBody.TLabel", background=colors["panel"], foreground=colors["text"], font=("Segoe UI", 10))
        style.configure("ProgressFile.TLabel", background=colors["panel"], foreground=colors["muted"], font=("Segoe UI", 9))
        style.configure("ProgressCount.TLabel", background=colors["panel"], foreground=colors["accent_dark"], font=("Segoe UI", 10, "bold"))
        style.configure(
            "Launch.Horizontal.TProgressbar",
            background=colors["header_badge"],
            troughcolor="#E8EEF5",
            bordercolor=colors["border"],
            lightcolor=colors["header_badge"],
            darkcolor=colors["header_badge"],
            thickness=12,
        )
        style.configure(
            "Launch.TEntry",
            fieldbackground="#FFFFFF",
            foreground=colors["text"],
            bordercolor=colors["border"],
            lightcolor=colors["border"],
            darkcolor=colors["border"],
            padding=(8, 6),
        )
        style.map(
            "Launch.TEntry",
            fieldbackground=[("readonly", "#FFFFFF")],
            foreground=[("readonly", colors["text"])],
        )
        style.configure(
            "LaunchBrowse.TButton",
            font=("Segoe UI", 10, "bold"),
            padding=(18, 10),
            foreground="#FFFFFF",
            background="#053D39",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchBrowse.TButton",
            background=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF")],
            bordercolor=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", colors["accent_dark"]), ("active", colors["accent_dark"])],
        )
        style.configure(
            "LaunchPrimary.TButton",
            font=("Segoe UI", 11, "bold"),
            padding=(20, 14),
            foreground="#FFFFFF",
            background="#053D39",
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.configure(
            "LaunchPrimaryHover.TButton",
            font=("Segoe UI", 11, "bold"),
            padding=(20, 14),
            foreground="#FFFFFF",
            background=colors["accent_dark"],
            bordercolor=colors["accent_dark"],
            lightcolor=colors["accent_dark"],
            darkcolor=colors["accent_dark"],
            focuscolor=colors["accent_dark"],
        )
        style.map(
            "LaunchPrimary.TButton",
            background=[("pressed", "#042F2C"), ("active", colors["accent_dark"]), ("disabled", "#B8C2CC")],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF"), ("disabled", "#EEF2F7")],
            bordercolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
        )
        style.map(
            "LaunchPrimaryHover.TButton",
            background=[("pressed", "#042F2C"), ("active", colors["accent_dark"]), ("disabled", "#B8C2CC")],
            foreground=[("pressed", "#FFFFFF"), ("active", "#FFFFFF"), ("disabled", "#EEF2F7")],
            bordercolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            lightcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
            darkcolor=[("pressed", "#042F2C"), ("active", colors["accent_dark"])],
        )
    except Exception:
        pass

    return colors


class ProgressWindow:
    def __init__(self, owner: Any):
        if tk is None or ttk is None:
            raise RuntimeError("Tkinter недоступний.")

        self.close_requested = False
        self.dialog = tk.Toplevel(owner)
        self.dialog.withdraw()
        install_frozen_executable_icon(self.dialog)
        self.dialog.title("SyncEW - Обробка")
        self.dialog.resizable(False, False)
        self.dialog.attributes("-topmost", True)
        self.dialog.protocol("WM_DELETE_WINDOW", self.request_close)
        colors = configure_launch_styles(self.dialog)
        self.dialog.configure(bg=colors["window"])
        install_dark_title_bar(self.dialog)

        container = ttk.Frame(self.dialog, style="LaunchRoot.TFrame")
        container.pack(fill="both", expand=True)

        header = tk.Frame(container, bg=colors["header"], padx=22, pady=10)
        header.pack(fill="x")
        header.grid_columnconfigure(0, weight=1)
        tk.Label(
            header,
            text="SyncEW",
            bg=colors["header"],
            fg="#FFFFFF",
            font=("Segoe UI Semibold", 15),
            anchor="w",
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            header,
            text="Синхронізація файлів виконується. Будь ласка, зачекайте.",
            bg=colors["header"],
            fg="#D7FBF5",
            font=("Segoe UI", 9),
            anchor="w",
        ).grid(row=1, column=0, sticky="w", pady=(3, 4))
        badge_shell = tk.Frame(header, bg=colors["accent_dark"], padx=1, pady=1)
        badge = tk.Label(
            badge_shell,
            text="ОБРОБКА",
            bg="#053D39",
            fg="#FFFFFF",
            font=("Segoe UI", 12, "bold"),
            padx=10,
            pady=5,
        )
        badge.pack()
        badge_shell.place(relx=1.0, x=5, y=5, anchor="ne")

        body = ttk.Frame(container, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
        body.pack(fill="both", expand=True)
        panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
        panel_shell.pack(fill="both", expand=True)
        frame = ttk.Frame(panel_shell, style="LaunchPanel.TFrame", padding=(16, 14, 16, 12))
        frame.pack(fill="both", expand=True)
        frame.grid_columnconfigure(0, weight=1)

        self.header_var = tk.StringVar(value="Підготовка...")
        self.detail_var = tk.StringVar(value="Будь ласка, зачекайте")
        self.file_var = tk.StringVar(value="")

        ttk.Label(frame, textvariable=self.header_var, style="ProgressHeader.TLabel").grid(
            row=0, column=0, sticky="w"
        )
        ttk.Label(frame, textvariable=self.detail_var, style="ProgressBody.TLabel", wraplength=500).grid(
            row=1, column=0, pady=(8, 6), sticky="w"
        )
        ttk.Label(frame, textvariable=self.file_var, style="ProgressFile.TLabel", wraplength=500).grid(
            row=2, column=0, pady=(0, 12), sticky="w"
        )
        self.progress = ttk.Progressbar(
            frame,
            orient="horizontal",
            mode="indeterminate",
            maximum=100,
            style="Launch.Horizontal.TProgressbar",
        )
        self.progress.grid(row=3, column=0, sticky="we", pady=(4, 0))
        self.progress.start(12)
        self.dialog.update_idletasks()
        width = max(560, self.dialog.winfo_reqwidth())
        height = max(248, self.dialog.winfo_reqheight())
        self.dialog.geometry(f"{width}x{height}")
        self.dialog.deiconify()
        self.dialog.lift()
        self.refresh()

    def update(self, header=None, detail=None, current=None, total=None, file_name=None):
        try:
            if not self.dialog.winfo_exists():
                return
            if header is not None:
                self.header_var.set(header)
            if detail is not None:
                self.detail_var.set(detail)
            if file_name is not None:
                self.file_var.set(file_name)
            self.refresh()
        except tk.TclError:
            return

    def refresh(self):
        try:
            if self.close_requested or not _widget_exists(self.dialog):
                return
            self.dialog.update_idletasks()
            self.dialog.update()
        except Exception:
            pass

    def close(self):
        self.close_requested = True
        try:
            self.progress.stop()
        except Exception:
            pass
        try:
            if _widget_exists(self.dialog):
                self.dialog.destroy()
        except Exception:
            pass

    def request_close(self):
        self.close()

    def show_success_then_close(self, detail: str, file_name: str = "", delay_ms: int = 900):
        if self.close_requested or not _widget_exists(self.dialog):
            return
        try:
            self.progress.stop()
        except Exception:
            pass
        self.update(
            header="Готово",
            detail=detail,
            current=1,
            total=1,
            file_name=file_name,
        )
        deadline = time.monotonic() + delay_ms / 1000.0
        while time.monotonic() < deadline:
            if not _widget_exists(self.dialog):
                return
            self.refresh()
            time.sleep(0.05)
        self.close()


ToggleSwitchBase = ttk.Frame if ttk is not None else object


class ToggleSwitch(ToggleSwitchBase):
    def __init__(self, owner: Any, variable: Any):
        super().__init__(owner)
        self.variable = variable
        self.state_var = tk.StringVar()

        style = ttk.Style(self)
        try:
            parent_style = owner.cget("style")
            if parent_style:
                self.configure(style=parent_style)
                canvas_bg = style.lookup(parent_style, "background")
            else:
                canvas_bg = ""
        except Exception:
            canvas_bg = ""
        canvas_bg = canvas_bg or style.lookup("LaunchPanel.TFrame", "background") or self.winfo_toplevel().cget("bg")
        state_label_style = "ToggleState.TLabel"
        try:
            style.configure(
                state_label_style,
                background="#F4F4F8",
                foreground="#15202B",
                font=("Segoe UI", 9),
            )
        except Exception:
            state_label_style = ""

        self.canvas = tk.Canvas(
            self,
            width=52,
            height=28,
            bd=0,
            highlightthickness=0,
            bg=canvas_bg,
            cursor="hand2",
        )
        self.canvas.grid(row=0, column=0, padx=(0, 8), sticky="w")

        self.state_label = ttk.Label(self, textvariable=self.state_var, style=state_label_style)
        self.state_label.grid(row=0, column=1, sticky="w")

        for widget in (self, self.canvas, self.state_label):
            widget.bind("<Button-1>", self.toggle)

        self.variable.trace_add("write", self._redraw)
        self._redraw()

    def toggle(self, _event=None) -> None:
        self.variable.set(not self.variable.get())

    def _redraw(self, *_args) -> None:
        is_enabled = bool(self.variable.get())
        track_color = "#0F766E" if is_enabled else "#CBD5E1"
        knob_color = "#FFFFFF"
        self.state_var.set("Увімкнено" if is_enabled else "Вимкнено")

        self.canvas.delete("all")

        left = 2
        top = 2
        right = 50
        bottom = 26
        radius = (bottom - top) // 2

        self.canvas.create_oval(left, top, left + 2 * radius, bottom, fill=track_color, outline=track_color)
        self.canvas.create_oval(right - 2 * radius, top, right, bottom, fill=track_color, outline=track_color)
        self.canvas.create_rectangle(left + radius, top, right - radius, bottom, fill=track_color, outline=track_color)

        knob_diameter = (2 * radius) - 2
        knob_left = right - knob_diameter - 1 if is_enabled else left + 1
        self.canvas.create_oval(
            knob_left,
            top + 1,
            knob_left + knob_diameter,
            bottom - 1,
            fill=knob_color,
            outline=knob_color,
        )


def _widget_exists(widget: Any | None) -> bool:
    if tk is None or widget is None:
        return False
    try:
        return bool(widget.winfo_exists())
    except (tk.TclError, RuntimeError):
        return False


def _destroy_widget(widget: Any | None) -> None:
    if not _widget_exists(widget):
        return
    try:
        widget.destroy()
    except Exception:
        pass


def _prepare_dialog_parent(root: Any) -> None:
    if not _widget_exists(root):
        return
    try:
        try:
            current_grab = root.grab_current()
            if current_grab is not None and current_grab is not root:
                current_grab.grab_release()
        except Exception:
            pass
        if hasattr(root, "deiconify") and not root.winfo_ismapped():
            root.deiconify()
        root.lift()
        root.attributes("-topmost", True)
        root.update()
    except Exception:
        pass


def _show_topmost_message(
    kind: str,
    title: str,
    text: str,
    parent: Any | None = None,
    reveal_parent: bool = True,
) -> None:
    if tk is None or messagebox is None:
        stream = sys.stderr if kind == "error" else sys.stdout
        print(f"{title}: {text}", file=stream)
        return

    if _widget_exists(parent):
        if reveal_parent:
            _prepare_dialog_parent(parent)
        else:
            try:
                current_grab = parent.grab_current()
                if current_grab is not None and current_grab is not parent:
                    current_grab.grab_release()
            except Exception:
                pass
            try:
                parent.attributes("-topmost", True)
                parent.update_idletasks()
            except Exception:
                pass
        if kind == "info":
            messagebox.showinfo(title, text, parent=parent)
        elif kind == "warning":
            messagebox.showwarning(title, text, parent=parent)
        else:
            messagebox.showerror(title, text, parent=parent)
        return

    root = tk.Tk()
    root.withdraw()
    install_frozen_executable_icon(root)
    install_dark_title_bar(root)
    root.attributes("-topmost", True)
    try:
        if kind == "info":
            messagebox.showinfo(title, text, parent=root)
        elif kind == "warning":
            messagebox.showwarning(title, text, parent=root)
        else:
            messagebox.showerror(title, text, parent=root)
    finally:
        _destroy_widget(root)


def format_path_for_entry_display(path: str) -> str:
    path = path.strip()
    if not path:
        return ""

    normalized = Path(path)
    drive_prefix = ""
    if normalized.drive and len(normalized.drive) == 2 and normalized.drive[1] == ":":
        drive_prefix = f"{normalized.drive}/"

    name = normalized.name
    parent_name = normalized.parent.name
    if parent_name and name:
        display_path = f"{drive_prefix}.../{parent_name}/{name}"
    elif name:
        display_path = f"{drive_prefix}{name}" if drive_prefix else f".../{name}"
    elif drive_prefix:
        display_path = drive_prefix
    else:
        display_path = path.replace("\\", "/")

    return display_path


def fit_path_for_entry_display(path: str, entry: Any) -> str:
    display_path = format_path_for_entry_display(path)
    if not display_path:
        return ""

    entry.update_idletasks()
    font_name = entry.cget("font") or "TkDefaultFont"
    try:
        font = tkfont.nametofont(font_name)
    except tk.TclError:
        try:
            font = tkfont.nametofont("TkDefaultFont")
        except tk.TclError:
            return display_path

    available_width = max(120, entry.winfo_width() - 14)
    if font.measure(display_path) <= available_width:
        return display_path

    parts = display_path.split("/")
    if len(parts) >= 3:
        prefix = "/".join(parts[:-1]) + "/..."
        shortened_name = parts[-1]
        while len(shortened_name) > 1 and font.measure(prefix + shortened_name) > available_width:
            shortened_name = shortened_name[1:]
        return prefix + shortened_name

    ellipsis = "..."
    shortened_path = display_path
    while len(shortened_path) > 1 and font.measure(ellipsis + shortened_path) > available_width:
        shortened_path = shortened_path[1:]
    return ellipsis + shortened_path


def refresh_path_entry_display(
    source_var: Any,
    display_var: Any,
    entry: Any,
    sheet_var: Any | None = None,
) -> None:
    if not _widget_exists(entry):
        return
    display_var.set(fit_path_for_entry_display(source_var.get(), entry))


def choose_file(
    root: Any,
    source_var: Any,
    display_var: Any,
    entry: Any,
    sheet_var: Any | None = None,
    after_select: Any | None = None,
) -> None:
    _prepare_dialog_parent(root)
    path = filedialog.askopenfilename(
        filetypes=[
            ("Word / Excel", "*.docx *.xlsx *.xlsm"),
            ("Word", "*.docx"),
            ("Excel", "*.xlsx *.xlsm"),
            ("All files", "*.*"),
        ],
        parent=root,
    )
    if path:
        source_var.set(path)
        if after_select is not None:
            after_select(path)
        refresh_path_entry_display(source_var, display_var, entry, sheet_var)
    _prepare_dialog_parent(root)


def bind_primary_button_hover(button: Any) -> None:
    def on_enter(_event) -> None:
        button.configure(style="LaunchPrimaryHover.TButton")
        try:
            button.focus_force()
        except Exception:
            pass

    def on_leave(_event) -> None:
        button.configure(style="LaunchPrimary.TButton")
        try:
            button.winfo_toplevel().focus_force()
        except Exception:
            pass

    button.bind("<Enter>", on_enter, add="+")
    button.bind("<Leave>", on_leave, add="+")


def run_gui() -> Optional[Tuple[Path, Path, bool, Optional[str], Optional[str]]]:
    if tk is None or ttk is None or filedialog is None:
        raise RuntimeError("Tkinter недоступний: неможливо показати вікно налаштувань.")

    selected: Dict[str, object] = {
        "source": "",
        "target": "",
        "source_sheet": None,
        "target_sheet": None,
        "highlight": False,
        "confirmed": False,
    }

    def on_start() -> None:
        selected["source"] = source_var.get().strip()
        selected["target"] = target_var.get().strip()
        selected["source_sheet"] = normalize_optional_sheet_name(source_sheet_var.get())
        selected["target_sheet"] = normalize_optional_sheet_name(target_sheet_var.get())
        selected["highlight"] = bool(highlight_var.get())
        if not selected["source"] or not selected["target"]:
            _show_topmost_message(
                "error",
                "Помилка",
                "Потрібно вказати шлях до джерела та шлях до цілі.",
                parent=root,
            )
            return

        try:
            source_path = Path(str(selected["source"])).expanduser().resolve()
            target_path = Path(str(selected["target"])).expanduser().resolve()
        except Exception as exc:
            _show_topmost_message("error", "Помилка", str(exc), parent=root)
            return

        if source_path == target_path:
            if not is_excel_file(source_path):
                _show_topmost_message(
                    "error",
                    "Помилка",
                    "Однаковий шлях дозволений лише для синхронізації аркушів однієї Excel-книги.",
                    parent=root,
                )
                return
            if not selected["source_sheet"] or not selected["target_sheet"]:
                _show_topmost_message(
                    "error",
                    "Помилка",
                    "Для однієї Excel-книги оберіть конкретний аркуш джерела та конкретний аркуш цілі.",
                    parent=root,
                )
                return
            if same_excel_sheet_name(str(selected["source_sheet"]), str(selected["target_sheet"])):
                _show_topmost_message(
                    "error",
                    "Помилка",
                    "Для однієї Excel-книги аркуш джерела і аркуш цілі мають бути різними.",
                    parent=root,
                )
                return

        selected["confirmed"] = True
        root.destroy()

    def on_cancel() -> None:
        selected["confirmed"] = False
        root.destroy()

    root = tk.Tk()
    root.withdraw()
    install_frozen_executable_icon(root)
    root.title("SyncEW - НАЛАШТУВАННЯ:")
    root.attributes("-topmost", True)
    root.resizable(False, False)
    root.protocol("WM_DELETE_WINDOW", on_cancel)
    colors = configure_launch_styles(root)
    root.configure(bg=colors["window"])
    install_dark_title_bar(root)

    source_var = tk.StringVar(value="")
    target_var = tk.StringVar(value="")
    source_display_var = tk.StringVar(value="")
    target_display_var = tk.StringVar(value="")
    source_sheet_var = tk.StringVar(value="")
    target_sheet_var = tk.StringVar(value="")
    highlight_var = tk.BooleanVar(value=False)
    sheet_names_cache: Dict[str, List[str]] = {}
    sheet_names_by_slot: Dict[str, List[str]] = {"source": [], "target": []}
    sheet_menu_cache: Dict[Tuple[str, str], Any] = {}

    def invalidate_sheet_menu_cache(slot_prefix: str) -> None:
        keys_to_remove = [key for key in sheet_menu_cache if key[0] == slot_prefix]
        for key in keys_to_remove:
            menu = sheet_menu_cache.pop(key)
            try:
                menu.destroy()
            except Exception:
                pass

    def set_sheet_button_enabled(button: Any, enabled: bool) -> None:
        if not _widget_exists(button):
            return
        try:
            button.configure(
                state="normal" if enabled else "disabled",
                fg=colors["text"] if enabled else colors["muted"],
                disabledforeground=colors["muted"],
            )
        except Exception:
            pass

    def load_sheet_names_for_path(path: str, file_label: str) -> Optional[List[str]]:
        path = path.strip()
        if not path or Path(path).suffix.lower() not in EXCEL_EXTENSIONS:
            return []

        cached = sheet_names_cache.get(path)
        if cached is not None:
            return list(cached)

        try:
            sheet_names = get_excel_sheet_names(Path(path))
        except Exception as exc:
            _show_topmost_message(
                "error",
                "Помилка",
                f"Не вдалося прочитати список аркушів у {file_label}:\n{exc}",
                parent=root,
            )
            return None

        if not sheet_names:
            _show_topmost_message(
                "warning",
                "Помилка",
                f"У {file_label} не знайдено жодного аркуша.",
                parent=root,
            )
            return None

        sheet_names_cache[path] = list(sheet_names)
        return sheet_names

    def on_excel_file_selected(
        slot_prefix: str,
        path: str,
        sheet_var: Any,
        sheet_button: Any,
        file_label: str,
    ) -> None:
        sheet_var.set("")
        sheet_names_by_slot[slot_prefix] = []
        invalidate_sheet_menu_cache(slot_prefix)

        if Path(path).suffix.lower() not in EXCEL_EXTENSIONS:
            set_sheet_button_enabled(sheet_button, False)
            return

        sheet_names = load_sheet_names_for_path(path, file_label)
        if sheet_names:
            sheet_names_by_slot[slot_prefix] = sheet_names
            set_sheet_button_enabled(sheet_button, True)
        else:
            set_sheet_button_enabled(sheet_button, False)

    def choose_sheet(
        slot_prefix: str,
        file_var: Any,
        sheet_var: Any,
        anchor_button: Any,
        file_label: str,
    ) -> None:
        path = file_var.get().strip()
        if not path or not Path(path).exists():
            _show_topmost_message(
                "warning",
                "Перевірка",
                f"Спершу оберіть коректний файл для {file_label}.",
                parent=root,
            )
            return
        if Path(path).suffix.lower() not in EXCEL_EXTENSIONS:
            _show_topmost_message(
                "warning",
                "Перевірка",
                f"Аркуші доступні лише для Excel-файлу у {file_label}.",
                parent=root,
            )
            return

        sheet_names = sheet_names_by_slot.get(slot_prefix) or load_sheet_names_for_path(path, file_label)
        if not sheet_names:
            return
        sheet_names_by_slot[slot_prefix] = sheet_names

        selected_sheet = sheet_var.get().strip()
        menu_items = [("", ALL_EXCEL_SHEETS_DISPLAY)] + [(sheet_name, sheet_name) for sheet_name in sheet_names]
        cache_key = (slot_prefix, path)
        menu = sheet_menu_cache.get(cache_key)

        if menu is None:
            menu = tk.Menu(root, tearoff=0)

            def add_sheet_option(value: str, label: str) -> None:
                menu.add_command(label=label, command=lambda sheet=value: sheet_var.set(sheet))

            for sheet_value, sheet_label in menu_items:
                add_sheet_option(sheet_value, sheet_label)
            sheet_menu_cache[cache_key] = menu

        for index, (sheet_value, sheet_label) in enumerate(menu_items):
            visible_label = f"✓ {sheet_label}" if sheet_value == selected_sheet else sheet_label
            menu.entryconfigure(index, label=visible_label)

        root.update_idletasks()
        x = anchor_button.winfo_rootx()
        y = anchor_button.winfo_rooty() + anchor_button.winfo_height()
        try:
            menu.tk_popup(x, y)
        finally:
            menu.grab_release()

    def bind_sheet_button_hover(button: Any) -> None:
        def on_enter(_event=None) -> None:
            if str(button.cget("state")) != "disabled":
                button.configure(bg="#E8EEF5")

        def on_leave(_event=None) -> None:
            button.configure(bg="#FFFFFF")

        button.bind("<Enter>", on_enter, add="+")
        button.bind("<Leave>", on_leave, add="+")

    def build_file_field(
        parent: Any,
        row: int,
        display_var: Any,
        sheet_command: Any,
    ) -> Tuple[Any, Any]:
        field_shell = tk.Frame(parent, bg=colors["border"], bd=0, padx=1, pady=1)
        field_shell.grid(row=row, column=1, padx=(0, 15), pady=7, sticky="we")
        field_shell.grid_columnconfigure(0, weight=1)

        field_panel = tk.Frame(field_shell, bg="#FFFFFF", bd=0)
        field_panel.grid(row=0, column=0, sticky="we")
        field_panel.grid_columnconfigure(0, weight=1)

        entry = tk.Entry(
            field_panel,
            bg="#FFFFFF",
            bd=0,
            fg=colors["text"],
            font=("Segoe UI", 10),
            highlightthickness=0,
            insertbackground=colors["text"],
            readonlybackground="#FFFFFF",
            relief="flat",
            state="readonly",
            textvariable=display_var,
            width=48,
        )
        entry.grid(row=0, column=0, padx=(8, 6), pady=6, sticky="we")

        separator = tk.Frame(field_panel, bg=colors["border"], width=1)
        separator.grid(row=0, column=1, sticky="ns")

        sheet_button = tk.Button(
            field_panel,
            activebackground="#FFFFFF",
            activeforeground=colors["text"],
            bg="#FFFFFF",
            bd=0,
            command=sheet_command,
            cursor="hand2",
            disabledforeground=colors["muted"],
            fg=colors["muted"],
            font=("Segoe UI", 9, "bold"),
            highlightthickness=0,
            padx=5,
            pady=4,
            relief="flat",
            state="disabled",
            text="Ар.",
            width=4,
        )
        sheet_button.grid(row=0, column=2, sticky="ns")
        bind_sheet_button_hover(sheet_button)
        return entry, sheet_button

    container = ttk.Frame(root, style="LaunchRoot.TFrame")
    container.grid(row=0, column=0, sticky="nsew")
    container.grid_columnconfigure(0, weight=1)

    header = tk.Frame(container, bg=colors["header"], padx=22, pady=12)
    header.grid(row=0, column=0, sticky="we")
    header.grid_columnconfigure(0, weight=1)
    tk.Label(
        header,
        text="SyncEW",
        bg=colors["header"],
        fg="#FFFFFF",
        font=("Segoe UI Semibold", 16),
        anchor="w",
    ).grid(row=0, column=0, sticky="w")
    tk.Label(
        header,
        text="Підготовка синхронізації DOCX / XLSX / XLSM...",
        bg=colors["header"],
        fg="#D7FBF5",
        font=("Segoe UI", 9),
        anchor="w",
    ).grid(row=1, column=0, sticky="w", pady=(2, 6))
    badge_shell = tk.Frame(header, bg=colors["accent_dark"], padx=1, pady=1)
    badge = tk.Label(
        badge_shell,
        text="НАЛАШТУВАННЯ",
        bg="#053D39",
        fg="#FFFFFF",
        font=("Segoe UI", 12, "bold"),
        padx=10,
        pady=5,
    )
    badge.pack()
    badge_shell.place(relx=1.0, x=5, y=5, anchor="ne")

    body = ttk.Frame(container, style="LaunchRoot.TFrame", padding=(18, 16, 18, 18))
    body.grid(row=1, column=0, sticky="nsew")
    body.grid_columnconfigure(0, weight=1)
    panel_shell = tk.Frame(body, bg=colors["border"], padx=1, pady=1)
    panel_shell.grid(row=0, column=0, sticky="we")
    settings = ttk.Frame(panel_shell, style="LaunchPanel.TFrame", padding=(16, 14, 16, 12))
    settings.pack(fill="both", expand=True)
    settings.grid_columnconfigure(1, weight=1)

    ttk.Label(settings, text="ФАЙЛИ", style="LaunchSection.TLabel").grid(
        row=0, column=0, columnspan=3, sticky="w"
    )
    ttk.Label(
        settings,
        text='У двох файлах має бути заголовок: "Ключ", та однакові заголовки згідно яких оновлюються дані...',
        style="LaunchMuted.TLabel",
        wraplength=560,
        justify="left",
    ).grid(row=1, column=0, columnspan=3, pady=(4, 10), sticky="w")

    ttk.Label(settings, text="ДЖЕРЕЛО:", style="LaunchField.TLabel").grid(
        row=2, column=0, padx=(0, 8), pady=7, sticky="e"
    )
    source_entry, source_sheet_button = build_file_field(
        settings,
        2,
        source_display_var,
        lambda: choose_sheet("source", source_var, source_sheet_var, source_sheet_button, "джерела"),
    )
    ttk.Button(
        settings,
        text="ОБРАТИ",
        width=13,
        style="LaunchBrowse.TButton",
        command=lambda: choose_file(
            root,
            source_var,
            source_display_var,
            source_entry,
            source_sheet_var,
            lambda path: on_excel_file_selected("source", path, source_sheet_var, source_sheet_button, "джерелі"),
        ),
    ).grid(row=2, column=2, pady=7, sticky="nsew")

    ttk.Label(settings, text="ЦІЛЬ:", style="LaunchField.TLabel").grid(
        row=3, column=0, padx=(0, 8), pady=7, sticky="e"
    )
    target_entry, target_sheet_button = build_file_field(
        settings,
        3,
        target_display_var,
        lambda: choose_sheet("target", target_var, target_sheet_var, target_sheet_button, "цілі"),
    )
    ttk.Button(
        settings,
        text="ОБРАТИ",
        width=13,
        style="LaunchBrowse.TButton",
        command=lambda: choose_file(
            root,
            target_var,
            target_display_var,
            target_entry,
            target_sheet_var,
            lambda path: on_excel_file_selected("target", path, target_sheet_var, target_sheet_button, "цілі"),
        ),
    ).grid(row=3, column=2, pady=7, sticky="nsew")

    ttk.Separator(settings, orient="horizontal").grid(
        row=4, column=0, columnspan=3, sticky="we", pady=(10, 12)
    )
    highlight_row = ttk.Frame(settings, style="LaunchPanel.TFrame")
    highlight_row.grid(row=5, column=0, columnspan=3, sticky="w")
    ttk.Label(highlight_row, text="Замальовування даних", style="LaunchField.TLabel").grid(
        row=0, column=0, sticky="w"
    )
    ToggleSwitch(highlight_row, highlight_var).grid(row=0, column=1, padx=(10, 0), sticky="w")

    ttk.Label(
        settings,
        text="Коли увімкнено, фактично змінені клітинки замальовуються жовтим...",
        style="LaunchMuted.TLabel",
        wraplength=560,
        justify="left",
    ).grid(row=6, column=0, columnspan=3, pady=(8, 0), sticky="w")

    footer = ttk.Frame(body, style="LaunchRoot.TFrame")
    footer.grid(row=1, column=0, pady=(10, 0), sticky="we")
    footer.grid_columnconfigure(0, weight=1)
    start_button = ttk.Button(
        footer,
        text="ЗАПУСТИТИ",
        style="LaunchPrimary.TButton",
        command=on_start,
        padding=(18, 18),
    )
    start_button.grid(row=0, column=0, sticky="we")
    bind_primary_button_hover(start_button)

    source_var.trace_add("write", lambda *_args: refresh_path_entry_display(source_var, source_display_var, source_entry, source_sheet_var))
    target_var.trace_add("write", lambda *_args: refresh_path_entry_display(target_var, target_display_var, target_entry, target_sheet_var))
    source_sheet_var.trace_add("write", lambda *_args: refresh_path_entry_display(source_var, source_display_var, source_entry, source_sheet_var))
    target_sheet_var.trace_add("write", lambda *_args: refresh_path_entry_display(target_var, target_display_var, target_entry, target_sheet_var))
    source_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(source_var, source_display_var, source_entry, source_sheet_var))
    target_entry.bind("<Configure>", lambda _event: refresh_path_entry_display(target_var, target_display_var, target_entry, target_sheet_var))

    root.update_idletasks()
    root.geometry(f"{root.winfo_reqwidth()}x{root.winfo_reqheight()}")
    root.deiconify()
    root.lift()
    root.grab_set()
    root.focus_force()
    root.mainloop()

    if not bool(selected["confirmed"]):
        return None

    return (
        Path(str(selected["source"])).expanduser().resolve(),
        Path(str(selected["target"])).expanduser().resolve(),
        bool(selected["highlight"]),
        normalize_optional_sheet_name(selected.get("source_sheet")),  # type: ignore[arg-type]
        normalize_optional_sheet_name(selected.get("target_sheet")),  # type: ignore[arg-type]
    )


class ProgressUpdateProxy:
    def __init__(self, updates: "queue.Queue[Dict[str, Any]]"):
        self.updates = updates

    def update(self, header=None, detail=None, current=None, total=None, file_name=None):
        self.updates.put(
            {
                "header": header,
                "detail": detail,
                "current": current,
                "total": total,
                "file_name": file_name,
            }
        )


def apply_pending_progress_updates(
    progress: ProgressWindow,
    updates: "queue.Queue[Dict[str, Any]]",
) -> None:
    while True:
        try:
            payload = updates.get_nowait()
        except queue.Empty:
            return
        progress.update(**payload)


def run_sync_with_progress(
    source_path: Path,
    target_path: Path,
    highlight_changes: bool,
    source_sheet_name: Optional[str],
    target_sheet_name: Optional[str],
    progress: ProgressWindow,
) -> SyncSummary:
    updates: "queue.Queue[Dict[str, Any]]" = queue.Queue()
    result: Dict[str, Any] = {"summary": None, "error": None}
    progress_proxy = ProgressUpdateProxy(updates)

    def worker() -> None:
        try:
            result["summary"] = sync_files(
                source_path,
                target_path,
                highlight_changes=highlight_changes,
                source_sheet_name=source_sheet_name,
                target_sheet_name=target_sheet_name,
                progress=progress_proxy,
            )
        except Exception as exc:
            result["error"] = exc

    thread = threading.Thread(target=worker, daemon=True)
    thread.start()

    while thread.is_alive():
        apply_pending_progress_updates(progress, updates)
        progress.refresh()
        time.sleep(0.03)

    thread.join()
    apply_pending_progress_updates(progress, updates)

    if result["error"] is not None:
        raise result["error"]
    if result["summary"] is None:
        raise RuntimeError("Синхронізація завершилася без результату.")
    return result["summary"]


def main() -> None:
    ui_root: Any | None = None
    progress: Optional[ProgressWindow] = None

    try:
        selection = run_gui()
        if selection is None:
            return

        source_path, target_path, highlight_changes, source_sheet_name, target_sheet_name = selection
        ui_root = tk.Tk()
        ui_root.withdraw()
        install_frozen_executable_icon(ui_root)
        install_dark_title_bar(ui_root)
        ui_root.attributes("-topmost", True)

        progress = ProgressWindow(ui_root)
        summary = run_sync_with_progress(
            source_path,
            target_path,
            highlight_changes,
            source_sheet_name,
            target_sheet_name,
            progress,
        )
        progress.show_success_then_close(
            "Синхронізацію завершено.",
            file_name=f"Ціль: {format_excel_endpoint_name(summary.target_path, summary.target_sheet_name)}",
        )
        progress = None

        message = build_summary_message(summary)
        print(message)
        _show_topmost_message("info", "Готово", message, parent=ui_root, reveal_parent=False)
    except Exception as exc:
        try:
            message_parent = ui_root
            if progress is not None and _widget_exists(progress.dialog):
                message_parent = progress.dialog
            _show_topmost_message("error", "Помилка", str(exc), parent=message_parent)
        except Exception:
            print(f"Помилка: {exc}", file=sys.stderr)
    finally:
        if progress is not None:
            progress.close()
        if ui_root is not None:
            _destroy_widget(ui_root)


if __name__ == "__main__":
    main()
